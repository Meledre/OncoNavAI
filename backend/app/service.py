from __future__ import annotations

import base64
import copy
import hashlib
import json
import re
import time
import uuid
from datetime import date, datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from typing import Any
import zipfile

from backend.app.casefacts.extractor import extract_case_facts
from backend.app.casefacts.extractor_v2 import extract_case_facts_v2
from backend.app.config import Settings, is_strict_release_profile, normalize_reasoning_mode
from backend.app.drugs.dictionary_loader import (
    DrugDictionaryBundle,
    load_drug_dictionary_bundle_from_path,
    load_drug_dictionary_bundle_from_text,
)
from backend.app.drugs.extractor import extract_drugs_and_regimens
from backend.app.drugs.models import (
    DoctorDrugSafety,
    DrugExtractedInn,
    DrugSafetySignal,
    DrugSafetyWarning,
    DrugUnresolvedCandidate,
    build_patient_drug_safety,
)
from backend.app.drugs.safety_provider import DrugSafetyProvider
from backend.app.exceptions import NotFoundError, ValidationError
from backend.app.icd10.infer import infer_icd10_code
from backend.app.icd10.reference_loader import parse_icd10_reference_entries_from_chunks
from backend.app.importers.text_extract import extract_text
from backend.app.guidelines.nosology_mapper import (
    apply_unknown_nosology_fallback,
    infer_cancer_type_for_guideline,
    is_nosology_mapped,
)
from backend.app.guidelines.source_registry import (
    DEFAULT_AUTO_SOURCE_IDS,
    OFFICIAL_SOURCE_RULES,
    evaluate_release_validity,
    is_pubmed_url,
    normalize_source_set_id,
    normalize_source_set_ids,
    resolve_primary_source_url,
    resolve_official_doc_hints,
)
from backend.app.guidelines.sync_minzdrav import (
    KNOWN_MINZDRAV_PDFS,
    download_minzdrav_pdf_with_url,
    resolve_minzdrav_pdf_from_page,
)
from backend.app.guidelines.sync_russco import discover_russco_2025_documents, download_russco_pdf
from backend.app.llm.generate_doctor_report import build_doctor_report_with_fallback, map_strict_to_public_report
from backend.app.llm.generate_patient_explain import (
    build_patient_explain_with_fallback,
    map_strict_to_pack_patient_v1_2,
    map_strict_to_public_patient,
)
from backend.app.llm.normalize_plan import normalize_plan
from backend.app.llm.query_bundle_planner import build_query_bundle_with_llm
from backend.app.llm.provider_router import LLMEndpoint, LLMProviderRouter
from backend.app.llm.prompt_registry import PromptRegistry
from backend.app.llm.route_planner import plan_nosology_route_with_llm
from backend.app.planning.next_steps import build_next_steps_plan_sections, flatten_plan_for_diff
from backend.app.rag.embedder import build_embedder
from backend.app.rag.engine import build_retriever
from backend.app.rag.ingest_pdf import extract_pdf_chunks, file_sha256
from backend.app.rag.kb_version import compute_kb_version
from backend.app.rag.qdrant_client import LocalQdrantLikeIndex, QdrantRestIndex
from backend.app.rag.query_bundle import build_query_bundle
from backend.app.rag.reranker import Reranker
from backend.app.reporting.citation_linker import attach_issue_citations, attach_plan_citations, build_citations_from_chunks
from backend.app.reporting.compat_doctor_v1_1 import project_doctor_report_v1_1, validate_doctor_projection_v1_1
from backend.app.reporting.compat_patient_projection import (
    project_patient_explain_alt_profile,
    validate_patient_projection_alt,
)
from backend.app.reporting.guided_report_builder import build_guided_report
from backend.app.reporting.patient_context_builder import build_patient_context_from_analyze_response
from backend.app.reporting.timeline_reconciler import reconcile_timeline_signals
from backend.app.reporting.v1_2_builder import (
    build_consilium_md,
    build_disease_context,
    build_run_meta,
    build_timeline,
)
from backend.app.routing.nosology_router import NosologyRouteDecision, resolve_nosology_route
from backend.app.rules.diff_engine import compute_diff
from backend.app.rules.data_sufficiency import evaluate_data_sufficiency
from backend.app.rules.drug_safety_rules import build_drug_safety_signals
from backend.app.rules.evidence_guard import enforce_retrieved_evidence
from backend.app.rules.gastric_rules import apply_gastric_rules
from backend.app.rules.min_case_requirements import evaluate_min_case_requirements
from backend.app.rules.ru_text_normalizer import normalize_ru_clinical_text
from backend.app.rules.sanity_checks import auto_repair_report, run_sanity_checks
from backend.app.schemas.contracts import (
    validate_analyze_request,
    validate_analyze_response,
    validate_doctor_report,
)
from backend.app.schemas.case_import import (
    SUPPORTED_CASE_DATA_MODES,
    SUPPORTED_CASE_IMPORT_PROFILES,
    normalize_case_import_payload,
)
from backend.app.schemas.analyze_bridge import (
    DIALECT_PACK_V2,
    is_pack_v0_2_request,
    normalize_analyze_request,
    serialize_analyze_response,
)
from backend.app.security.logging import get_logger, safe_log
from backend.app.security.pii_detector import contains_pii, redact_pii
from backend.app.security.rate_limit import RateLimiter
from backend.app.security.rbac import ensure_role
from backend.app.storage import DocRecord, SQLiteStore

_ICD10_PATTERN = re.compile(r"\b[CD]\d{2}(?:\.[0-9A-Z]{1,2})?\b", re.IGNORECASE)
_STAGE_PATTERN = re.compile(r"(?:стадия|stage)\s*[:\-]?\s*(IV|III|II|I)\b", re.IGNORECASE)
_YEAR_PATTERN = re.compile(r"(?<!\d)(19[0-9]{2}|20[0-9]{2})(?:\s*г(?:\.|од[ау]?)?)?(?!\d)", re.IGNORECASE)
_BIRTH_YEAR_PATTERN = re.compile(
    r"(?:дата\s+рождени\w*|д\.?\s*р\.?)\D{0,20}(?:\d{1,2}[./-]\d{1,2}[./-])?(19[0-9]{2}|20[0-9]{2})(?:\s*г(?:\.|од[ау]?)?)?",
    re.IGNORECASE,
)
_AGE_PATTERN = re.compile(r"(?:возраст|age)\s*[:\-]?\s*(\d{1,3})\b|(\d{1,3})\s*(?:лет|года|год)\b", re.IGNORECASE)
_LINE_PATTERN = re.compile(r"(?:line|линия)\s*[:#-]?\s*(\d{1,2})|(\d{1,2})\s*[- ]?(?:я|й)?\s*линия", re.IGNORECASE)
_CYCLE_PATTERN = re.compile(r"(?:cycle|курс)\s*[:#-]?\s*(\d{1,3})", re.IGNORECASE)
_HEIGHT_PATTERN = re.compile(r"(?:рост|height)\s*[:\-]?\s*(\d{2,3}(?:[.,]\d{1,2})?)\s*(?:см|cm)?\b", re.IGNORECASE)
_WEIGHT_PATTERN = re.compile(r"(?:вес|масса(?:\s+тела)?|weight)\s*[:\-]?\s*(\d{2,3}(?:[.,]\d{1,2})?)\s*(?:кг|kg)?\b", re.IGNORECASE)
_ECOG_PATTERN = re.compile(r"(?:ecog|эко?г)\s*[:\-]?\s*([0-5])\b", re.IGNORECASE)
_BIOMARKER_PATTERN = re.compile(
    r"\b(EGFR|HER2|PD[-\s]?L1(?:_CPS)?|MSI|BRCA1?|ALK|ROS1)\b\s*[:=]\s*([A-Za-z0-9+%._/-]+)",
    re.IGNORECASE,
)
_SUPPORT_TOKEN_PATTERN = re.compile(r"[A-Za-zА-Яа-яЁё0-9][A-Za-zА-Яа-яЁё0-9+._/-]{3,}")
_SUPPORT_STOPWORDS = {
    "требует",
    "оценка",
    "риска",
    "нужна",
    "отдельного",
    "консилиума",
    "with",
    "without",
    "from",
    "this",
    "that",
    "issue",
    "case",
    "critical",
    "warning",
    "important",
    "пациент",
    "пациента",
    "клинический",
    "клиническая",
    "данных",
    "следующего",
    "выбора",
    "лечения",
}


class OncoService:
    @staticmethod
    def _build_llm_endpoint(url: str, model: str, api_key: str) -> LLMEndpoint | None:
        normalized_url = str(url or "").strip()
        normalized_model = str(model or "").strip()
        normalized_key = str(api_key or "").strip()
        if not normalized_url or not normalized_model:
            return None
        # Skip unauthenticated public OpenAI endpoint to avoid unnecessary traffic.
        if "api.openai.com" in normalized_url.lower() and not normalized_key:
            return None
        return LLMEndpoint(normalized_url, normalized_model, normalized_key)

    @staticmethod
    def _strict_profile_token(settings: Settings) -> str:
        return "strict_full" if is_strict_release_profile(getattr(settings, "release_profile", "compat")) else "compat"

    @classmethod
    def _ensure_strict_release_requirements(
        cls,
        *,
        settings: Settings,
        primary_endpoint: LLMEndpoint | None,
    ) -> None:
        if cls._strict_profile_token(settings) != "strict_full":
            return
        failures: list[str] = []
        if not bool(getattr(settings, "llm_generation_enabled", False)):
            failures.append("LLM_GENERATION_ENABLED must be true")
        if str(getattr(settings, "vector_backend", "")).strip().lower() != "qdrant":
            failures.append("VECTOR_BACKEND must be qdrant")
        if not str(getattr(settings, "qdrant_url", "")).strip():
            failures.append("QDRANT_URL is required")
        if str(getattr(settings, "embedding_backend", "")).strip().lower() != "openai":
            failures.append("EMBEDDING_BACKEND must be openai")
        if not str(getattr(settings, "embedding_url", "")).strip():
            failures.append("EMBEDDING_URL is required")
        if not str(getattr(settings, "embedding_model", "")).strip():
            failures.append("EMBEDDING_MODEL is required")
        if not str(getattr(settings, "embedding_api_key", "")).strip():
            failures.append("EMBEDDING_API_KEY is required")
        if str(getattr(settings, "reranker_backend", "")).strip().lower() != "llm":
            failures.append("RERANKER_BACKEND must be llm")
        if normalize_reasoning_mode(str(getattr(settings, "reasoning_mode", ""))) != "llm_rag_only":
            failures.append("ONCOAI_REASONING_MODE must be llm_rag_only")
        if str(getattr(settings, "llm_fallback_url", "")).strip():
            failures.append("LLM_FALLBACK_URL must be empty in strict_full")
        if primary_endpoint is None:
            failures.append("LLM primary endpoint is required in strict_full")
        if failures:
            raise ValidationError("STRICT_PROFILE_CONFIG_ERROR: " + "; ".join(failures))

    def _build_execution_meta(self) -> dict[str, Any]:
        strict_mode = bool(self._strict_fail_closed)
        return {
            "execution_profile": str(getattr(self.settings, "release_profile", "compat") or "compat"),
            "reasoning_mode": str(getattr(self, "_reasoning_mode", "compat") or "compat"),
            "strict_mode": strict_mode,
            "retrieval_backend": str(getattr(self.settings, "vector_backend", "local") or "local"),
            "embedding_backend": str(getattr(self.settings, "embedding_backend", "hash") or "hash"),
            "reranker_backend": str(getattr(self.settings, "reranker_backend", "lexical") or "lexical"),
            "fail_closed": strict_mode,
        }

    def _is_release_ready_chunk_for_citation(self, chunk: dict[str, Any]) -> bool:
        if not isinstance(chunk, dict):
            return False
        source_id = normalize_source_set_id(str(chunk.get("source_set") or ""))
        if source_id not in OFFICIAL_SOURCE_RULES:
            return False
        doc_id = str(chunk.get("doc_id") or "").strip()
        doc_version = str(chunk.get("doc_version") or "").strip()
        if not doc_id or not doc_version:
            return False
        version_row = self.store.get_guideline_version_by_doc(doc_id=doc_id, doc_version=doc_version)
        if not isinstance(version_row, dict):
            return False
        return self._is_version_release_ready_for_retrieval(
            source_id=source_id,
            doc_id=doc_id,
            cancer_type=str(chunk.get("cancer_type") or ""),
            version_row=version_row,
        )

    def _sanitize_citations(
        self,
        *,
        citations: list[dict[str, Any]],
        chunk_to_citation: dict[str, str],
        reranked_chunks: list[dict[str, Any]],
    ) -> tuple[list[dict[str, Any]], dict[str, str]]:
        chunk_by_id = {
            str(chunk.get("chunk_id") or "").strip(): chunk
            for chunk in reranked_chunks
            if isinstance(chunk, dict) and str(chunk.get("chunk_id") or "").strip()
        }
        citation_to_chunk = {
            str(citation_id): str(chunk_id)
            for chunk_id, citation_id in (chunk_to_citation or {}).items()
            if str(chunk_id).strip() and str(citation_id).strip()
        }
        filtered_citations: list[dict[str, Any]] = []
        filtered_map: dict[str, str] = {}
        seen_citation_ids: set[str] = set()
        for citation in citations:
            if not isinstance(citation, dict):
                continue
            citation_id = str(citation.get("citation_id") or "").strip()
            if not citation_id or citation_id in seen_citation_ids:
                continue
            source_id = normalize_source_set_id(str(citation.get("source_id") or ""))
            file_uri = str(citation.get("file_uri") or "").strip().lower()
            if source_id in {"", "legacy_source"}:
                continue
            if file_uri in {"", "about:blank"}:
                continue
            chunk_id = citation_to_chunk.get(citation_id, "")
            chunk = chunk_by_id.get(chunk_id) if chunk_id else None
            if self._strict_fail_closed:
                if source_id not in OFFICIAL_SOURCE_RULES:
                    continue
                if not isinstance(chunk, dict) or not self._is_release_ready_chunk_for_citation(chunk):
                    continue
            normalized_citation = dict(citation)
            normalized_citation["source_id"] = source_id
            filtered_citations.append(normalized_citation)
            if chunk_id:
                filtered_map[chunk_id] = citation_id
            seen_citation_ids.add(citation_id)
        return filtered_citations, filtered_map

    @staticmethod
    def _filter_plan_sections_with_citations(plan_sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
        cleaned_sections: list[dict[str, Any]] = []
        for section in plan_sections:
            if not isinstance(section, dict):
                continue
            steps = section.get("steps") if isinstance(section.get("steps"), list) else []
            cleaned_steps: list[dict[str, Any]] = []
            for step in steps:
                if not isinstance(step, dict):
                    continue
                citation_ids = [str(item) for item in (step.get("citation_ids") if isinstance(step.get("citation_ids"), list) else []) if str(item).strip()]
                if not citation_ids:
                    continue
                step_copy = dict(step)
                step_copy["citation_ids"] = list(dict.fromkeys(citation_ids))
                cleaned_steps.append(step_copy)
            if not cleaned_steps:
                continue
            section_copy = dict(section)
            section_copy["steps"] = cleaned_steps
            cleaned_sections.append(section_copy)
        return cleaned_sections

    @staticmethod
    def _attach_drug_signal_citations(
        *,
        drug_safety_payload: dict[str, Any],
        citation_ids: list[str],
        strict_fail_closed: bool,
    ) -> dict[str, Any]:
        payload = copy.deepcopy(drug_safety_payload if isinstance(drug_safety_payload, dict) else {})
        signals = payload.get("signals") if isinstance(payload.get("signals"), list) else []
        default_citation_id = str(citation_ids[0]) if citation_ids else ""
        filtered_signals: list[dict[str, Any]] = []
        for raw_signal in signals:
            if not isinstance(raw_signal, dict):
                continue
            signal = dict(raw_signal)
            linked = signal.get("citation_ids") if isinstance(signal.get("citation_ids"), list) else []
            normalized_ids = [str(item) for item in linked if str(item).strip()]
            if not normalized_ids and default_citation_id:
                normalized_ids = [default_citation_id]
            signal["citation_ids"] = normalized_ids
            if not str(signal.get("source_origin") or "").strip():
                signal["source_origin"] = "rule_engine"
            severity = str(signal.get("severity") or "").strip().lower()
            if severity in {"critical", "warning"} and not normalized_ids:
                if strict_fail_closed:
                    continue
                signal["severity"] = "info"
                signal["kind"] = "missing_data"
                details = str(signal.get("details") or "").strip()
                note = "Сигнал понижен: отсутствует трассируемая ссылка на источник."
                signal["details"] = f"{details} {note}".strip()
            filtered_signals.append(signal)
        payload["signals"] = filtered_signals
        return payload

    def __init__(self, settings: Settings) -> None:
        self.settings = settings
        self._reasoning_mode = normalize_reasoning_mode(str(getattr(settings, "reasoning_mode", "llm_rag_only")))
        self._llm_rag_only_mode = self._reasoning_mode == "llm_rag_only"
        self._strict_fail_closed = self._strict_profile_token(settings) == "strict_full"
        self.settings.data_dir.mkdir(parents=True, exist_ok=True)
        self.settings.docs_dir.mkdir(parents=True, exist_ok=True)
        self.settings.reports_dir.mkdir(parents=True, exist_ok=True)

        self.logger = get_logger()
        self.store = SQLiteStore(settings.db_path)
        self._bootstrap_reference_seeds()
        self._drug_dictionary_entries: list[dict[str, Any]] = []
        self._drug_regimen_aliases: list[dict[str, Any]] = []
        self._drug_synonyms_extra: dict[str, Any] = {}
        self._bootstrap_drug_dictionary_seed()
        self._reload_drug_dictionary_cache()
        self._regimen_aliases = self._load_regimen_aliases()
        self.drug_safety_provider = DrugSafetyProvider(
            store=self.store,
            cache_ttl_hours=int(getattr(self.settings, "drug_safety_cache_ttl_hours", 24 * 14)),
            request_timeout_sec=int(getattr(self.settings, "drug_safety_request_timeout_sec", 12)),
            openfda_base_url=str(getattr(self.settings, "drug_safety_openfda_base_url", "https://api.fda.gov")),
        )
        self.prompt_registry = PromptRegistry(settings.prompt_registry_dir)

        llm_enabled = settings.llm_generation_enabled
        primary = (
            self._build_llm_endpoint(settings.llm_primary_url, settings.llm_primary_model, settings.llm_primary_api_key)
            if llm_enabled
            else None
        )
        fallback = (
            self._build_llm_endpoint(settings.llm_fallback_url, settings.llm_fallback_model, settings.llm_fallback_api_key)
            if llm_enabled
            else None
        )
        self._ensure_strict_release_requirements(settings=settings, primary_endpoint=primary)
        if self._strict_fail_closed and fallback is not None:
            raise ValidationError("STRICT_PROFILE_CONFIG_ERROR: strict_full forbids fallback LLM endpoint")
        if self._llm_rag_only_mode and fallback is not None:
            raise ValidationError("STRICT_LLM_RAG_ONLY: fallback LLM endpoint is forbidden")
        self.llm_router = LLMProviderRouter(primary=primary, fallback=fallback)
        self.embedder = build_embedder(
            backend=settings.embedding_backend,
            url=settings.embedding_url,
            model=settings.embedding_model,
            api_key=settings.embedding_api_key,
            fail_closed=self._strict_fail_closed,
        )

        local_index = LocalQdrantLikeIndex(self.store, embedder=self.embedder)
        if settings.vector_backend == "qdrant" and settings.qdrant_url:
            self.index = QdrantRestIndex(
                qdrant_url=settings.qdrant_url,
                collection=settings.qdrant_collection,
                fallback_index=local_index,
                embedder=self.embedder,
                fail_closed=self._strict_fail_closed,
            )
        else:
            if self._strict_fail_closed:
                raise ValidationError("STRICT_PROFILE_CONFIG_ERROR: VECTOR_BACKEND=qdrant and QDRANT_URL are required")
            self.index = local_index
        self.retriever, self._retrieval_engine, self._retrieval_fallback_reason = build_retriever(
            requested_engine=settings.rag_engine,
            index=self.index,
            top_k=settings.retrieval_top_k,
            fail_closed=self._strict_fail_closed,
        )
        self.reranker = Reranker(
            top_n=settings.rerank_top_n,
            backend=settings.reranker_backend,
            llm_router=self.llm_router,
            fail_closed=self._strict_fail_closed,
        )
        self.rate_limiter = RateLimiter(max_requests_per_minute=settings.rate_limit_per_minute)
        self._kb_version = compute_kb_version(self.store.list_docs())
        self._routing_cache_ttl_seconds = 15 * 60
        self._routing_cache: dict[str, tuple[float, NosologyRouteDecision]] = {}

    @staticmethod
    def _builtin_disease_registry_seed() -> list[dict[str, Any]]:
        detailed_entries: list[dict[str, Any]] = [
            {
                "schema_version": "1.0",
                "disease_id": "a76e5701-e3b1-54fd-a4b8-001bcd63de6e",
                "icd10_codes": ["C16"],
                "disease_name_ru": "Рак желудка",
                "disease_name_en": "Gastric cancer",
                "common_synonyms": ["рак желудка", "gastric cancer", "stomach cancer"],
                "active": True,
                "updated_at": "2026-02-20T00:00:00Z",
            },
            {
                "schema_version": "1.0",
                "disease_id": "2efcb0a0-2b4a-5f44-a247-9e1c6d9a7f42",
                "icd10_codes": ["C34"],
                "disease_name_ru": "Немелкоклеточный рак легкого",
                "disease_name_en": "Non-small cell lung cancer",
                "common_synonyms": ["рак легкого", "нмрл", "nsclc", "non-small cell lung cancer"],
                "active": True,
                "updated_at": "2026-02-20T00:00:00Z",
            },
            {
                "schema_version": "1.0",
                "disease_id": "9d9d8f58-2a2d-5c9d-b43d-7d4af8854d38",
                "icd10_codes": ["C50"],
                "disease_name_ru": "Рак молочной железы",
                "disease_name_en": "Breast cancer",
                "common_synonyms": ["рак молочной железы", "рмж", "breast cancer"],
                "active": True,
                "updated_at": "2026-02-20T00:00:00Z",
            },
            {
                "schema_version": "1.0",
                "disease_id": "c8b1f6d0-4b6f-53cf-9e7d-6df58cc1ad5f",
                "icd10_codes": ["C18", "C19", "C20"],
                "disease_name_ru": "Колоректальный рак",
                "disease_name_en": "Colorectal cancer",
                "common_synonyms": ["колоректальный рак", "рак толстой кишки", "рак прямой кишки", "crc"],
                "active": True,
                "updated_at": "2026-02-21T00:00:00Z",
            },
            {
                "schema_version": "1.0",
                "disease_id": "b53b53b7-f1e4-58ef-8d3d-5846df8f9a10",
                "icd10_codes": ["C61"],
                "disease_name_ru": "Рак предстательной железы",
                "disease_name_en": "Prostate cancer",
                "common_synonyms": ["рак предстательной железы", "рак простаты", "prostate cancer"],
                "active": True,
                "updated_at": "2026-02-21T00:00:00Z",
            },
            {
                "schema_version": "1.0",
                "disease_id": "e4d29126-54ce-56cb-88dc-2dcf4954eaf9",
                "icd10_codes": ["C64"],
                "disease_name_ru": "Почечно-клеточный рак",
                "disease_name_en": "Renal cell carcinoma",
                "common_synonyms": ["почечно-клеточный рак", "рак почки", "renal cell carcinoma"],
                "active": True,
                "updated_at": "2026-02-21T00:00:00Z",
            },
            {
                "schema_version": "1.0",
                "disease_id": "d80c5e16-28df-5f1d-b88b-f76795db4c59",
                "icd10_codes": ["C67"],
                "disease_name_ru": "Рак мочевого пузыря",
                "disease_name_en": "Bladder cancer",
                "common_synonyms": ["рак мочевого пузыря", "urothelial carcinoma", "bladder cancer"],
                "active": True,
                "updated_at": "2026-02-21T00:00:00Z",
            },
            {
                "schema_version": "1.0",
                "disease_id": "c0a0a03b-040b-5314-9802-abef422d53b5",
                "icd10_codes": ["C71"],
                "disease_name_ru": "Первичные опухоли головного мозга",
                "disease_name_en": "Primary malignant brain tumors",
                "common_synonyms": ["первичная опухоль головного мозга", "глиома", "glioblastoma", "brain primary cancer"],
                "active": True,
                "updated_at": "2026-02-21T00:00:00Z",
            },
            {
                "schema_version": "1.0",
                "disease_id": "7a2bf75a-b89e-5fb9-bc16-ee6eae6c27b8",
                "icd10_codes": ["C79.3"],
                "disease_name_ru": "Метастазы в ЦНС",
                "disease_name_en": "Secondary malignant neoplasm of brain (CNS metastases)",
                "common_synonyms": ["метастазы в головной мозг", "метастазы в цнс", "brain metastases", "cns metastases"],
                "active": True,
                "updated_at": "2026-02-21T00:00:00Z",
            },
        ]
        covered_codes = {
            str(code).strip().upper()
            for item in detailed_entries
            for code in (item.get("icd10_codes") if isinstance(item.get("icd10_codes"), list) else [])
            if str(code).strip()
        }
        generated_entries: list[dict[str, Any]] = []
        for idx in range(98):
            prefix = f"C{idx:02d}"
            if prefix in covered_codes:
                continue
            generated_entries.append(
                {
                    "schema_version": "1.0",
                    "disease_id": str(uuid.uuid5(uuid.NAMESPACE_URL, f"oncoai:disease:icd10:{prefix}")),
                    "icd10_codes": [prefix],
                    "disease_name_ru": f"Злокачественное новообразование ({prefix})",
                    "disease_name_en": f"Malignant neoplasm ({prefix})",
                    "common_synonyms": [
                        f"рак {prefix}",
                        f"онкология {prefix}",
                        f"malignant neoplasm {prefix}",
                    ],
                    "active": True,
                    "updated_at": "2026-02-23T00:00:00Z",
                }
            )
        return [*detailed_entries, *generated_entries]

    def _discover_seed_dirs(self) -> list[Path]:
        seed_dirs: list[Path] = []
        seen: set[str] = set()
        for candidate in (
            self.settings.project_root / "docs" / "contracts" / "onco_json_pack_v1" / "seeds",
            Path(__file__).resolve().parents[2] / "docs" / "contracts" / "onco_json_pack_v1" / "seeds",
        ):
            key = str(candidate.resolve()) if candidate.exists() else str(candidate)
            if key in seen:
                continue
            seen.add(key)
            if candidate.exists():
                seed_dirs.append(candidate)
        return seed_dirs

    def _bootstrap_reference_seeds(self) -> None:
        seed_dirs = self._discover_seed_dirs()

        loaded_sources = 0
        loaded_diseases = 0
        for seed_dir in seed_dirs:
            sources_path = seed_dir / "guideline_sources.seed.json"
            if sources_path.exists():
                try:
                    payload = json.loads(sources_path.read_text(encoding="utf-8"))
                    if isinstance(payload, list):
                        for item in payload:
                            if isinstance(item, dict):
                                self.store.upsert_guideline_source(item)
                                loaded_sources += 1
                except Exception:  # noqa: BLE001
                    continue

            diseases_path = seed_dir / "disease_registry.seed.json"
            if diseases_path.exists():
                try:
                    payload = json.loads(diseases_path.read_text(encoding="utf-8"))
                    if isinstance(payload, list):
                        for item in payload:
                            if isinstance(item, dict):
                                self.store.upsert_disease_registry_entry(item)
                                loaded_diseases += 1
                except Exception:  # noqa: BLE001
                    continue

        # Ensure minimum disease coverage for core demo ICD10 groups
        # even when external seed files are not mounted in runtime container.
        for item in self._builtin_disease_registry_seed():
            self.store.upsert_disease_registry_entry(item)
            loaded_diseases += 1

        if loaded_sources or loaded_diseases:
            safe_log(
                self.logger,
                "governance.seed_bootstrap",
                {
                    "sources_loaded": loaded_sources,
                    "diseases_loaded": loaded_diseases,
                },
            )

    def _reload_drug_dictionary_cache(self) -> None:
        self._drug_dictionary_entries = self.store.list_drug_dictionary_entries(limit=50000)
        self._drug_regimen_aliases = self.store.list_drug_regimen_aliases(limit=50000)
        self._drug_synonyms_extra = self._load_latest_drug_synonyms_extra()

    def _load_latest_drug_synonyms_extra(self) -> dict[str, Any]:
        versions = self.store.list_drug_dictionary_versions(limit=1)
        if not versions:
            return {}
        metadata = versions[0].get("metadata") if isinstance(versions[0], dict) else {}
        if not isinstance(metadata, dict):
            return {}
        synonyms = metadata.get("synonyms_extra")
        return dict(synonyms) if isinstance(synonyms, dict) else {}

    def _apply_drug_dictionary_bundle(
        self,
        *,
        bundle: DrugDictionaryBundle,
        origin: str,
        loaded_by: str,
    ) -> dict[str, Any]:
        updated_at = datetime.now(timezone.utc).isoformat()
        replaced = self.store.replace_drug_dictionary(
            source_version=bundle.version,
            entries=bundle.entries,
            regimens=bundle.regimen_aliases,
            updated_at=updated_at,
        )
        self.store.save_drug_dictionary_version(
            version_id=bundle.version,
            sha256=bundle.sha256,
            loaded_at=updated_at,
            metadata={
                "schema": bundle.schema,
                "notes": bundle.notes,
                "origin": origin,
                "loaded_by": loaded_by,
                "synonyms_extra": bundle.synonyms_extra,
            },
        )
        self._reload_drug_dictionary_cache()
        return {
            "version": bundle.version,
            "sha256": bundle.sha256,
            "entries": int(replaced.get("entries") or 0),
            "regimens": int(replaced.get("regimens") or 0),
            "loaded_at": updated_at,
        }

    def _bootstrap_drug_dictionary_seed(self) -> None:
        seed_filename = "drug_dictionary_ru_inn.v1.2.json"
        seed_path: Path | None = None
        for seed_dir in self._discover_seed_dirs():
            candidate = seed_dir / seed_filename
            if candidate.exists():
                seed_path = candidate
                break
        if seed_path is None:
            return
        try:
            bundle = load_drug_dictionary_bundle_from_path(seed_path)
            loaded = self._apply_drug_dictionary_bundle(
                bundle=bundle,
                origin=str(seed_path),
                loaded_by="bootstrap",
            )
            safe_log(
                self.logger,
                "drug_dictionary.bootstrap",
                {
                    "version": loaded["version"],
                    "entries": loaded["entries"],
                    "regimens": loaded["regimens"],
                },
            )
        except Exception:  # noqa: BLE001
            return

    def _load_regimen_aliases(self) -> dict[str, str]:
        aliases: dict[str, str] = {}
        for seed_dir in self._discover_seed_dirs():
            regimens_path = seed_dir / "regimen_aliases.seed.json"
            if not regimens_path.exists():
                continue
            try:
                payload = json.loads(regimens_path.read_text(encoding="utf-8"))
            except Exception:  # noqa: BLE001
                continue

            regimens = payload.get("regimens") if isinstance(payload, dict) else []
            if not isinstance(regimens, list):
                continue
            for item in regimens:
                if not isinstance(item, dict):
                    continue
                regimen_id = str(item.get("regimen_id") or "").strip().upper()
                if not regimen_id:
                    continue
                raw_aliases = item.get("aliases") if isinstance(item.get("aliases"), list) else []
                for alias in [regimen_id, *raw_aliases]:
                    if not isinstance(alias, str):
                        continue
                    key = alias.strip().lower()
                    if key:
                        aliases[key] = regimen_id

        if aliases:
            return aliases
        # Safe defaults for environments without the seed file.
        return {
            "xelox": "XELOX",
            "capox": "CAPOX",
            "mfolfox6": "FOLFOX",
            "folfox": "FOLFOX",
            "folfox6": "FOLFOX",
        }

    def _governance_snapshot(self) -> dict[str, Any]:
        sources = self.store.list_guideline_sources()
        documents = self.store.list_guideline_documents()
        versions = self.store.list_guideline_versions()
        disease_registry = self.store.list_disease_registry()
        ingestion_runs = self.store.list_ingestion_runs(limit=5)
        return {
            "sources_total": len(sources),
            "documents_total": len(documents),
            "versions_total": len(versions),
            "disease_registry_total": len(disease_registry),
            "sources": sources,
            "disease_registry": disease_registry,
            "latest_ingestion_runs": ingestion_runs,
        }

    @staticmethod
    def _normalize_case_data_mode(value: Any) -> str:
        mode = str(value or "DEID").strip().upper()
        if mode not in SUPPORTED_CASE_DATA_MODES:
            return "DEID"
        return mode

    def _apply_deid_redaction(
        self,
        *,
        case_json: dict[str, Any],
    ) -> tuple[dict[str, Any], bool]:
        if not bool(getattr(self.settings, "case_import_deid_redact_pii", True)):
            return case_json, False

        redacted = copy.deepcopy(case_json)
        changed = False

        def redact_text(value: Any) -> str:
            nonlocal changed
            original = str(value or "")
            masked = redact_pii(original)
            if masked != original:
                changed = True
            return masked

        if isinstance(redacted.get("notes"), str):
            redacted["notes"] = redact_text(redacted.get("notes"))

        patient = redacted.get("patient")
        if isinstance(patient, dict):
            for key, value in list(patient.items()):
                if not isinstance(value, str):
                    continue
                key_lower = str(key).strip().lower()
                if any(marker in key_lower for marker in ("name", "phone", "email", "contact")):
                    patient[key] = redact_text(value)

        diagnoses = redacted.get("diagnoses")
        if isinstance(diagnoses, list):
            for diagnosis in diagnoses:
                if not isinstance(diagnosis, dict):
                    continue
                timeline = diagnosis.get("timeline")
                if isinstance(timeline, list):
                    for event in timeline:
                        if not isinstance(event, dict):
                            continue
                        for text_field in ("label", "details"):
                            if isinstance(event.get(text_field), str):
                                event[text_field] = redact_text(event.get(text_field))

        source_refs = redacted.get("source_refs")
        if isinstance(source_refs, list):
            for item in source_refs:
                if not isinstance(item, dict):
                    continue
                for text_field in ("snippet", "note"):
                    if isinstance(item.get(text_field), str):
                        item[text_field] = redact_text(item.get(text_field))

        return redacted, changed

    @staticmethod
    def _fallback_case_payload(
        case_id: str,
        import_profile: str,
        notes: str,
        now: str,
        *,
        data_mode: str = "DEID",
    ) -> dict[str, Any]:
        diagnosis_id = str(uuid.uuid4())
        disease_id = str(uuid.uuid5(uuid.NAMESPACE_URL, f"oncoai:disease:{case_id}"))
        diagnosis: dict[str, Any] = {
            "diagnosis_id": diagnosis_id,
            "disease_id": disease_id,
        }
        return {
            "schema_version": "1.0",
            "case_id": case_id,
            "created_at": now,
            "updated_at": now,
            "data_mode": data_mode,
            "import_profile": import_profile,
            "patient": {"sex": "unknown"},
            "diagnoses": [diagnosis],
            "attachments": [],
            "notes": notes,
        }

    @staticmethod
    def _case_import_llm_output_schema() -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "patient": {
                    "type": "object",
                    "properties": {
                        "sex": {"type": ["string", "null"]},
                        "birth_year": {"type": ["integer", "null"]},
                        "height_cm": {"type": ["number", "null"]},
                        "weight_kg": {"type": ["number", "null"]},
                        "ecog": {"type": ["integer", "null"]},
                    },
                    "required": ["sex", "birth_year", "height_cm", "weight_kg", "ecog"],
                    "additionalProperties": False,
                },
                "diagnosis": {
                    "type": "object",
                    "properties": {
                        "icd10": {"type": ["string", "null"]},
                        "histology": {"type": ["string", "null"]},
                        "stage_group": {"type": ["string", "null"]},
                        "regimen": {"type": ["string", "null"]},
                        "line": {"type": ["integer", "null"]},
                        "cycle": {"type": ["integer", "null"]},
                        "biomarkers": {
                            "type": ["array", "null"],
                            "items": {
                                "type": "object",
                                "properties": {
                                    "name": {"type": ["string", "null"]},
                                    "value": {"type": ["string", "null"]},
                                },
                                "required": ["name", "value"],
                                "additionalProperties": False,
                            },
                        },
                    },
                    "required": ["icd10", "histology", "stage_group", "regimen", "line", "cycle", "biomarkers"],
                    "additionalProperties": False,
                },
                "missing_required_fields": {"type": "array", "items": {"type": "string"}},
                "warnings": {"type": "array", "items": {"type": "string"}},
            },
            "required": ["patient", "diagnosis", "missing_required_fields", "warnings"],
            "additionalProperties": False,
        }

    def _require_primary_llm_for_strict_mode(self, *, context: str) -> None:
        if self.llm_router.primary is None:
            raise ValidationError(f"STRICT_LLM_RAG_ONLY: primary LLM provider is required for {context}")

    def _call_case_import_llm(
        self,
        *,
        import_profile: str,
        case_text: str,
    ) -> dict[str, Any]:
        self._require_primary_llm_for_strict_mode(context=f"{import_profile} case import")
        prompt = (
            "Ты извлекаешь структурированный онкологический кейс из клинического текста.\n"
            "Верни строго JSON по схеме.\n"
            "Не добавляй markdown.\n"
            "Если поле отсутствует в тексте, укажи null и добавь путь поля в missing_required_fields.\n"
            "warnings: короткие русские строки о неопределенностях.\n"
            f"import_profile={import_profile}\n"
            f"case_text={case_text[:12000]}\n"
        )
        payload, path = self.llm_router.generate_json(
            prompt=prompt,
            output_schema=self._case_import_llm_output_schema(),
            schema_name="case_import_structured_v1",
        )
        if str(path or "").strip().lower() != "primary":
            raise ValidationError("STRICT_LLM_RAG_ONLY: case import must use primary LLM path")
        if not isinstance(payload, dict):
            raise ValidationError("STRICT_LLM_RAG_ONLY: LLM returned empty case import payload")
        return payload

    def _normalize_case_from_llm_payload(
        self,
        *,
        llm_payload: dict[str, Any],
        import_profile: str,
        case_id: str,
        now: str,
        data_mode: str,
        notes_text: str,
        attachment: dict[str, Any] | None = None,
    ) -> tuple[dict[str, Any], list[str], list[dict[str, str]]]:
        patient_payload = llm_payload.get("patient") if isinstance(llm_payload.get("patient"), dict) else {}
        diagnosis_payload = llm_payload.get("diagnosis") if isinstance(llm_payload.get("diagnosis"), dict) else {}

        patient: dict[str, Any] = {"sex": self._normalize_sex(patient_payload.get("sex"))}
        birth_year_raw = patient_payload.get("birth_year")
        if isinstance(birth_year_raw, int) and 1900 <= birth_year_raw <= datetime.now(timezone.utc).year:
            patient["birth_year"] = birth_year_raw
        height_raw = patient_payload.get("height_cm")
        if isinstance(height_raw, (int, float)):
            patient["height_cm"] = float(height_raw)
        weight_raw = patient_payload.get("weight_kg")
        if isinstance(weight_raw, (int, float)):
            patient["weight_kg"] = float(weight_raw)
        ecog_raw = patient_payload.get("ecog")
        if isinstance(ecog_raw, int) and 0 <= ecog_raw <= 5:
            patient["ecog"] = ecog_raw

        diagnosis: dict[str, Any] = {
            "diagnosis_id": str(uuid.uuid4()),
            "disease_id": str(uuid.uuid5(uuid.NAMESPACE_URL, f"oncoai:disease:{case_id}")),
            "biomarkers": [],
            "timeline": [],
        }
        icd10 = str(diagnosis_payload.get("icd10") or "").strip().upper()
        if icd10:
            diagnosis["icd10"] = icd10
            diagnosis["disease_id"] = self._disease_id_from_icd10(icd10, case_id)
        histology = str(diagnosis_payload.get("histology") or "").strip()
        if histology:
            diagnosis["histology"] = histology
        stage_group = str(diagnosis_payload.get("stage_group") or "").strip().upper()
        if stage_group:
            diagnosis["stage"] = {"system": "UNKNOWN", "stage_group": stage_group}
        regimen = str(diagnosis_payload.get("regimen") or "").strip()
        line = diagnosis_payload.get("line")
        cycle = diagnosis_payload.get("cycle")
        if regimen:
            last_plan: dict[str, Any] = {"date": now[:10], "precision": "day", "regimen": regimen}
            if isinstance(line, int) and line > 0:
                last_plan["line"] = line
            if isinstance(cycle, int) and cycle > 0:
                last_plan["cycle"] = cycle
            diagnosis["last_plan"] = last_plan
        biomarkers_raw = diagnosis_payload.get("biomarkers")
        if isinstance(biomarkers_raw, list):
            biomarkers: list[dict[str, str]] = []
            for item in biomarkers_raw:
                if not isinstance(item, dict):
                    continue
                name = str(item.get("name") or "").strip()
                value = str(item.get("value") or "").strip()
                if name and value:
                    biomarkers.append({"name": name, "value": value})
            diagnosis["biomarkers"] = biomarkers

        missing_required_fields = [
            str(item).strip()
            for item in (
                llm_payload.get("missing_required_fields")
                if isinstance(llm_payload.get("missing_required_fields"), list)
                else []
            )
            if str(item).strip()
        ]
        if not icd10 and "diagnoses[0].icd10" not in missing_required_fields:
            missing_required_fields.append("diagnoses[0].icd10")
        if not regimen and "diagnoses[0].last_plan.regimen" not in missing_required_fields:
            missing_required_fields.append("diagnoses[0].last_plan.regimen")

        warnings: list[dict[str, str]] = []
        raw_warnings = llm_payload.get("warnings") if isinstance(llm_payload.get("warnings"), list) else []
        for item in raw_warnings:
            text = str(item).strip()
            if text:
                warnings.append({"code": "LLM_IMPORT_WARNING", "message": text})

        attachments = [attachment] if isinstance(attachment, dict) else []
        case_json = {
            "schema_version": "1.0",
            "case_id": case_id,
            "created_at": now,
            "updated_at": now,
            "data_mode": data_mode,
            "import_profile": import_profile,
            "patient": patient,
            "diagnoses": [diagnosis],
            "attachments": attachments,
            "notes": notes_text[:4000],
        }
        return case_json, sorted(set(missing_required_fields)), warnings

    def _build_case_json_from_import_llm(
        self,
        *,
        import_profile: str,
        payload: dict[str, Any],
        case_id: str,
        now: str,
        data_mode: str,
    ) -> tuple[dict[str, Any], list[str]]:
        if import_profile == "FREE_TEXT":
            notes_text = str(payload.get("free_text") or "").strip()
        else:
            template_obj = payload.get("custom_template") if isinstance(payload.get("custom_template"), dict) else {}
            notes_text = json.dumps(template_obj, ensure_ascii=False).strip()
        if not notes_text:
            raise ValidationError(f"{import_profile} requires non-empty clinical text for llm_rag_only mode")
        llm_payload = self._call_case_import_llm(import_profile=import_profile, case_text=notes_text)
        case_json, missing_required_fields, _warnings = self._normalize_case_from_llm_payload(
            llm_payload=llm_payload,
            import_profile=import_profile,
            case_id=case_id,
            now=now,
            data_mode=data_mode,
            notes_text=notes_text,
        )
        return case_json, missing_required_fields

    def _build_case_json_from_kin_pdf_llm(
        self,
        *,
        payload: dict[str, Any],
        case_id: str,
        now: str,
        data_mode: str,
    ) -> tuple[dict[str, Any], list[str], list[dict[str, str]]]:
        text = self._extract_kin_text(payload)
        if not text:
            raise ValidationError("KIN_PDF requires `kin_pdf_text` or `kin_pdf.text`/`kin_pdf.pages`")
        llm_payload = self._call_case_import_llm(import_profile="KIN_PDF", case_text=text)
        filename = str(payload.get("filename") or "kin_case.pdf").strip() or "kin_case.pdf"
        attachment_doc_id = str(uuid.uuid5(uuid.NAMESPACE_URL, f"oncoai:case:{case_id}:kin_pdf"))
        attachment = {
            "doc_id": attachment_doc_id,
            "kind": "pdf",
            "filename": filename,
            "hash": self._hash_text_payload(text),
        }
        return self._normalize_case_from_llm_payload(
            llm_payload=llm_payload,
            import_profile="KIN_PDF",
            case_id=case_id,
            now=now,
            data_mode=data_mode,
            notes_text=text,
            attachment=attachment,
        )

    def _build_case_json_from_import(
        self,
        *,
        import_profile: str,
        payload: dict[str, Any],
        case_id: str,
        now: str,
        data_mode: str,
    ) -> tuple[dict[str, Any], list[str]]:
        if import_profile == "FREE_TEXT":
            free_text = str(payload.get("free_text") or "").strip()
            missing_fields: list[str] = []
            if not free_text:
                free_text = "No free-text details provided."
                missing_fields.append("free_text")
            lowered = free_text.lower()
            sex = "unknown"
            if "муж" in lowered or "male" in lowered:
                sex = "male"
            elif "жен" in lowered or "female" in lowered:
                sex = "female"

            patient: dict[str, Any] = {"sex": sex}
            birth_year = self._extract_year(free_text)
            if birth_year is not None:
                patient["birth_year"] = birth_year
            else:
                age_year = self._extract_birth_year_from_age(free_text)
                if age_year is not None:
                    patient["birth_year"] = age_year
            height_cm = self._extract_height_cm(free_text)
            if height_cm is not None:
                patient["height_cm"] = height_cm
            weight_kg = self._extract_weight_kg(free_text)
            if weight_kg is not None:
                patient["weight_kg"] = weight_kg
            ecog = self._extract_ecog(free_text)
            if ecog is not None:
                patient["ecog"] = ecog

            icd10 = self._extract_first_icd10(free_text)
            stage_group = self._extract_stage_group(free_text)
            regimen = self._extract_regimen_name(free_text)
            line, cycle = self._extract_line_cycle(free_text)

            biomarkers = self._extract_biomarkers_from_text(free_text)
            histology = ""
            if "аденокарц" in lowered or "adenocarc" in lowered:
                histology = "adenocarcinoma"
            elif "плоскоклет" in lowered or "squamous" in lowered:
                histology = "squamous_cell_carcinoma"

            if not icd10:
                missing_fields.append("diagnoses[0].icd10")

            diagnosis: dict[str, Any] = {
                "diagnosis_id": str(uuid.uuid4()),
                "disease_id": self._disease_id_from_icd10(icd10, case_id),
                "biomarkers": biomarkers,
                "timeline": [],
            }
            if icd10:
                diagnosis["icd10"] = icd10
            if histology:
                diagnosis["histology"] = histology
            if stage_group:
                diagnosis["stage"] = {"system": "UNKNOWN", "stage_group": stage_group}
            if regimen:
                last_plan: dict[str, Any] = {
                    "date": now[:10],
                    "precision": "day",
                    "regimen": regimen,
                }
                if line is not None:
                    last_plan["line"] = line
                if cycle is not None:
                    last_plan["cycle"] = cycle
                diagnosis["last_plan"] = last_plan

            case_json = {
                "schema_version": "1.0",
                "case_id": case_id,
                "created_at": now,
                "updated_at": now,
                "data_mode": data_mode,
                "import_profile": import_profile,
                "patient": patient,
                "diagnoses": [diagnosis],
                "attachments": [],
                "notes": free_text,
            }
            return case_json, missing_fields

        template = payload.get("custom_template")
        template_obj = template if isinstance(template, dict) else {}
        missing_fields = [] if template_obj else ["custom_template"]
        notes = (
            str(template_obj.get("notes") or "").strip()
            if template_obj
            else "No custom template details provided."
        )
        if not notes and template_obj:
            notes = json.dumps(template_obj, ensure_ascii=False)

        case_json = self._fallback_case_payload(case_id, import_profile, notes, now, data_mode=data_mode)
        patient = case_json["patient"]
        if isinstance(template_obj.get("sex"), str) and str(template_obj["sex"]).strip():
            patient["sex"] = str(template_obj["sex"]).strip()
        if isinstance(template_obj.get("birth_year"), int):
            patient["birth_year"] = int(template_obj["birth_year"])
        if isinstance(template_obj.get("ecog"), int):
            patient["ecog"] = int(template_obj["ecog"])

        diagnosis = case_json["diagnoses"][0]
        if isinstance(template_obj.get("icd10"), str) and str(template_obj["icd10"]).strip():
            diagnosis["icd10"] = str(template_obj["icd10"]).strip().upper()
        if isinstance(template_obj.get("histology"), str) and str(template_obj["histology"]).strip():
            diagnosis["histology"] = str(template_obj["histology"]).strip()
        if isinstance(template_obj.get("stage"), str) and str(template_obj["stage"]).strip():
            diagnosis["stage"] = {"system": "UNKNOWN", "stage_group": str(template_obj["stage"]).strip()}
        if isinstance(template_obj.get("regimen"), str) and str(template_obj["regimen"]).strip():
            diagnosis["last_plan"] = {
                "date": now[:10],
                "precision": "day",
                "regimen": str(template_obj["regimen"]).strip(),
                "line": 1,
            }
        return case_json, missing_fields

    def _normalize_imported_case(
        self,
        *,
        case_json: dict[str, Any],
        import_profile: str,
        case_id: str,
        now: str,
        data_mode: str,
    ) -> tuple[dict[str, Any], list[str]]:
        normalized = dict(case_json)
        missing_fields: list[str] = []

        normalized["schema_version"] = "1.0"
        normalized["case_id"] = str(normalized.get("case_id") or case_id)
        normalized["import_profile"] = str(normalized.get("import_profile") or import_profile)
        normalized["created_at"] = str(normalized.get("created_at") or now)
        normalized["updated_at"] = str(normalized.get("updated_at") or now)
        normalized["data_mode"] = self._normalize_case_data_mode(data_mode)

        patient = normalized.get("patient")
        if not isinstance(patient, dict):
            patient = {"sex": "unknown"}
            missing_fields.append("patient")
        elif not isinstance(patient.get("sex"), str) or not str(patient.get("sex")).strip():
            patient = dict(patient)
            patient["sex"] = "unknown"
            missing_fields.append("patient.sex")
        normalized["patient"] = patient

        diagnoses = normalized.get("diagnoses")
        if not isinstance(diagnoses, list) or not diagnoses:
            diagnoses = [
                {
                    "diagnosis_id": str(uuid.uuid4()),
                    "disease_id": str(uuid.uuid5(uuid.NAMESPACE_URL, f"oncoai:disease:{normalized['case_id']}")),
                }
            ]
            missing_fields.append("diagnoses")
        normalized["diagnoses"] = diagnoses

        attachments = normalized.get("attachments")
        if not isinstance(attachments, list):
            attachments = []
            missing_fields.append("attachments")
        normalized["attachments"] = attachments

        if "notes" in normalized:
            normalized["notes"] = str(normalized.get("notes") or "")
        return normalized, missing_fields

    @staticmethod
    def _build_failed_case_import_run(
        *,
        import_run_id: str,
        case_id: str,
        import_profile: str,
        started_at: str,
        finished_at: str,
        error_code: str,
        error_message: str,
    ) -> dict[str, Any]:
        return {
            "schema_version": "1.0",
            "import_run_id": import_run_id,
            "case_id": case_id,
            "import_profile": import_profile,
            "started_at": started_at,
            "finished_at": finished_at,
            "status": "FAILED",
            "confidence": 0.0,
            "missing_required_fields": [],
            "warnings": [],
            "errors": [{"code": error_code, "message": error_message}],
        }

    @staticmethod
    def _hash_text_payload(value: str) -> str:
        digest = hashlib.sha256(value.encode("utf-8")).hexdigest()
        return f"sha256:{digest}"

    @staticmethod
    def _hash_json_payload(value: dict[str, Any]) -> str:
        encoded = json.dumps(value, ensure_ascii=False, sort_keys=True).encode("utf-8")
        digest = hashlib.sha256(encoded).hexdigest()
        return f"sha256:{digest}"

    @staticmethod
    def _normalize_sex(value: Any) -> str:
        text = str(value or "").strip().lower()
        if text in {"male", "m", "man", "муж", "м"}:
            return "male"
        if text in {"female", "f", "woman", "жен", "ж"}:
            return "female"
        if text in {"other"}:
            return "other"
        return "unknown"

    @staticmethod
    def _extract_first_icd10(text: str) -> str:
        match = _ICD10_PATTERN.search(text)
        return match.group(0).upper() if match else ""

    def _infer_icd10_from_text(self, text: str, *, explicit_code: str = "") -> dict[str, Any]:
        return infer_icd10_code(
            text=str(text or ""),
            explicit_code=str(explicit_code or ""),
            disease_registry=self.store.list_disease_registry(),
            icd10_reference=self.store.list_icd10_reference(limit=10000),
        )

    def _apply_icd10_inference_to_case(
        self,
        *,
        case_json: dict[str, Any],
        missing_required_fields: list[str],
        warnings: list[dict[str, str]],
    ) -> None:
        diagnoses = case_json.get("diagnoses")
        if not isinstance(diagnoses, list) or not diagnoses:
            return
        diagnosis = diagnoses[0] if isinstance(diagnoses[0], dict) else None
        if not isinstance(diagnosis, dict):
            return

        existing_icd10 = str(diagnosis.get("icd10") or "").strip().upper()
        notes = str(case_json.get("notes") or "")
        histology = str(diagnosis.get("histology") or "")
        explicit_hint = existing_icd10
        inference = self._infer_icd10_from_text(
            "\n".join(item for item in [notes, histology] if item),
            explicit_code=explicit_hint,
        )
        inferred_code = str(inference.get("code") or "").strip().upper()
        if inferred_code:
            diagnosis["icd10"] = inferred_code
            diagnosis["disease_id"] = self._disease_id_from_icd10(inferred_code, str(case_json.get("case_id") or ""))
            if "diagnoses[0].icd10" in missing_required_fields:
                missing_required_fields[:] = [item for item in missing_required_fields if item != "diagnoses[0].icd10"]
            method = str(inference.get("method") or "")
            if method and method != "explicit":
                reason = str(inference.get("reason") or method)
                warnings.append(
                    {
                        "code": "ICD10_INFERRED",
                        "message": f"ICD-10 код определен автоматически ({inferred_code}; method={method}; reason={reason}).",
                    }
                )
            return

        if "diagnoses[0].icd10" not in missing_required_fields:
            missing_required_fields.append("diagnoses[0].icd10")

    @staticmethod
    def _extract_stage_group(text: str) -> str:
        match = _STAGE_PATTERN.search(text)
        return match.group(1).upper() if match else ""

    def _extract_regimen_name(self, text: str) -> str:
        lowered = text.lower()
        for alias in sorted(self._regimen_aliases.keys(), key=len, reverse=True):
            pattern = rf"(?<![a-zA-Z0-9]){re.escape(alias)}(?![a-zA-Z0-9])"
            match = re.search(pattern, lowered, flags=re.IGNORECASE)
            if match:
                return match.group(0).upper()
        return ""

    @staticmethod
    def _extract_year(text: str) -> int | None:
        birth_match = _BIRTH_YEAR_PATTERN.search(text)
        if birth_match:
            year = int(birth_match.group(1))
            return year if 1900 <= year <= datetime.now(timezone.utc).year else None

        years = [int(match.group(1)) for match in _YEAR_PATTERN.finditer(text)]
        if not years:
            return None
        # In longitudinal oncology narratives dates of therapy dominate;
        # the earliest plausible year is usually the patient's birth year.
        year = min(years)
        return year if 1900 <= year <= datetime.now(timezone.utc).year else None

    @staticmethod
    def _extract_birth_year_from_age(text: str) -> int | None:
        match = _AGE_PATTERN.search(text)
        if not match:
            return None
        candidate = match.group(1) or match.group(2)
        if not candidate or not candidate.isdigit():
            return None
        age = int(candidate)
        if age < 1 or age > 120:
            return None
        current_year = datetime.now(timezone.utc).year
        birth_year = current_year - age
        return birth_year if 1900 <= birth_year <= current_year else None

    @staticmethod
    def _extract_height_cm(text: str) -> float | None:
        match = _HEIGHT_PATTERN.search(text)
        if not match:
            return None
        value = str(match.group(1) or "").strip().replace(",", ".")
        try:
            height = float(value)
        except ValueError:
            return None
        if not (90.0 <= height <= 250.0):
            return None
        return height

    @staticmethod
    def _extract_weight_kg(text: str) -> float | None:
        match = _WEIGHT_PATTERN.search(text)
        if not match:
            return None
        value = str(match.group(1) or "").strip().replace(",", ".")
        try:
            weight = float(value)
        except ValueError:
            return None
        if not (20.0 <= weight <= 350.0):
            return None
        return weight

    @staticmethod
    def _extract_ecog(text: str) -> int | None:
        match = _ECOG_PATTERN.search(text)
        if not match:
            return None
        try:
            ecog = int(str(match.group(1) or "").strip())
        except ValueError:
            return None
        return ecog if 0 <= ecog <= 5 else None

    @staticmethod
    def _extract_biomarkers_from_text(text: str) -> list[dict[str, str]]:
        biomarkers: list[dict[str, str]] = []
        seen: set[str] = set()
        for match in _BIOMARKER_PATTERN.finditer(text):
            name = match.group(1).upper().replace(" ", "")
            value = match.group(2).strip()
            key = f"{name}:{value.lower()}"
            if key in seen:
                continue
            seen.add(key)
            biomarkers.append({"name": name, "value": value})
        return biomarkers[:12]

    @staticmethod
    def _extract_line_cycle(text: str) -> tuple[int | None, int | None]:
        if not text:
            return None, None
        line_match = _LINE_PATTERN.search(text)
        cycle_match = _CYCLE_PATTERN.search(text)
        line: int | None = None
        cycle: int | None = None
        if line_match:
            candidate = line_match.group(1) or line_match.group(2)
            if candidate and candidate.isdigit():
                line = int(candidate)
        if cycle_match:
            candidate = cycle_match.group(1)
            if candidate and candidate.isdigit():
                cycle = int(candidate)
        return line, cycle

    @staticmethod
    def _resource_date(resource: dict[str, Any]) -> str:
        for key in ("authoredOn", "effectiveDateTime", "performedDateTime", "recordedDate"):
            value = resource.get(key)
            if isinstance(value, str) and value.strip():
                return value[:10]
        effective_period = resource.get("effectivePeriod")
        if isinstance(effective_period, dict):
            start = effective_period.get("start")
            if isinstance(start, str) and start.strip():
                return start[:10]
        performed_period = resource.get("performedPeriod")
        if isinstance(performed_period, dict):
            start = performed_period.get("start")
            if isinstance(start, str) and start.strip():
                return start[:10]
        return ""

    @staticmethod
    def _resource_note_text(resource: dict[str, Any]) -> str:
        notes = resource.get("note")
        if not isinstance(notes, list):
            return ""
        parts: list[str] = []
        for item in notes:
            if not isinstance(item, dict):
                continue
            text = item.get("text")
            if isinstance(text, str) and text.strip():
                parts.append(text.strip())
        return " ".join(parts)

    @staticmethod
    def _resource_dosage_text(resource: dict[str, Any]) -> str:
        dosage = resource.get("dosageInstruction")
        if not isinstance(dosage, list):
            return ""
        parts: list[str] = []
        for item in dosage:
            if not isinstance(item, dict):
                continue
            text = item.get("text")
            if isinstance(text, str) and text.strip():
                parts.append(text.strip())
        return " ".join(parts)

    @staticmethod
    def _extract_medication_name(resource: dict[str, Any]) -> str:
        medication_code = (
            resource.get("medicationCodeableConcept")
            if isinstance(resource.get("medicationCodeableConcept"), dict)
            else {}
        )
        regimen_raw = str(medication_code.get("text") or "").strip()
        if regimen_raw:
            return regimen_raw
        coding = medication_code.get("coding") if isinstance(medication_code.get("coding"), list) else []
        for coding_item in coding:
            if not isinstance(coding_item, dict):
                continue
            regimen_raw = str(coding_item.get("display") or coding_item.get("code") or "").strip()
            if regimen_raw:
                return regimen_raw
        return ""

    @staticmethod
    def _extract_procedure_label(resource: dict[str, Any]) -> str:
        code = resource.get("code") if isinstance(resource.get("code"), dict) else {}
        label = str(code.get("text") or "").strip()
        if label:
            return label
        coding = code.get("coding") if isinstance(code.get("coding"), list) else []
        for coding_item in coding:
            if not isinstance(coding_item, dict):
                continue
            label = str(coding_item.get("display") or coding_item.get("code") or "").strip()
            if label:
                return label
        return "Procedure"

    def _disease_id_from_icd10(self, icd10: str, case_id: str) -> str:
        target = icd10.upper().strip()
        for entry in self.store.list_disease_registry():
            codes = entry.get("icd10_codes")
            if not isinstance(codes, list):
                continue
            for raw_code in codes:
                code = str(raw_code or "").upper().strip()
                if not code:
                    continue
                if target == code or target.startswith(f"{code}.") or code.startswith(f"{target}."):
                    disease_id = str(entry.get("disease_id") or "").strip()
                    if disease_id:
                        return disease_id
        seed = target or case_id
        return str(uuid.uuid5(uuid.NAMESPACE_URL, f"oncoai:disease:{seed}"))

    def _build_case_json_from_fhir_bundle(
        self,
        *,
        payload: dict[str, Any],
        case_id: str,
        now: str,
        data_mode: str,
    ) -> tuple[dict[str, Any], list[str], list[dict[str, str]]]:
        bundle = payload.get("fhir_bundle")
        if not isinstance(bundle, dict):
            raise ValidationError("FHIR_BUNDLE requires `fhir_bundle` object payload")
        if str(bundle.get("resourceType") or "") != "Bundle":
            raise ValidationError("FHIR_BUNDLE requires `fhir_bundle.resourceType` = `Bundle`")
        entries = bundle.get("entry")
        if not isinstance(entries, list) or not entries:
            raise ValidationError("FHIR_BUNDLE requires non-empty `fhir_bundle.entry`")

        resources: list[dict[str, Any]] = []
        for entry in entries:
            if isinstance(entry, dict) and isinstance(entry.get("resource"), dict):
                resources.append(entry["resource"])
        if not resources:
            raise ValidationError("FHIR_BUNDLE has no valid `entry[].resource` objects")

        patient_resource = next(
            (item for item in resources if str(item.get("resourceType") or "") == "Patient"),
            {},
        )
        patient = {
            "sex": self._normalize_sex(patient_resource.get("gender")),
        }
        birth_date = str(patient_resource.get("birthDate") or "")
        if len(birth_date) >= 4 and birth_date[:4].isdigit():
            birth_year = int(birth_date[:4])
            if 1900 <= birth_year <= datetime.now(timezone.utc).year:
                patient["birth_year"] = birth_year

        condition = next(
            (item for item in resources if str(item.get("resourceType") or "") == "Condition"),
            {},
        )
        icd10 = ""
        code_obj = condition.get("code") if isinstance(condition.get("code"), dict) else {}
        coding = code_obj.get("coding") if isinstance(code_obj.get("coding"), list) else []
        for coding_item in coding:
            if not isinstance(coding_item, dict):
                continue
            candidate = self._extract_first_icd10(str(coding_item.get("code") or ""))
            if candidate:
                icd10 = candidate
                break
        if not icd10:
            icd10 = self._extract_first_icd10(str(code_obj.get("text") or ""))

        stage_group = ""
        stage_items = condition.get("stage") if isinstance(condition.get("stage"), list) else []
        for stage_item in stage_items:
            if not isinstance(stage_item, dict):
                continue
            summary = stage_item.get("summary") if isinstance(stage_item.get("summary"), dict) else {}
            stage_group = self._extract_stage_group(str(summary.get("text") or ""))
            if stage_group:
                break

        medication_resources = [
            item
            for item in resources
            if str(item.get("resourceType") or "") in {"MedicationRequest", "MedicationStatement"}
        ]
        medication_candidates: list[dict[str, Any]] = []
        for medication in medication_resources:
            regimen_raw = self._extract_medication_name(medication)
            regimen = self._extract_regimen_name(regimen_raw) or regimen_raw
            note_text = self._resource_note_text(medication)
            dosage_text = self._resource_dosage_text(medication)
            combined_line_text = " ".join(part for part in (regimen_raw, note_text, dosage_text) if part)
            line, cycle = self._extract_line_cycle(combined_line_text)
            date = self._resource_date(medication) or now[:10]
            medication_candidates.append(
                {
                    "date": date,
                    "regimen": regimen,
                    "line": line,
                    "cycle": cycle,
                    "details": " ".join(part for part in (note_text, dosage_text) if part).strip(),
                }
            )
        medication_candidates.sort(key=lambda item: item.get("date") or "")
        selected_medication = medication_candidates[-1] if medication_candidates else {}
        regimen = str(selected_medication.get("regimen") or "").strip()
        line = selected_medication.get("line") if isinstance(selected_medication.get("line"), int) else None
        cycle = selected_medication.get("cycle") if isinstance(selected_medication.get("cycle"), int) else None

        biomarkers: list[dict[str, str]] = []
        for observation in resources:
            if str(observation.get("resourceType") or "") != "Observation":
                continue
            code = observation.get("code") if isinstance(observation.get("code"), dict) else {}
            name = str(code.get("text") or "").strip()
            if not name:
                coding = code.get("coding") if isinstance(code.get("coding"), list) else []
                for coding_item in coding:
                    if not isinstance(coding_item, dict):
                        continue
                    name = str(coding_item.get("display") or coding_item.get("code") or "").strip()
                    if name:
                        break
            if not name:
                continue
            value = ""
            if isinstance(observation.get("valueString"), str):
                value = str(observation.get("valueString"))
            elif isinstance(observation.get("valueCodeableConcept"), dict):
                value = str(observation["valueCodeableConcept"].get("text") or "")
            elif isinstance(observation.get("valueQuantity"), dict):
                raw = observation["valueQuantity"].get("value")
                value = str(raw) if raw is not None else ""
            value = value.strip()
            if not value:
                continue
            biomarkers.append({"name": name, "value": value})
        biomarkers = biomarkers[:12]

        timeline: list[dict[str, Any]] = []
        for procedure in [item for item in resources if str(item.get("resourceType") or "") == "Procedure"]:
            label = self._extract_procedure_label(procedure)
            note_text = self._resource_note_text(procedure)
            combined = f"{label} {note_text}".lower()
            event_type = "surgery"
            if any(marker in combined for marker in ("diagnostic", "биопс", "staging", "laparoscopy")):
                event_type = "diagnostic"
            timeline.append(
                {
                    "event_id": str(uuid.uuid4()),
                    "date": self._resource_date(procedure) or now[:10],
                    "precision": "day",
                    "type": event_type,
                    "label": label,
                    "details": note_text or label,
                }
            )

        for medication in medication_candidates:
            if not medication.get("regimen"):
                continue
            details = str(medication.get("details") or "").strip()
            if medication.get("line") is not None:
                details = f"{details}; line {medication['line']}".strip("; ").strip()
            if medication.get("cycle") is not None:
                details = f"{details}; cycle {medication['cycle']}".strip("; ").strip()
            timeline.append(
                {
                    "event_id": str(uuid.uuid4()),
                    "date": str(medication.get("date") or now[:10])[:10],
                    "precision": "day",
                    "type": "systemic_therapy",
                    "label": str(medication.get("regimen")),
                    "details": details or str(medication.get("regimen")),
                }
            )

        timeline.sort(key=lambda item: str(item.get("date") or ""))
        authored_on = str(selected_medication.get("date") or now[:10])[:10]
        missing_fields: list[str] = []
        if not icd10:
            missing_fields.append("diagnoses[0].icd10")
        if not regimen:
            missing_fields.append("diagnoses[0].last_plan.regimen")

        diagnosis: dict[str, Any] = {
            "diagnosis_id": str(uuid.uuid4()),
            "disease_id": self._disease_id_from_icd10(icd10, case_id),
            "biomarkers": biomarkers,
            "timeline": timeline,
        }
        if icd10:
            diagnosis["icd10"] = icd10
        histology = str(code_obj.get("text") or "").strip()
        if histology:
            diagnosis["histology"] = histology
        if stage_group:
            diagnosis["stage"] = {"system": "UNKNOWN", "stage_group": stage_group}
        if regimen:
            last_plan: dict[str, Any] = {
                "date": authored_on,
                "precision": "day",
                "regimen": regimen,
            }
            if line is not None:
                last_plan["line"] = line
            if cycle is not None:
                last_plan["cycle"] = cycle
            diagnosis["last_plan"] = last_plan

        notes_parts = []
        if isinstance(payload.get("notes"), str) and str(payload.get("notes")).strip():
            notes_parts.append(str(payload["notes"]).strip())
        notes_parts.append("Imported from FHIR bundle")

        attachment_doc_id = str(uuid.uuid5(uuid.NAMESPACE_URL, f"oncoai:case:{case_id}:fhir_bundle"))
        case_json: dict[str, Any] = {
            "schema_version": "1.0",
            "case_id": case_id,
            "created_at": now,
            "updated_at": now,
            "data_mode": data_mode,
            "import_profile": "FHIR_BUNDLE",
            "patient": patient,
            "diagnoses": [diagnosis],
            "attachments": [
                {
                    "doc_id": attachment_doc_id,
                    "kind": "fhir_bundle",
                    "filename": str(payload.get("filename") or "fhir_bundle.json"),
                    "hash": self._hash_json_payload(bundle),
                }
            ],
            "notes": ". ".join(notes_parts),
        }
        warnings: list[dict[str, str]] = []
        if not biomarkers:
            warnings.append({"code": "BIOMARKERS_NOT_FOUND", "message": "No biomarkers extracted from FHIR bundle"})
        return case_json, missing_fields, warnings

    def _extract_kin_text(self, payload: dict[str, Any]) -> str:
        direct = payload.get("kin_pdf_text")
        if isinstance(direct, str) and direct.strip():
            return direct.strip()
        raw = payload.get("kin_pdf")
        if isinstance(raw, dict):
            for key in ("text", "ocr_text", "extracted_text"):
                value = raw.get(key)
                if isinstance(value, str) and value.strip():
                    return value.strip()
            pages = raw.get("pages")
            if isinstance(pages, list):
                parts = [str(item).strip() for item in pages if isinstance(item, str) and item.strip()]
                if parts:
                    return "\n".join(parts)
        return ""

    def _build_case_json_from_kin_pdf(
        self,
        *,
        payload: dict[str, Any],
        case_id: str,
        now: str,
        data_mode: str,
    ) -> tuple[dict[str, Any], list[str], list[dict[str, str]]]:
        text = self._extract_kin_text(payload)
        if not text:
            raise ValidationError("KIN_PDF requires `kin_pdf_text` or `kin_pdf.text`/`kin_pdf.pages`")

        sex = "unknown"
        lowered = text.lower()
        if "муж" in lowered or "male" in lowered:
            sex = "male"
        elif "жен" in lowered or "female" in lowered:
            sex = "female"

        patient: dict[str, Any] = {"sex": sex}
        birth_year = self._extract_year(text)
        if birth_year is not None:
            patient["birth_year"] = birth_year

        icd10 = self._extract_first_icd10(text)
        stage_group = self._extract_stage_group(text)
        regimen = self._extract_regimen_name(text)
        line, cycle = self._extract_line_cycle(text)
        biomarkers = self._extract_biomarkers_from_text(text)
        missing_fields: list[str] = []
        if not icd10:
            missing_fields.append("diagnoses[0].icd10")
        if not regimen:
            missing_fields.append("diagnoses[0].last_plan.regimen")

        diagnosis: dict[str, Any] = {
            "diagnosis_id": str(uuid.uuid4()),
            "disease_id": self._disease_id_from_icd10(icd10, case_id),
            "biomarkers": biomarkers,
            "timeline": [],
        }
        if icd10:
            diagnosis["icd10"] = icd10
        if stage_group:
            diagnosis["stage"] = {"system": "UNKNOWN", "stage_group": stage_group}
        if regimen:
            last_plan: dict[str, Any] = {
                "date": now[:10],
                "precision": "day",
                "regimen": regimen,
            }
            if line is not None:
                last_plan["line"] = line
            if cycle is not None:
                last_plan["cycle"] = cycle
            diagnosis["last_plan"] = last_plan

        attachment_doc_id = str(uuid.uuid5(uuid.NAMESPACE_URL, f"oncoai:case:{case_id}:kin_pdf"))
        filename = str(payload.get("filename") or "kin_case.pdf")
        case_json = {
            "schema_version": "1.0",
            "case_id": case_id,
            "created_at": now,
            "updated_at": now,
            "data_mode": data_mode,
            "import_profile": "KIN_PDF",
            "patient": patient,
            "diagnoses": [diagnosis],
            "attachments": [
                {
                    "doc_id": attachment_doc_id,
                    "kind": "pdf",
                    "filename": filename,
                    "hash": self._hash_text_payload(text),
                }
            ],
            "notes": text[:4000],
        }
        warnings: list[dict[str, str]] = []
        if len(text) < 120:
            warnings.append({"code": "LOW_TEXT_VOLUME", "message": "KIN_PDF text is short; extraction confidence reduced"})
        if not biomarkers:
            warnings.append({"code": "BIOMARKERS_NOT_FOUND", "message": "No biomarkers extracted from KIN_PDF text"})
        return case_json, missing_fields, warnings

    def _resolve_pack_case_reference(self, payload: dict[str, Any]) -> dict[str, Any]:
        if not is_pack_v0_2_request(payload):
            return payload
        case = payload.get("case")
        if not isinstance(case, dict):
            return payload
        if isinstance(case.get("case_json"), dict):
            return payload
        case_id = str(case.get("case_id") or "").strip()
        if not case_id:
            return payload

        stored_case = self.store.get_case(case_id)
        if not stored_case:
            raise ValidationError(f"CASE_NOT_FOUND: {case_id}")

        resolved_payload = copy.deepcopy(payload)
        # Pack v0.2 contract requires exactly one case selector: either case_id or case_json.
        # After lookup by case_id we switch to canonical inline case_json to keep schema oneOf valid.
        resolved_payload["case"] = {"case_json": stored_case}
        return resolved_payload

    def health(self) -> dict[str, Any]:
        vector_preflight = self._build_vector_preflight_snapshot()
        return {
            "status": "ok",
            "service": "oncoai-backend",
            "kb_version": self._kb_version,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "vector_preflight": vector_preflight,
        }

    def _configured_embedding_vector_size(self) -> int | None:
        backend = str(self.settings.embedding_backend or "").strip().lower()
        if backend != "openai":
            return 64
        model = str(self.settings.embedding_model or "").strip().lower()
        if not model:
            return 64
        if model in {"text-embedding-3-small", "text-embedding-ada-002"}:
            return 1536
        if model == "text-embedding-3-large":
            return 3072
        return None

    def _build_vector_preflight_snapshot(self) -> dict[str, Any]:
        backend = str(self.settings.vector_backend or "local").strip().lower()
        if backend != "qdrant":
            return {
                "backend": backend or "local",
                "status": "not_applicable",
            }
        expected_size = self._configured_embedding_vector_size()
        if isinstance(self.index, QdrantRestIndex):
            return self.index.preflight_vector_alignment(expected_vector_size=expected_size)
        return {
            "backend": "qdrant",
            "status": "index_not_qdrant",
            "expected_vector_size": expected_size,
        }

    @staticmethod
    def _sync_summary_with_visible_issues(doctor_report: dict[str, Any]) -> dict[str, Any]:
        report = dict(doctor_report)
        issues_count = len(report.get("issues", []))
        report["summary"] = (
            f"Found {issues_count} potential issue(s) while checking treatment plan against indexed guidance."
        )
        return report

    @staticmethod
    def _build_pack_summary_md(
        *,
        issues: list[dict[str, Any]],
        query_type: str,
        insufficient_status: bool,
    ) -> str:
        critical = 0
        warning = 0
        for issue in issues:
            if not isinstance(issue, dict):
                continue
            severity = str(issue.get("severity") or "").strip().lower()
            if severity == "critical":
                critical += 1
            elif severity == "warning":
                warning += 1

        question_label = "следующих шагов лечения" if query_type == "NEXT_STEPS" else "проверки последнего этапа лечения"
        if critical > 0:
            return (
                f"Выявлены критические замечания ({critical}) по результатам {question_label}. "
                "Требуется приоритетный клинический пересмотр тактики."
            )
        if warning > 0:
            return (
                f"Выявлены предупреждения ({warning}) по результатам {question_label}. "
                "Рекомендуется уточнить план и закрыть дефицитные данные."
            )
        if insufficient_status:
            return (
                f"Критических расхождений не выявлено, но по результатам {question_label} "
                "остаются ограничения из-за неполноты клинических данных."
            )
        return f"Критических расхождений не выявлено по результатам {question_label}."

    @staticmethod
    def _build_verification_summary(
        *,
        issues: list[dict[str, Any]],
        insufficient_status: bool,
    ) -> dict[str, Any]:
        counts = {
            "ok": 0,
            "not_compliant": 0,
            "needs_data": 0,
            "risk": 0,
        }
        for issue in issues:
            if not isinstance(issue, dict):
                continue
            kind = str(issue.get("kind") or "").strip().lower()
            severity = str(issue.get("severity") or "").strip().lower()
            if kind == "contraindication" or severity == "critical":
                counts["risk"] += 1
            elif kind in {"deviation", "inconsistency"}:
                counts["not_compliant"] += 1
            elif kind == "missing_data":
                counts["needs_data"] += 1
            elif kind:
                counts["needs_data"] += 1

        if insufficient_status and counts["needs_data"] == 0:
            counts["needs_data"] = 1

        if sum(counts.values()) == 0:
            counts["ok"] = 1

        if counts["risk"] > 0:
            category = "RISK"
            status_line = "Выявлены риски/критичные вопросы, требующие приоритетного клинического обсуждения."
        elif counts["not_compliant"] > 0:
            category = "NOT_COMPLIANT"
            status_line = "Обнаружены потенциальные несоответствия текущей тактики клиническим рекомендациям."
        elif counts["needs_data"] > 0:
            category = "NEEDS_DATA"
            status_line = "Для окончательной оценки необходимо уточнить недостающие клинические данные."
        else:
            category = "OK"
            status_line = "Существенных расхождений по доступным данным не обнаружено."

        return {
            "category": category,
            "status_line": status_line,
            "counts": counts,
        }

    def _clear_routing_cache(self) -> None:
        self._routing_cache.clear()

    @staticmethod
    def _routing_cache_key(case_payload: dict[str, Any], language: str, requested_source_ids: list[str]) -> str:
        diagnosis = case_payload.get("diagnosis") if isinstance(case_payload.get("diagnosis"), dict) else {}
        notes = str(case_payload.get("notes") or "")[:512]
        histology = str(diagnosis.get("histology") or "")[:128]
        icd10 = str(diagnosis.get("icd10") or "")[:32]
        disease_id = str(diagnosis.get("disease_id") or "")[:64]
        cancer_type = str(case_payload.get("cancer_type") or "")[:64]
        sources = ",".join(sorted(normalize_source_set_ids([str(item).strip() for item in requested_source_ids if str(item).strip()])))
        payload = "|".join(
            [
                language,
                icd10,
                disease_id,
                cancer_type,
                histology,
                notes,
                sources,
            ]
        )
        digest = hashlib.sha256(payload.encode("utf-8")).hexdigest()
        return f"route:{digest}"

    def _resolve_nosology_route_with_cache(
        self,
        *,
        case_payload: dict[str, Any],
        language: str,
        requested_source_ids: list[str],
    ) -> NosologyRouteDecision:
        now = time.time()
        cache_key = self._routing_cache_key(
            case_payload=case_payload,
            language=language,
            requested_source_ids=requested_source_ids,
        )
        cached = self._routing_cache.get(cache_key)
        if cached:
            timestamp, decision = cached
            if now - timestamp < self._routing_cache_ttl_seconds:
                return decision
            self._routing_cache.pop(cache_key, None)

        if self._llm_rag_only_mode:
            active_routes = self.store.list_nosology_routes(language=language, active_only=True)
            decision = plan_nosology_route_with_llm(
                llm_router=self.llm_router,
                case_payload=case_payload,
                language=language,
                requested_source_ids=requested_source_ids,
                available_routes=active_routes,
            )
        else:
            decision = resolve_nosology_route(
                store=self.store,
                case_payload=case_payload,
                language=language,
                requested_source_ids=requested_source_ids,
            )
        self._routing_cache[cache_key] = (now, decision)
        return decision

    @staticmethod
    def _normalize_string_list(value: Any, *, uppercase: bool = False, lowercase: bool = False) -> list[str]:
        raw_list: list[Any]
        if isinstance(value, list):
            raw_list = value
        elif isinstance(value, str):
            raw_list = [item.strip() for item in value.split(",")]
        else:
            raw_list = []
        normalized: list[str] = []
        for item in raw_list:
            text = str(item).strip()
            if not text:
                continue
            if uppercase:
                text = text.upper()
            if lowercase:
                text = text.lower()
            normalized.append(text)
        deduped: list[str] = []
        seen: set[str] = set()
        for item in normalized:
            if item in seen:
                continue
            seen.add(item)
            deduped.append(item)
        return deduped

    @staticmethod
    def _icd10_prefix_from_cancer_type(cancer_type: str) -> str:
        normalized = str(cancer_type or "").strip().lower()
        mapping = {
            "gastric_cancer": "C16",
            "nsclc_egfr": "C34",
            "breast_hr+/her2-": "C50",
        }
        return mapping.get(normalized, "")

    @staticmethod
    def _cancer_type_from_icd10_prefix(icd10_prefix: str) -> str:
        normalized = str(icd10_prefix or "").strip().upper()
        mapping = {
            "C16": "gastric_cancer",
            "C34": "nsclc_egfr",
            "C50": "breast_hr+/her2-",
        }
        mapped = mapping.get(normalized)
        if mapped:
            return mapped
        if re.fullmatch(r"C\d{2}", normalized):
            return f"oncology_{normalized.lower()}"
        return "unknown"

    def _derive_upload_route_hints(
        self,
        *,
        metadata: dict[str, Any],
    ) -> tuple[list[str], list[str], str]:
        disease_id = str(metadata.get("disease_id") or "").strip()
        icd10_prefixes = self._normalize_string_list(metadata.get("icd10_prefixes"), uppercase=True)
        keywords = self._normalize_string_list(metadata.get("nosology_keywords"), lowercase=True)
        cancer_type = str(metadata.get("cancer_type") or "").strip()

        if not icd10_prefixes:
            guessed_prefix = self._icd10_prefix_from_cancer_type(cancer_type)
            if guessed_prefix:
                icd10_prefixes = [guessed_prefix]

        diseases = self.store.list_disease_registry()
        matched_disease: dict[str, Any] | None = None
        if disease_id:
            matched_disease = next((item for item in diseases if str(item.get("disease_id") or "") == disease_id), None)
        elif icd10_prefixes:
            for candidate in diseases:
                codes = candidate.get("icd10_codes") if isinstance(candidate.get("icd10_codes"), list) else []
                normalized_codes = {str(code).strip().upper().split(".")[0] for code in codes if str(code).strip()}
                if any(prefix in normalized_codes for prefix in icd10_prefixes):
                    matched_disease = candidate
                    break

        if matched_disease and not disease_id:
            disease_id = str(matched_disease.get("disease_id") or "").strip()
        if matched_disease and not keywords:
            synonyms = matched_disease.get("common_synonyms") if isinstance(matched_disease.get("common_synonyms"), list) else []
            keywords = self._normalize_string_list(synonyms, lowercase=True)

        if not keywords and cancer_type:
            keywords = [cancer_type.replace("_", " ").lower()]

        return icd10_prefixes, keywords, disease_id

    def _upsert_doc_nosology_routes(
        self,
        *,
        metadata: dict[str, Any],
        uploaded_at: str,
    ) -> int:
        source_id = normalize_source_set_id(str(metadata.get("source_set") or ""))
        doc_id = str(metadata.get("doc_id") or "").strip()
        language = str(metadata.get("language") or "ru").strip().lower() or "ru"
        cancer_type = str(metadata.get("cancer_type") or "").strip() or "unknown"
        if not source_id or not doc_id:
            return 0

        icd10_prefixes, keywords, disease_id = self._derive_upload_route_hints(metadata=metadata)
        if not icd10_prefixes and not keywords:
            return 0

        normalized_cancer_type = cancer_type
        if normalized_cancer_type.lower() in {"", "unknown"}:
            for prefix in icd10_prefixes:
                mapped = self._cancer_type_from_icd10_prefix(prefix)
                if mapped != "unknown":
                    normalized_cancer_type = mapped
                    break
        if normalized_cancer_type.lower() in {"", "unknown"} and disease_id:
            lowered_disease = disease_id.lower()
            if "gastric" in lowered_disease:
                normalized_cancer_type = "gastric_cancer"
            elif "breast" in lowered_disease:
                normalized_cancer_type = "breast_hr+/her2-"
            elif "lung" in lowered_disease or "nsclc" in lowered_disease:
                normalized_cancer_type = "nsclc_egfr"
        if not normalized_cancer_type:
            normalized_cancer_type = "unknown"
        if str(normalized_cancer_type).strip().lower() in {"", "unknown", "other"}:
            normalized_cancer_type = "general_oncology"

        prefixes = icd10_prefixes or ["*"]
        keyword_values = keywords or ["*"]
        created = 0
        for prefix in prefixes:
            for keyword in keyword_values:
                route_id = str(
                    uuid.uuid5(
                        uuid.NAMESPACE_URL,
                        f"oncoai:route:{language}:{prefix}:{keyword}:{disease_id}:{normalized_cancer_type}:{source_id}:{doc_id}",
                    )
                )
                self.store.upsert_nosology_route(
                    {
                        "route_id": route_id,
                        "language": language,
                        "icd10_prefix": prefix,
                        "keyword": keyword,
                        "disease_id": disease_id or "unknown_disease",
                        "cancer_type": normalized_cancer_type,
                        "source_id": source_id,
                        "doc_id": doc_id,
                        "priority": 100,
                        "active": 1,
                        "updated_at": uploaded_at,
                    }
                )
                created += 1
        return created

    def _extract_case_texts(self, payload: dict[str, Any]) -> list[str]:
        case = payload.get("case", {})
        treatment = payload.get("treatment_plan", {})
        diagnosis = case.get("diagnosis") if isinstance(case.get("diagnosis"), dict) else {}
        return [
            str(case.get("notes", "")),
            str(treatment.get("plan_text", "")),
            json.dumps(case.get("patient", {}), ensure_ascii=False),
            str(diagnosis.get("histology", "")),
            str(diagnosis.get("icd10", "")),
        ]

    def _redact_analyze_payload_pii(self, payload: dict[str, Any]) -> bool:
        changed = False
        case = payload.get("case")
        if not isinstance(case, dict):
            return False

        def redact_text(value: Any) -> str:
            nonlocal changed
            original = str(value or "")
            masked = redact_pii(original)
            if masked != original:
                changed = True
            return masked

        if isinstance(case.get("notes"), str):
            case["notes"] = redact_text(case.get("notes"))

        patient = case.get("patient")
        if isinstance(patient, dict):
            for key, value in list(patient.items()):
                if isinstance(value, str):
                    patient[key] = redact_text(value)

        diagnosis = case.get("diagnosis")
        if isinstance(diagnosis, dict):
            for field in ("histology", "icd10", "disease_id"):
                if isinstance(diagnosis.get(field), str):
                    diagnosis[field] = redact_text(diagnosis.get(field))

        treatment_plan = payload.get("treatment_plan")
        if isinstance(treatment_plan, dict):
            if isinstance(treatment_plan.get("plan_text"), str):
                treatment_plan["plan_text"] = redact_text(treatment_plan.get("plan_text"))
            plan_structured = treatment_plan.get("plan_structured")
            if isinstance(plan_structured, list):
                for step in plan_structured:
                    if not isinstance(step, dict):
                        continue
                    for field in ("name", "details", "description"):
                        if isinstance(step.get(field), str):
                            step[field] = redact_text(step.get(field))

        return changed

    @staticmethod
    def _decode_content_base64(value: Any) -> bytes:
        payload = str(value or "").strip()
        if not payload:
            raise ValidationError("content_base64 is required")
        try:
            return base64.b64decode(payload, validate=True)
        except Exception as exc:  # noqa: BLE001
            raise ValidationError("Invalid content_base64 payload") from exc

    @staticmethod
    def _guess_mime_type(filename: str, mime_type: str) -> str:
        normalized_mime = str(mime_type or "").strip().lower()
        if normalized_mime:
            return normalized_mime
        name = str(filename or "").strip().lower()
        if name.endswith(".pdf"):
            return "application/pdf"
        if name.endswith(".docx"):
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if name.endswith(".md"):
            return "text/markdown"
        return "text/plain"

    def _build_case_import_payload_from_file(self, payload: dict[str, Any]) -> dict[str, Any]:
        filename = str(payload.get("filename") or "").strip()
        if not filename:
            raise ValidationError("filename is required")
        content = self._decode_content_base64(payload.get("content_base64"))
        mime_type = self._guess_mime_type(filename=filename, mime_type=str(payload.get("mime_type") or ""))
        extraction = extract_text(content, filename=filename, mime=mime_type)

        file_kind = str(extraction.get("file_kind") or "").strip().lower()
        import_profile = "KIN_PDF" if file_kind == "pdf" else "FREE_TEXT"
        normalized_payload: dict[str, Any] = {
            "schema_version": "1.0",
            "import_profile": import_profile,
            "filename": filename,
            "data_mode": str(payload.get("data_mode") or "DEID"),
        }
        if payload.get("case_id"):
            normalized_payload["case_id"] = str(payload.get("case_id"))

        extracted_text = str(extraction.get("text") or "").strip()
        if import_profile == "KIN_PDF":
            normalized_payload["kin_pdf_text"] = extracted_text
            normalized_payload["kin_pdf"] = {
                "pages": [extracted_text],
                "page_map": extraction.get("page_map") if isinstance(extraction.get("page_map"), dict) else {},
                "sha256": extraction.get("sha256"),
            }
        else:
            normalized_payload["free_text"] = extracted_text
            normalized_payload["notes"] = f"Imported from {filename}"
        extraction_warnings = extraction.get("warnings")
        if isinstance(extraction_warnings, list):
            normalized_payload["__file_warnings"] = [
                {"code": str(item.get("code") or "").strip(), "message": str(item.get("message") or "").strip()}
                for item in extraction_warnings
                if isinstance(item, dict) and str(item.get("code") or "").strip() and str(item.get("message") or "").strip()
            ]

        return normalized_payload

    def _step_1_validate_and_guard_input(
        self,
        payload: dict[str, Any],
        role: str,
        client_id: str,
    ) -> str:
        ensure_role(role, {"clinician", "admin"})
        self.rate_limiter.check(client_id, route="/analyze")
        validate_analyze_request(payload)
        schema_version = str(payload.get("schema_version", "0.1"))
        case = payload.get("case") if isinstance(payload.get("case"), dict) else {}
        data_mode = self._normalize_case_data_mode(case.get("data_mode"))
        if data_mode == "FULL" and not bool(getattr(self.settings, "case_import_allow_full_mode", False)):
            raise ValidationError(
                "FULL_MODE_DISABLED: configure CASE_IMPORT_ALLOW_FULL_MODE=true for private FULL-mode deployment"
            )
        if data_mode != "FULL":
            if contains_pii(self._extract_case_texts(payload)):
                self._redact_analyze_payload_pii(payload)
                if contains_pii(self._extract_case_texts(payload)):
                    raise ValidationError("PII_DETECTED: remove personal data before analysis")
        return schema_version

    @staticmethod
    def _parse_iso_date_token(value: Any) -> date | None:
        token = str(value or "").strip()
        if not token:
            return None
        candidates = [token]
        if len(token) >= 10:
            short = token[:10]
            if short != token:
                candidates.append(short)
        for candidate in candidates:
            try:
                return date.fromisoformat(candidate)
            except ValueError:
                continue
        return None

    @staticmethod
    def _extract_version_effective_window(version_row: dict[str, Any] | None) -> tuple[date | None, date | None]:
        payload = version_row if isinstance(version_row, dict) else {}
        metadata = payload.get("metadata") if isinstance(payload.get("metadata"), dict) else {}
        raw_from = (
            payload.get("effective_from")
            or metadata.get("effective_from")
            or metadata.get("valid_from")
            or ""
        )
        raw_to = (
            payload.get("effective_to")
            or metadata.get("effective_to")
            or metadata.get("valid_to")
            or ""
        )
        return (
            OncoService._parse_iso_date_token(raw_from),
            OncoService._parse_iso_date_token(raw_to),
        )

    def _is_version_release_ready_for_retrieval(
        self,
        *,
        source_id: str,
        doc_id: str,
        cancer_type: str,
        version_row: dict[str, Any] | None,
    ) -> bool:
        normalized_source = normalize_source_set_id(source_id)
        if normalized_source not in OFFICIAL_SOURCE_RULES:
            return True
        payload = version_row if isinstance(version_row, dict) else {}
        metadata = payload.get("metadata") if isinstance(payload.get("metadata"), dict) else {}
        source_url = resolve_primary_source_url(
            source_url=str(metadata.get("source_url") or ""),
            source_page_url=str(metadata.get("source_page_url") or ""),
            source_pdf_url=str(metadata.get("source_pdf_url") or ""),
        )
        validity = evaluate_release_validity(
            source_set=normalized_source,
            source_url=source_url,
            status=str(payload.get("status") or metadata.get("status") or ""),
            doc_id=str(doc_id or ""),
            nosology_mapped=is_nosology_mapped(str(cancer_type or "")),
        )
        return bool(validity.get("is_valid"))

    def _pick_doc_version_for_as_of(
        self,
        *,
        docs: list[dict[str, Any]],
        source_id: str,
        doc_id: str,
        as_of_date: date,
    ) -> dict[str, Any] | None:
        candidates: list[tuple[tuple[date, date, str], dict[str, Any]]] = []
        for doc in docs:
            if not isinstance(doc, dict):
                continue
            candidate_doc_id = str(doc.get("doc_id") or "").strip()
            candidate_doc_version = str(doc.get("doc_version") or "").strip()
            if not candidate_doc_id or not candidate_doc_version:
                continue
            version_row = self.store.get_guideline_version_by_doc(
                doc_id=candidate_doc_id,
                doc_version=candidate_doc_version,
            )
            if not self._is_version_release_ready_for_retrieval(
                source_id=source_id,
                doc_id=doc_id,
                cancer_type=str(doc.get("cancer_type") or ""),
                version_row=version_row,
            ):
                continue
            effective_from, effective_to = self._extract_version_effective_window(version_row)
            if effective_from and as_of_date < effective_from:
                continue
            if effective_to and as_of_date > effective_to:
                continue
            uploaded_date = self._parse_iso_date_token(doc.get("uploaded_at")) or date.min
            sort_key = (
                effective_from or uploaded_date,
                uploaded_date,
                candidate_doc_version,
            )
            candidates.append(
                (
                    sort_key,
                    {
                        "source_id": source_id,
                        "doc_id": candidate_doc_id,
                        "doc_version": candidate_doc_version,
                        "effective_from": effective_from.isoformat() if effective_from else "",
                        "effective_to": effective_to.isoformat() if effective_to else "",
                    },
                )
            )
        if not candidates:
            return None
        candidates.sort(key=lambda item: item[0], reverse=True)
        return candidates[0][1]

    def _build_guideline_snapshot_for_date(
        self,
        *,
        language: str,
        route_pairs: list[tuple[str, str]],
        source_sets: list[str],
        doc_ids: list[str],
        as_of_date: date,
    ) -> dict[str, Any]:
        docs = [
            doc
            for doc in self.store.list_docs()
            if str(doc.get("language") or "").strip().lower() == language
        ]
        allowed_sources = set(normalize_source_set_ids(source_sets))
        allowed_doc_ids = {str(item).strip() for item in doc_ids if str(item).strip()}
        docs_by_pair: dict[tuple[str, str], list[dict[str, Any]]] = {}
        for doc in docs:
            source_id = normalize_source_set_id(str(doc.get("source_set") or ""))
            doc_id = str(doc.get("doc_id") or "").strip()
            pair = (source_id, doc_id)
            if not source_id or not doc_id:
                continue
            if allowed_sources and source_id not in allowed_sources:
                continue
            if allowed_doc_ids and doc_id not in allowed_doc_ids:
                continue
            docs_by_pair.setdefault(pair, []).append(doc)

        requested_pairs: list[tuple[str, str]]
        if route_pairs:
            requested_pairs = []
            seen_pairs: set[tuple[str, str]] = set()
            for source_id, doc_id in route_pairs:
                pair = (normalize_source_set_id(str(source_id)), str(doc_id).strip())
                if not pair[0] or not pair[1] or pair in seen_pairs:
                    continue
                seen_pairs.add(pair)
                requested_pairs.append(pair)
        else:
            requested_pairs = sorted(docs_by_pair.keys())

        selected_versions: list[dict[str, Any]] = []
        missing_pairs: list[dict[str, str]] = []
        for source_id, doc_id in requested_pairs:
            selected = self._pick_doc_version_for_as_of(
                docs=docs_by_pair.get((source_id, doc_id), []),
                source_id=source_id,
                doc_id=doc_id,
                as_of_date=as_of_date,
            )
            if selected is None:
                missing_pairs.append({"source_id": source_id, "doc_id": doc_id})
                continue
            selected_versions.append(selected)

        selected_source_ids = sorted(
            normalize_source_set_ids(
                [str(item.get("source_id") or "").strip() for item in selected_versions if str(item.get("source_id") or "").strip()]
            )
        )
        selected_doc_ids = sorted(
            {
                str(item.get("doc_id") or "").strip()
                for item in selected_versions
                if str(item.get("doc_id") or "").strip()
            }
        )
        route_triples = [
            (
                normalize_source_set_id(str(item.get("source_id") or "")),
                str(item.get("doc_id") or "").strip(),
                str(item.get("doc_version") or "").strip(),
            )
            for item in selected_versions
            if str(item.get("source_id") or "").strip()
            and str(item.get("doc_id") or "").strip()
            and str(item.get("doc_version") or "").strip()
        ]

        return {
            "as_of_date": as_of_date.isoformat(),
            "requested_pairs": requested_pairs,
            "selected_versions": selected_versions,
            "route_triples": route_triples,
            "source_ids": selected_source_ids,
            "doc_ids": selected_doc_ids,
            "missing_pairs": missing_pairs,
            "has_coverage": bool(route_triples) and not missing_pairs,
        }

    def _step_2_normalize_and_prepare_retrieval(
        self,
        payload: dict[str, Any],
        *,
        as_of_date: str | None = None,
        historical_assessment_requested: bool = False,
    ) -> tuple[
        list[dict[str, Any]],
        dict[str, Any],
        str,
        list[str],
        list[str],
        list[tuple[str, str]],
        list[tuple[str, str, str]],
        dict[str, Any],
    ]:
        treatment_plan = payload["treatment_plan"]
        plan_structured = normalize_plan(
            plan_text=treatment_plan.get("plan_text", ""),
            existing=treatment_plan.get("plan_structured"),
        )

        case_payload = payload["case"] if isinstance(payload.get("case"), dict) else {}
        language = str(case_payload.get("language") or "ru").strip().lower()
        if language not in {"ru", "en"}:
            language = "ru"

        kb_filters = payload.get("kb_filters") or {}
        requested_source_sets: list[str] = []
        raw_source_sets = kb_filters.get("source_sets") if isinstance(kb_filters, dict) else None
        if isinstance(raw_source_sets, list):
            requested_source_sets = normalize_source_set_ids([str(item).strip() for item in raw_source_sets if str(item).strip()])
        if not requested_source_sets and isinstance(kb_filters, dict):
            source_set = str(kb_filters.get("source_set") or "").strip()
            if source_set:
                requested_source_sets = [normalize_source_set_id(source_set)]
        source_mode = str(kb_filters.get("source_mode") or "").strip().upper() if isinstance(kb_filters, dict) else ""
        router_requested_sources = requested_source_sets
        if source_mode == "AUTO":
            # AUTO mode keeps source_ids as an allowlist for retrieval, but routing itself
            # should not be marked as a manual override.
            router_requested_sources = []

        try:
            decision = self._resolve_nosology_route_with_cache(
                case_payload=case_payload,
                language=language,
                requested_source_ids=router_requested_sources,
            )
        except RuntimeError as exc:
            if self._llm_rag_only_mode:
                raise ValidationError(f"STRICT_LLM_RAG_ONLY: route planning failed: {exc}") from exc
            raise
        ambiguous_brain_scope = str(decision.match_strategy or "").strip().lower() == "ambiguous_brain_scope"
        resolved_cancer_type = str(
            decision.resolved_cancer_type
            or case_payload.get("cancer_type")
            or "unknown"
        ).strip()
        if resolved_cancer_type and isinstance(case_payload, dict):
            case_payload["cancer_type"] = resolved_cancer_type

        filters = {
            "cancer_type": resolved_cancer_type,
            "language": language,
        }

        explicit_doc_ids: list[str] = []
        if isinstance(kb_filters, dict):
            raw_doc_ids = kb_filters.get("doc_ids")
            if isinstance(raw_doc_ids, list):
                explicit_doc_ids = [str(item).strip() for item in raw_doc_ids if str(item).strip()]

        decision_source_sets = normalize_source_set_ids([str(item).strip() for item in decision.source_ids if str(item).strip()])
        if requested_source_sets:
            requested_allowlist = set(requested_source_sets)
            source_sets = [item for item in decision_source_sets if item in requested_allowlist]
        else:
            source_sets = decision_source_sets
        doc_ids = explicit_doc_ids or [str(item).strip() for item in decision.doc_ids if str(item).strip()]
        route_pairs = [
            (normalize_source_set_id(str(source_id)), str(doc_id).strip())
            for source_id, doc_id in decision.route_pairs
            if str(source_id).strip() and str(doc_id).strip()
        ]
        if source_sets:
            allowed_sources = set(source_sets)
            route_pairs = [pair for pair in route_pairs if pair[0] in allowed_sources]
        if doc_ids:
            allowed_docs = set(doc_ids)
            route_pairs = [pair for pair in route_pairs if pair[1] in allowed_docs]
        if not route_pairs and source_sets and doc_ids:
            route_pairs = [(source_set, doc_id) for source_set in source_sets for doc_id in doc_ids]
        if ambiguous_brain_scope:
            source_sets = []
            doc_ids = []
            route_pairs = []
        no_release_ready_docs = bool(requested_source_sets) and not route_pairs and not doc_ids
        if no_release_ready_docs or ambiguous_brain_scope:
            # Keep retrieval constrained to an impossible doc_id rather than silently
            # broadening to every chunk in requested sources.
            filters["doc_id"] = "__ambiguous_brain_scope__" if ambiguous_brain_scope else "__no_release_ready_docs__"

        if isinstance(kb_filters, dict):
            passthrough_filter_keys = {"doc_version"}
            for key, value in kb_filters.items():
                if not value or key in {"source_sets", "source_set", "doc_ids", "source_mode"}:
                    continue
                if key not in passthrough_filter_keys:
                    continue
                if isinstance(value, str):
                    trimmed = value.strip()
                    if trimmed:
                        filters[key] = trimmed
                elif isinstance(value, (int, float)):
                    filters[key] = value

        route_triples: list[tuple[str, str, str]] = []
        historical_meta: dict[str, Any] | None = None
        requested_historical_date = self._parse_iso_date_token(as_of_date)
        today_date = date.today()
        if requested_historical_date is not None and requested_historical_date <= today_date:
            current_snapshot = self._build_guideline_snapshot_for_date(
                language=language,
                route_pairs=route_pairs,
                source_sets=source_sets,
                doc_ids=doc_ids,
                as_of_date=today_date,
            )
            as_of_snapshot = self._build_guideline_snapshot_for_date(
                language=language,
                route_pairs=route_pairs,
                source_sets=source_sets,
                doc_ids=doc_ids,
                as_of_date=requested_historical_date,
            )
            route_triples = [
                (
                    normalize_source_set_id(str(source_id)),
                    str(doc_id).strip(),
                    str(doc_version).strip(),
                )
                for source_id, doc_id, doc_version in as_of_snapshot.get("route_triples", [])
                if str(source_id).strip() and str(doc_id).strip() and str(doc_version).strip()
            ]
            if route_triples:
                source_sets = sorted({source_id for source_id, _doc_id, _doc_version in route_triples if source_id})
                doc_ids = sorted({doc_id for _source_id, doc_id, _doc_version in route_triples if doc_id})
                route_pairs = sorted({(source_id, doc_id) for source_id, doc_id, _doc_version in route_triples})
                filters.pop("doc_version", None)
            else:
                # Keep retrieval constrained when historical snapshot has no usable versions.
                filters["doc_id"] = "__historical_sources_unavailable__"
            historical_meta = {
                "requested_as_of_date": requested_historical_date.isoformat(),
                "current_snapshot": current_snapshot,
                "as_of_snapshot": as_of_snapshot,
                "as_of_filter_applied": bool(route_triples),
            }
        elif historical_assessment_requested:
            current_snapshot = self._build_guideline_snapshot_for_date(
                language=language,
                route_pairs=route_pairs,
                source_sets=source_sets,
                doc_ids=doc_ids,
                as_of_date=today_date,
            )
            historical_meta = {
                "requested_as_of_date": str(as_of_date or "").strip() or None,
                "current_snapshot": current_snapshot,
                "as_of_snapshot": {},
                "as_of_filter_applied": False,
                "future_requested": bool(requested_historical_date and requested_historical_date > today_date),
            }

        query = treatment_plan.get("plan_text", "")
        routing_meta = {
            "resolved_disease_id": str(decision.resolved_disease_id or "unknown_disease"),
            "resolved_cancer_type": resolved_cancer_type or "unknown",
            "match_strategy": str(decision.match_strategy or "default_sources_fallback"),
            "source_ids": sorted({item for item in source_sets if item}),
            "doc_ids": sorted({item for item in doc_ids if item}),
            "candidate_chunks": 0,
            "no_release_ready_docs": no_release_ready_docs,
            "ambiguous_brain_scope": ambiguous_brain_scope,
        }
        if historical_meta is not None:
            routing_meta["historical"] = historical_meta
        return plan_structured, filters, query, source_sets, doc_ids, route_pairs, route_triples, routing_meta

    def _step_3_retrieve_rerank_and_probe_llm(
        self,
        query: str,
        query_bundle: list[str] | None,
        filters: dict[str, Any],
        source_sets: list[str] | None = None,
        doc_ids: list[str] | None = None,
        route_pairs: list[tuple[str, str]] | None = None,
        route_triples: list[tuple[str, str, str]] | None = None,
    ) -> tuple[list[dict[str, Any]], list[dict[str, Any]], str, int]:
        best_by_chunk: dict[str, dict[str, Any]] = {}
        queries = [str(item).strip() for item in (query_bundle or []) if str(item).strip()]
        if not queries:
            queries = [str(query or "").strip() or "next treatment steps"]
        elif str(query or "").strip() and str(query).strip() not in queries:
            queries.insert(0, str(query).strip())

        def retrieve_with_cancer_fallback(scoped_filters: dict[str, Any], *, query_text: str) -> list[dict[str, Any]]:
            items = self.retriever.retrieve(query=query_text, filters=scoped_filters)
            if self._llm_rag_only_mode:
                return items
            if items:
                return items
            if not scoped_filters.get("cancer_type"):
                return items
            relaxed_filters = dict(scoped_filters)
            relaxed_filters.pop("cancer_type", None)
            return self.retriever.retrieve(query=query_text, filters=relaxed_filters)

        def collect(items: list[dict[str, Any]]) -> None:
            for item in items:
                chunk_id = str(item.get("chunk_id") or "")
                if not chunk_id:
                    continue
                current = best_by_chunk.get(chunk_id)
                current_score = float(current.get("score", 0.0)) if current else float("-inf")
                item_score = float(item.get("score", 0.0))
                if current is None or item_score > current_score:
                    best_by_chunk[chunk_id] = item

        routed_scope_count = 0
        if route_triples:
            routed_scope_count = len(route_triples)
        elif route_pairs:
            routed_scope_count = len(route_pairs)
        elif source_sets and doc_ids:
            routed_scope_count = len(source_sets) * len(doc_ids)
        routed_per_scope_limit = self.settings.retrieval_top_k
        if routed_scope_count > 0:
            routed_per_scope_limit = max(3, int(self.settings.retrieval_top_k / max(1, routed_scope_count)))

        if route_triples:
            for source_set, doc_id, doc_version in route_triples:
                scoped_filters = dict(filters)
                scoped_filters["source_set"] = source_set
                scoped_filters["doc_id"] = doc_id
                scoped_filters["doc_version"] = doc_version
                for query_text in queries:
                    collect(retrieve_with_cancer_fallback(scoped_filters, query_text=query_text)[:routed_per_scope_limit])
        elif route_pairs:
            for source_set, doc_id in route_pairs:
                scoped_filters = dict(filters)
                scoped_filters["source_set"] = source_set
                scoped_filters["doc_id"] = doc_id
                for query_text in queries:
                    collect(retrieve_with_cancer_fallback(scoped_filters, query_text=query_text)[:routed_per_scope_limit])
        elif source_sets and doc_ids:
            for source_set in source_sets:
                for doc_id in doc_ids:
                    scoped_filters = dict(filters)
                    scoped_filters["source_set"] = source_set
                    scoped_filters["doc_id"] = doc_id
                    for query_text in queries:
                        collect(
                            retrieve_with_cancer_fallback(scoped_filters, query_text=query_text)[:routed_per_scope_limit]
                        )
        elif source_sets:
            for source_set in source_sets:
                scoped_filters = dict(filters)
                scoped_filters["source_set"] = source_set
                for query_text in queries:
                    collect(retrieve_with_cancer_fallback(scoped_filters, query_text=query_text))
        elif doc_ids:
            for doc_id in doc_ids:
                scoped_filters = dict(filters)
                scoped_filters["doc_id"] = doc_id
                for query_text in queries:
                    collect(retrieve_with_cancer_fallback(scoped_filters, query_text=query_text))
        else:
            for query_text in queries:
                collect(retrieve_with_cancer_fallback(dict(filters), query_text=query_text))

        retrieved = sorted(best_by_chunk.values(), key=lambda row: float(row.get("score", 0.0)), reverse=True)
        reranked = self.reranker.rerank(query=query, retrieved=retrieved)

        baseline_candidates = len(retrieved)
        has_routing_narrowing = bool(route_triples or route_pairs or (source_sets and doc_ids))
        if has_routing_narrowing:
            baseline_by_chunk: dict[str, dict[str, Any]] = {}

            def collect_baseline(items: list[dict[str, Any]]) -> None:
                for item in items:
                    chunk_id = str(item.get("chunk_id") or "")
                    if not chunk_id:
                        continue
                    current = baseline_by_chunk.get(chunk_id)
                    current_score = float(current.get("score", 0.0)) if current else float("-inf")
                    item_score = float(item.get("score", 0.0))
                    if current is None or item_score > current_score:
                        baseline_by_chunk[chunk_id] = item

            baseline_sources = normalize_source_set_ids([str(item).strip() for item in (source_sets or []) if str(item).strip()])
            if not baseline_sources and route_pairs:
                baseline_sources = sorted(normalize_source_set_ids([str(source).strip() for source, _doc in route_pairs if str(source).strip()]))
            if not baseline_sources and route_triples:
                baseline_sources = sorted(
                    normalize_source_set_ids([str(source).strip() for source, _doc, _doc_version in route_triples if str(source).strip()])
                )
            if baseline_sources:
                for source_set in baseline_sources:
                    baseline_filters = dict(filters)
                    baseline_filters["source_set"] = source_set
                    for query_text in queries:
                        collect_baseline(retrieve_with_cancer_fallback(baseline_filters, query_text=query_text))
                baseline_candidates = max(len(baseline_by_chunk), len(retrieved))

        # Optional online LLM probe; disabled by default for predictable latency.
        llm_path = "deterministic"
        if self.settings.llm_probe_enabled and (self.llm_router.primary or self.llm_router.fallback):
            _unused_llm_result, llm_path = self.llm_router.generate_json(prompt=queries[0])
        safe_log(self.logger, "analyze.llm_path", {"path": llm_path})
        return retrieved, reranked, llm_path, baseline_candidates

    @staticmethod
    def _pack_issue_severity(value: str) -> str:
        token = str(value or "").strip().lower()
        if token == "critical":
            return "critical"
        if token in {"important", "warning", "warn"}:
            return "warning"
        return "info"

    @staticmethod
    def _pack_issue_kind(*, category: str, title: str) -> str:
        category_lc = str(category or "").strip().lower()
        title_lc = str(title or "").strip().lower()
        if "contra" in category_lc or "contra" in title_lc:
            return "contraindication"
        if "incons" in category_lc or "incons" in title_lc:
            return "inconsistency"
        if "missing" in category_lc or "missing" in title_lc or "data" in category_lc:
            return "missing_data"
        if category_lc and category_lc not in {"other", "note"}:
            return "deviation"
        return "other"

    @staticmethod
    def _pack_issue_id(seed: str) -> str:
        return str(uuid.uuid5(uuid.NAMESPACE_URL, f"oncoai:issue:{seed}"))

    @staticmethod
    def _map_generation_path_for_pack(path: str) -> str:
        normalized = str(path or "").strip().lower()
        if normalized in {"llm_primary", "primary"}:
            return "primary"
        if normalized in {"llm_fallback", "fallback"}:
            return "fallback"
        return "deterministic_only"

    @staticmethod
    def _map_fallback_reason_for_pack(reason: str | None) -> str:
        value = str(reason or "").strip().lower()
        if not value:
            return "none"
        if "no_docs" in value or "no retrieved" in value:
            return "no_docs"
        if "low_recall" in value:
            return "low_recall"
        if "not_configured" in value:
            return "llm_not_configured"
        if "no_valid_response" in value:
            return "llm_no_valid_response"
        if "invalid_json" in value or "invalid_response" in value:
            return "llm_invalid_json"
        if value.startswith("llm_"):
            return value
        if "timeout" in value:
            return "timeout"
        return "other"

    @staticmethod
    def _resolve_run_meta_llm_path(llm_path: str, report_generation_path: str) -> str:
        normalized_llm = str(llm_path or "").strip().lower()
        if normalized_llm in {"llm_primary", "primary"}:
            return "primary"
        if normalized_llm in {"llm_fallback", "fallback"}:
            return "fallback"
        mapped_generation = OncoService._map_generation_path_for_pack(report_generation_path)
        if mapped_generation in {"primary", "fallback"}:
            return mapped_generation
        return "deterministic"

    @staticmethod
    def _build_case_text_for_casefacts(normalized_payload: dict[str, Any], case_json: dict[str, Any] | None) -> str:
        case_payload = normalized_payload.get("case") if isinstance(normalized_payload.get("case"), dict) else {}
        treatment_plan = (
            normalized_payload.get("treatment_plan")
            if isinstance(normalized_payload.get("treatment_plan"), dict)
            else {}
        )
        case_json_notes = str(case_json.get("notes") or "").strip() if isinstance(case_json, dict) else ""
        if case_json_notes:
            # Preserve stable evidence span offsets relative to stored case notes.
            return OncoService._segment_case_narrative(case_json_notes)

        case_notes = str(case_payload.get("notes") or "").strip()
        if case_notes:
            return OncoService._segment_case_narrative(case_notes)

        plan_text = str(treatment_plan.get("plan_text") or "").strip()
        return OncoService._segment_case_narrative(plan_text)

    @staticmethod
    def _extract_case_page_map(case_json: dict[str, Any] | None) -> dict[int, tuple[int, int]]:
        if not isinstance(case_json, dict):
            return {}
        raw_page_map = case_json.get("page_map")
        if not isinstance(raw_page_map, dict):
            kin_pdf = case_json.get("kin_pdf")
            if isinstance(kin_pdf, dict) and isinstance(kin_pdf.get("page_map"), dict):
                raw_page_map = kin_pdf.get("page_map")
        if not isinstance(raw_page_map, dict):
            return {}
        normalized: dict[int, tuple[int, int]] = {}
        for key, value in raw_page_map.items():
            if not isinstance(value, (list, tuple)) or len(value) != 2:
                continue
            try:
                page = int(str(key))
                start = int(value[0])
                end = int(value[1])
            except (TypeError, ValueError):
                continue
            if page <= 0 or start < 0 or end < start:
                continue
            normalized[page] = (start, end)
        return normalized

    def _build_drug_safety(
        self,
        *,
        case_text: str,
        case_json: dict[str, Any] | None,
    ) -> DoctorDrugSafety:
        if not bool(getattr(self.settings, "oncoai_drug_safety_enabled", True)):
            return DoctorDrugSafety(
                status="unavailable",
                warnings=[DrugSafetyWarning(code="DRUG_SAFETY_DISABLED", message="Drug safety enrichment disabled by feature flag.")],
            )
        entries = self._drug_dictionary_entries if isinstance(self._drug_dictionary_entries, list) else []
        regimens = self._drug_regimen_aliases if isinstance(self._drug_regimen_aliases, list) else []
        if not entries:
            return DoctorDrugSafety(
                status="unavailable",
                warnings=[DrugSafetyWarning(code="DRUG_DICTIONARY_EMPTY", message="Локальный словарь препаратов не загружен.")],
            )

        extracted, unresolved = extract_drugs_and_regimens(
            case_text=case_text,
            entries=entries,
            regimens=regimens,
            synonyms_extra=self._drug_synonyms_extra,
            page_map=self._extract_case_page_map(case_json),
        )
        inns = [item.inn for item in extracted if str(item.inn).strip()]
        fetch_result = self.drug_safety_provider.get_profiles(inns) if inns else None

        status = "unavailable"
        profiles: list[Any] = []
        warnings: list[DrugSafetyWarning] = []
        if fetch_result is not None:
            status = fetch_result.status
            profiles = fetch_result.profiles
            warnings.extend(fetch_result.warnings)
        elif extracted:
            status = "partial"
            warnings.append(
                DrugSafetyWarning(
                    code="DRUG_SAFETY_NOT_FETCHED",
                    message="Safety-данные не запрошены: не удалось инициализировать провайдер.",
                )
            )

        signals = build_drug_safety_signals(
            extracted=extracted,
            profiles=profiles,
            unresolved=unresolved,
            case_text=case_text,
        )

        if unresolved and status == "ok":
            status = "partial"
        if extracted and status == "unavailable":
            status = "partial"
        return DoctorDrugSafety(
            status=status,
            extracted_inn=extracted,
            unresolved_candidates=unresolved,
            profiles=profiles,
            signals=signals,
            warnings=warnings,
        )

    @staticmethod
    def _drug_signal_to_issue(signal: DrugSafetySignal) -> dict[str, Any]:
        return {
            "severity": signal.severity,
            "kind": signal.kind,
            "summary": signal.summary,
            "details": signal.details,
            "field_path": "drug_safety",
        }

    @staticmethod
    def _segment_case_narrative(raw_text: str) -> str:
        text = str(raw_text or "").strip()
        if not text:
            return ""
        lowered = text.lower()
        literal_markers = [
            "для врача (краткое обоснование",
            "для пациента",
            "что нужно сделать:",
            "почему это неправильно:",
        ]
        cut_points = [lowered.find(marker) for marker in literal_markers if lowered.find(marker) >= 0]
        regex_markers = [
            r"рекомендац\w*\s*ai[\s\-‑–—]*помощник\w*",
            r"рекомендац\w*\s*ai[\s\-‑–—]*ассистент\w*",
            r"recommendation\s*ai",
        ]
        for pattern in regex_markers:
            match = re.search(pattern, lowered, flags=re.IGNORECASE)
            if match:
                cut_points.append(match.start())
        if not cut_points:
            return text
        cutoff = min(cut_points)
        if cutoff < 24:
            return text
        return text[:cutoff].strip()

    @staticmethod
    def _tokenize_support_text(text: str) -> set[str]:
        tokens = {
            token.lower()
            for token in _SUPPORT_TOKEN_PATTERN.findall(str(text or ""))
            if len(token) >= 4 and token.lower() not in _SUPPORT_STOPWORDS
        }
        return tokens

    @staticmethod
    def _casefacts_support_tokens(case_facts: dict[str, Any]) -> set[str]:
        pieces: list[str] = []
        if not isinstance(case_facts, dict):
            return set()
        for field in ("initial_stage", "current_stage"):
            stage = case_facts.get(field)
            if isinstance(stage, dict):
                pieces.append(str(stage.get("tnm") or ""))
                pieces.append(str(stage.get("stage_group") or ""))
                spans = stage.get("evidence_spans")
                if isinstance(spans, list):
                    pieces.extend(str(span.get("text") or "") for span in spans if isinstance(span, dict))
        biomarkers = case_facts.get("biomarkers")
        if isinstance(biomarkers, dict):
            pieces.append(str(biomarkers.get("her2") or ""))
            pieces.append(str(biomarkers.get("msi_status") or ""))
            pdl1_values = biomarkers.get("pd_l1_cps_values")
            if isinstance(pdl1_values, list):
                pieces.extend(str(item) for item in pdl1_values)
            spans = biomarkers.get("evidence_spans")
            if isinstance(spans, list):
                pieces.extend(str(span.get("text") or "") for span in spans if isinstance(span, dict))
        metastases = case_facts.get("metastases")
        if isinstance(metastases, list):
            for item in metastases:
                if not isinstance(item, dict):
                    continue
                pieces.append(str(item.get("site") or ""))
                spans = item.get("evidence_spans")
                if isinstance(spans, list):
                    pieces.extend(str(span.get("text") or "") for span in spans if isinstance(span, dict))
        history = case_facts.get("treatment_history")
        if isinstance(history, list):
            for item in history:
                if not isinstance(item, dict):
                    continue
                pieces.append(str(item.get("name") or ""))
                pieces.append(str(item.get("response") or ""))
                spans = item.get("evidence_spans")
                if isinstance(spans, list):
                    pieces.extend(str(span.get("text") or "") for span in spans if isinstance(span, dict))
        complications = case_facts.get("complications")
        if isinstance(complications, list):
            pieces.extend(str(item) for item in complications)
        return OncoService._tokenize_support_text("\n".join(item for item in pieces if item))

    @staticmethod
    def _issue_has_casefacts_support(issue: dict[str, Any], case_facts: dict[str, Any]) -> bool:
        issue_text = f"{str(issue.get('summary') or '')}\n{str(issue.get('details') or '')}"
        issue_tokens = OncoService._tokenize_support_text(issue_text)
        if not issue_tokens:
            return False
        support_tokens = OncoService._casefacts_support_tokens(case_facts)
        if not support_tokens:
            return False
        overlap = issue_tokens.intersection(support_tokens)
        return len(overlap) >= 1

    @staticmethod
    def _guard_pack_issues_against_support(
        *,
        issues: list[dict[str, Any]],
        case_facts: dict[str, Any],
        citations: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        real_citation_ids = {
            str(item.get("citation_id") or "")
            for item in citations
            if isinstance(item, dict) and str(item.get("chunk_id") or "") not in {"", "synthetic:no_evidence"}
        }
        guarded: list[dict[str, Any]] = []
        for raw_issue in issues:
            if not isinstance(raw_issue, dict):
                continue
            issue = dict(raw_issue)
            if str(issue.get("kind") or "") != "contraindication":
                guarded.append(issue)
                continue
            citation_ids = issue.get("citation_ids") if isinstance(issue.get("citation_ids"), list) else []
            has_retrieved_support = bool(real_citation_ids.intersection({str(item) for item in citation_ids if str(item)}))
            has_casefacts_support = OncoService._issue_has_casefacts_support(issue, case_facts)
            if has_retrieved_support or has_casefacts_support:
                guarded.append(issue)
                continue
            downgraded = dict(issue)
            downgraded["kind"] = "inconsistency"
            downgraded["severity"] = "warning"
            details = str(downgraded.get("details") or "").strip()
            note = "Флаг contraindication понижен: не найдено подтверждение в case_facts или retrieved evidence."
            downgraded["details"] = f"{details} {note}".strip()
            guarded.append(downgraded)
        return guarded

    @staticmethod
    def _normalize_pack_issue_texts(issues: list[dict[str, Any]]) -> list[dict[str, Any]]:
        normalized: list[dict[str, Any]] = []
        for raw in issues:
            if not isinstance(raw, dict):
                continue
            issue = dict(raw)
            issue["summary"] = normalize_ru_clinical_text(str(issue.get("summary") or ""))
            issue["details"] = normalize_ru_clinical_text(str(issue.get("details") or ""))
            normalized.append(issue)
        return normalized

    @staticmethod
    def _validate_comparative_claims_policy(doctor_report: dict[str, Any]) -> None:
        claims = doctor_report.get("comparative_claims")
        if claims is None:
            return
        if not isinstance(claims, list):
            raise ValidationError("doctor_report.comparative_claims must be array")
        for idx, claim in enumerate(claims, start=1):
            if not isinstance(claim, dict):
                raise ValidationError(f"doctor_report.comparative_claims[{idx}] must be object")
            if not bool(claim.get("comparative_superiority")):
                continue
            pubmed_id = str(claim.get("pubmed_id") or "").strip()
            pubmed_url = str(claim.get("pubmed_url") or "").strip()
            if not pubmed_id and not pubmed_url:
                raise ValidationError(
                    f"comparative claim #{idx} must include pubmed_id or pubmed_url when comparative_superiority=true"
                )
            if pubmed_url and not is_pubmed_url(pubmed_url):
                raise ValidationError(
                    f"comparative claim #{idx} pubmed_url must point to PubMed/NCBI domain"
                )

    @staticmethod
    def _build_sources_only_result(response: dict[str, Any]) -> dict[str, Any]:
        doctor_report = response.get("doctor_report") if isinstance(response.get("doctor_report"), dict) else {}
        citations = doctor_report.get("citations") if isinstance(doctor_report.get("citations"), list) else []
        items: list[dict[str, Any]] = []
        for index, citation in enumerate(citations[:20], start=1):
            if not isinstance(citation, dict):
                continue
            citation_id = str(citation.get("citation_id") or "").strip()
            source_id = str(citation.get("source_id") or "unknown").strip().lower() or "unknown"
            section_path = str(citation.get("section_path") or "guideline_fragment").strip() or "guideline_fragment"
            quote = str(citation.get("quote") or "").strip() or "Релевантный фрагмент клинической рекомендации."
            items.append(
                {
                    "item_id": f"{source_id}_{index}",
                    "title": f"{source_id.upper()}: {section_path}",
                    "summary": quote,
                    "citation_ids": [citation_id] if citation_id else [],
                    "source_ids": [source_id],
                }
            )
        if not items:
            items.append(
                {
                    "item_id": "sources_only_no_evidence",
                    "title": "Источники не найдены",
                    "summary": "Не удалось извлечь релевантные фрагменты рекомендаций из заданных источников.",
                    "citation_ids": [],
                    "source_ids": [],
                }
            )
        return {
            "mode": "SOURCES_ONLY",
            "items": items,
            "policy": "Терапевтические предложения отключены: только источники и цитируемые фрагменты.",
        }

    def _build_historical_assessment(
        self,
        *,
        as_of_date: str | None,
        historical_assessment_requested: bool,
        source_ids: list[str],
        citations: list[dict[str, Any]],
        routing_meta: dict[str, Any] | None = None,
    ) -> dict[str, Any] | None:
        if not historical_assessment_requested and not as_of_date:
            return None

        today_date = date.today()
        today = today_date.isoformat()
        requested_token = str(as_of_date or "").strip()
        requested_date = self._parse_iso_date_token(requested_token) if requested_token else None
        requested = requested_date.isoformat() if requested_date else (requested_token or today)
        status = "ok"
        reason_code = "ok"
        source_ids_clean = normalize_source_set_ids([str(item).strip() for item in source_ids if str(item).strip()])
        if not source_ids_clean:
            source_ids_clean = sorted(
                normalize_source_set_ids(
                    [
                        str(item.get("source_id") or "").strip()
                        for item in citations
                        if isinstance(item, dict) and str(item.get("source_id") or "").strip()
                    ]
                )
            )

        historical_meta = routing_meta.get("historical") if isinstance(routing_meta, dict) else {}
        historical_meta = historical_meta if isinstance(historical_meta, dict) else {}
        current_snapshot = historical_meta.get("current_snapshot")
        current_snapshot = current_snapshot if isinstance(current_snapshot, dict) else {}
        as_of_snapshot = historical_meta.get("as_of_snapshot")
        as_of_snapshot = as_of_snapshot if isinstance(as_of_snapshot, dict) else {}

        current_source_ids = normalize_source_set_ids(
            [
                str(item).strip()
                for item in (current_snapshot.get("source_ids") if isinstance(current_snapshot.get("source_ids"), list) else [])
                if str(item).strip()
            ]
        ) or list(source_ids_clean)
        as_of_source_ids = normalize_source_set_ids(
            [
                str(item).strip()
                for item in (as_of_snapshot.get("source_ids") if isinstance(as_of_snapshot.get("source_ids"), list) else [])
                if str(item).strip()
            ]
        ) or list(source_ids_clean)

        if historical_assessment_requested and not requested_token:
            status = "insufficient_data"
            reason_code = "missing_as_of_date"
        elif requested_date and requested_date > today_date:
            status = "insufficient_data"
            reason_code = "future_as_of_date"
        elif requested_token:
            selected_versions = (
                as_of_snapshot.get("selected_versions")
                if isinstance(as_of_snapshot.get("selected_versions"), list)
                else []
            )
            missing_pairs = (
                as_of_snapshot.get("missing_pairs")
                if isinstance(as_of_snapshot.get("missing_pairs"), list)
                else []
            )
            if selected_versions and missing_pairs:
                status = "insufficient_data"
                reason_code = "historical_sources_unavailable"
            elif not selected_versions:
                if historical_meta or not citations:
                    status = "insufficient_data"
                    reason_code = "historical_sources_unavailable"

        conflicts: list[dict[str, Any]] = []
        current_versions = (
            current_snapshot.get("selected_versions")
            if isinstance(current_snapshot.get("selected_versions"), list)
            else []
        )
        as_of_versions = (
            as_of_snapshot.get("selected_versions")
            if isinstance(as_of_snapshot.get("selected_versions"), list)
            else []
        )
        if current_versions and as_of_versions:
            current_map: dict[tuple[str, str], str] = {}
            as_of_map: dict[tuple[str, str], str] = {}
            for item in current_versions:
                if not isinstance(item, dict):
                    continue
                key = (
                    normalize_source_set_id(str(item.get("source_id") or "")),
                    str(item.get("doc_id") or "").strip(),
                )
                if not key[0] or not key[1]:
                    continue
                current_map[key] = str(item.get("doc_version") or "").strip()
            for item in as_of_versions:
                if not isinstance(item, dict):
                    continue
                key = (
                    normalize_source_set_id(str(item.get("source_id") or "")),
                    str(item.get("doc_id") or "").strip(),
                )
                if not key[0] or not key[1]:
                    continue
                as_of_map[key] = str(item.get("doc_version") or "").strip()
            for pair in sorted(set(current_map.keys()).intersection(set(as_of_map.keys()))):
                current_version = current_map.get(pair) or ""
                as_of_version = as_of_map.get(pair) or ""
                if current_version and as_of_version and current_version != as_of_version:
                    conflicts.append(
                        {
                            "type": "version_changed",
                            "source_id": pair[0],
                            "doc_id": pair[1],
                            "current_doc_version": current_version,
                            "as_of_doc_version": as_of_version,
                        }
                    )
            for pair in sorted(set(current_map.keys()).difference(set(as_of_map.keys()))):
                conflicts.append(
                    {
                        "type": "missing_historical_version",
                        "source_id": pair[0],
                        "doc_id": pair[1],
                        "current_doc_version": current_map.get(pair) or "",
                    }
                )

        return {
            "requested_as_of_date": requested,
            "status": status,
            "reason_code": reason_code,
            "current_guideline": {
                "as_of_date": today,
                "source_ids": current_source_ids,
                "note": "Оценка по текущему состоянию базы рекомендаций.",
            },
            "as_of_date_guideline": {
                "as_of_date": requested,
                "source_ids": as_of_source_ids,
                "note": (
                    "Historical snapshot рассчитан по доступным версиям документов."
                    if status == "ok"
                    else "Недостаточно данных для полной historical реконструкции."
                ),
            },
            "conflicts": conflicts,
        }

    @staticmethod
    def _apply_sources_only_mode(response: dict[str, Any]) -> dict[str, Any]:
        patched = copy.deepcopy(response)
        doctor_report = patched.get("doctor_report") if isinstance(patched.get("doctor_report"), dict) else {}
        doctor_report["plan"] = []
        doctor_report["issues"] = []
        doctor_report["consilium_md"] = (
            "## SOURCES_ONLY\n"
            "- Режим ответа: только цитируемые фрагменты рекомендаций.\n"
            "- Терапевтическая тактика не формируется.\n"
        )
        doctor_report["summary_md"] = "Режим sources-only: терапевтические предложения отключены."
        patched["doctor_report"] = doctor_report

        patient_explain = patched.get("patient_explain")
        if isinstance(patient_explain, dict):
            patient_explain["summary_plain"] = "Показаны только источники рекомендаций без назначения лечения."
            patient_explain["key_points"] = [
                "В этом режиме не формируются лечебные назначения.",
                "Проверьте источники и обсудите тактику с лечащим врачом.",
            ]
            patient_explain["questions_for_doctor"] = [
                "Какие пункты рекомендаций применимы к моему случаю?",
                "Какие данные нужно добавить для выбора тактики лечения?",
            ]
            patched["patient_explain"] = patient_explain

        patched["sources_only_result"] = OncoService._build_sources_only_result(patched)
        return patched

    @staticmethod
    def _enrich_plan_steps_with_evidence_metadata(plan_sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
        enriched: list[dict[str, Any]] = []
        for section in plan_sections:
            if not isinstance(section, dict):
                continue
            section_copy = dict(section)
            steps = section_copy.get("steps") if isinstance(section_copy.get("steps"), list) else []
            new_steps: list[dict[str, Any]] = []
            for step in steps:
                if not isinstance(step, dict):
                    continue
                step_copy = dict(step)
                citation_ids = step_copy.get("citation_ids") if isinstance(step_copy.get("citation_ids"), list) else []
                if not str(step_copy.get("evidence_level") or "").strip():
                    step_copy["evidence_level"] = "LoE-NotSpecified"
                if not str(step_copy.get("recommendation_strength") or "").strip():
                    step_copy["recommendation_strength"] = "GoR-Consensus"
                if step_copy.get("confidence") is None:
                    step_copy["confidence"] = 0.82 if citation_ids else 0.45
                new_steps.append(step_copy)
            section_copy["steps"] = new_steps
            enriched.append(section_copy)
        return enriched

    @staticmethod
    def _doc_release_validity(
        *,
        doc_id: str,
        source_set: str,
        source_url: str,
        status: str,
        cancer_type: str,
    ) -> dict[str, Any]:
        normalized_source = normalize_source_set_id(source_set)
        validity = evaluate_release_validity(
            source_set=normalized_source,
            source_url=source_url,
            status=status,
            doc_id=doc_id,
            nosology_mapped=is_nosology_mapped(cancer_type),
        )
        return {
            "is_valid": bool(validity.get("is_valid")),
            "validity_reason": str(validity.get("validity_reason") or "unknown"),
            "official_source": str(validity.get("official_source") or ""),
        }

    @staticmethod
    def _build_sync_validation_summary(records: list[dict[str, Any]]) -> dict[str, Any]:
        valid = 0
        invalid = 0
        manual = 0
        for item in records:
            reason = str(item.get("validity_reason") or "").strip().lower()
            if bool(item.get("is_valid")):
                valid += 1
            elif reason in {"nosology_unmapped", "missing_source_url", "status_not_release_ready"}:
                manual += 1
            else:
                invalid += 1
        return {
            "total": len(records),
            "valid": valid,
            "invalid": invalid,
            "needs_manual": manual,
        }

    def _delete_doc_everywhere(self, *, doc_id: str, doc_version: str) -> dict[str, Any]:
        doc = self.store.get_doc(doc_id, doc_version)
        self.index.replace_doc_chunks(doc_id=doc_id, doc_version=doc_version, chunks=[])
        deleted = self.store.delete_doc_bundle(doc_id=doc_id, doc_version=doc_version)
        file_removed = False
        if isinstance(doc, dict):
            file_path = Path(str(doc.get("file_path") or ""))
            try:
                if file_path.exists():
                    file_path.unlink()
                    file_removed = True
                    parent = file_path.parent
                    for _ in range(3):
                        if parent == self.settings.docs_dir or parent == parent.parent:
                            break
                        if any(parent.iterdir()):
                            break
                        parent.rmdir()
                        parent = parent.parent
            except Exception:  # noqa: BLE001
                file_removed = False
        deleted["file_removed"] = file_removed
        self._kb_version = compute_kb_version(self.store.list_docs())
        self._clear_routing_cache()
        return deleted

    @staticmethod
    def _infer_cancer_type(explicit_cancer_type: str, case_text: str) -> str:
        normalized = str(explicit_cancer_type or "").strip().lower()

        text = str(case_text or "").lower()
        inferred_from_text = "unknown"
        if re.search(r"рак\s+желудк|аденокарцином\w*\s+желудк|gastric|siewert|кардиоэзофаг", text):
            inferred_from_text = "gastric_cancer"
        elif re.search(r"рак\s+легк|nsclc|аденокарцином\w*\s+легк", text):
            inferred_from_text = "nsclc_egfr"
        elif re.search(r"рак\s+молочн\w*\s+желез|breast", text):
            inferred_from_text = "breast_hr+/her2-"

        if normalized and normalized != "unknown":
            if inferred_from_text != "unknown" and inferred_from_text != normalized:
                return inferred_from_text
            return normalized
        return inferred_from_text

    def _build_native_pack_v1_2_response_llm_only(
        self,
        *,
        bridge_context: Any,
        normalized_payload: dict[str, Any],
        retrieved: list[dict[str, Any]],
        reranked: list[dict[str, Any]],
        llm_path: str,
        report_generation_path: str,
        fallback_reason: str | None,
        routing_meta: dict[str, Any],
        started: float,
        legacy_doctor_report: dict[str, Any],
    ) -> dict[str, Any]:
        request_id = str(bridge_context.request_id or normalized_payload.get("request_id") or uuid.uuid4())
        query_type = str(bridge_context.query_type or "NEXT_STEPS").strip().upper()
        case_payload = normalized_payload.get("case") if isinstance(normalized_payload.get("case"), dict) else {}
        case_text = self._build_case_text_for_casefacts(normalized_payload=normalized_payload, case_json=bridge_context.case_json)
        cancer_type = self._infer_cancer_type(
            explicit_cancer_type=str(case_payload.get("cancer_type") or "unknown").strip(),
            case_text=case_text,
        )

        if bool(self.settings.oncoai_casefacts_enabled):
            case_facts = extract_case_facts(case_text=case_text, case_json=bridge_context.case_json).model_dump()
            case_facts["case_facts_v2"] = extract_case_facts_v2(
                case_text=case_text,
                case_json=bridge_context.case_json,
                drug_dictionary_entries=self._drug_dictionary_entries,
                drug_regimen_aliases=self._drug_regimen_aliases,
                drug_synonyms_extra=self._drug_synonyms_extra,
            ).model_dump()
        else:
            case_facts = {
                "initial_stage": {},
                "current_stage": {},
                "metastases": [],
                "biomarkers": {},
                "treatment_history": [],
                "complications": [],
                "case_facts_v2": {"patient": {}, "labs": [], "current_medications": [], "comorbidities": [], "tumor": {}, "therapy_timeline": [], "key_unknowns": []},
                "key_unknowns": [],
            }

        disease_context = build_disease_context(
            normalized_payload=normalized_payload,
            case_json=bridge_context.case_json,
            case_facts=case_facts,
        )
        timeline = build_timeline(
            bridge_context.case_json,
            case_facts=case_facts,
        )
        timeline_reconciliation = reconcile_timeline_signals(
            case_text=case_text,
            case_facts=case_facts,
            timeline=timeline,
        )
        timeline = build_timeline(
            bridge_context.case_json,
            case_facts=case_facts,
            timeline_reconciliation=timeline_reconciliation,
        )

        citations, chunk_to_citation = build_citations_from_chunks(
            reranked_chunks=reranked,
            max_citations=40,
            version_metadata_resolver=self.store.get_guideline_version_by_doc,
        )
        citations, chunk_to_citation = self._sanitize_citations(
            citations=citations,
            chunk_to_citation=chunk_to_citation,
            reranked_chunks=reranked,
        )

        llm_issues_raw = legacy_doctor_report.get("issues") if isinstance(legacy_doctor_report.get("issues"), list) else []
        pack_issues: list[dict[str, Any]] = []
        for item in llm_issues_raw:
            if not isinstance(item, dict):
                continue
            pack_issues.append(
                {
                    "issue_id": self._pack_issue_id(str(item.get("issue_id") or uuid.uuid4())),
                    "severity": self._pack_issue_severity(str(item.get("severity") or "")),
                    "kind": self._pack_issue_kind(
                        category=str(item.get("category") or ""),
                        title=str(item.get("title") or item.get("summary") or ""),
                    ),
                    "summary": normalize_ru_clinical_text(str(item.get("title") or item.get("summary") or "Potential issue")),
                    "details": normalize_ru_clinical_text(str(item.get("description") or item.get("details") or "")),
                    "field_path": "llm_report",
                }
            )
        fallback_citation_ids = [str(item.get("citation_id") or "") for item in citations if str(item.get("citation_id") or "").strip()]
        pack_issues = attach_issue_citations(
            issues=pack_issues,
            reranked_chunks=reranked,
            chunk_to_citation=chunk_to_citation,
            fallback_citation_ids=fallback_citation_ids,
        )
        if not citations:
            pack_issues = [
                issue
                for issue in pack_issues
                if str(issue.get("kind") or "").strip().lower() == "missing_data"
            ]

        insufficient_status = not bool(citations)
        insufficient_reason = (
            "Недостаточно подтвержденных цитат из загруженных и одобренных источников."
            if insufficient_status
            else "Достаточно подтверждающих данных."
        )
        insufficient_reason = normalize_ru_clinical_text(insufficient_reason)

        doctor_report = {
            "schema_version": "1.2",
            "report_id": self._pack_issue_id(str(legacy_doctor_report.get("report_id") or uuid.uuid4())),
            "request_id": request_id,
            "query_type": query_type if query_type in {"NEXT_STEPS", "CHECK_LAST_TREATMENT"} else "NEXT_STEPS",
            "disease_context": disease_context,
            "case_facts": case_facts,
            "timeline": timeline,
            "consilium_md": build_consilium_md(
                query_type=query_type,
                case_facts=case_facts,
                plan_sections=[],
                issues=pack_issues,
                sufficiency={"status": insufficient_status, "reason": insufficient_reason},
                source_ids=[str(item.get("source_id") or "") for item in citations if str(item.get("source_id") or "").strip()],
                has_real_evidence=bool(citations),
                timeline_reconciliation=timeline_reconciliation,
            ),
            "plan": [],
            "issues": pack_issues,
            "comparative_claims": [],
            "sanity_checks": [],
            "drug_safety": {
                "status": "unavailable",
                "extracted_inn": [],
                "unresolved_candidates": [],
                "profiles": [],
                "signals": [],
                "warnings": [],
            },
            "citations": citations,
            "generated_at": datetime.now(timezone.utc).isoformat(),
        }
        doctor_report["summary"] = normalize_ru_clinical_text(str(legacy_doctor_report.get("summary") or ""))
        if not doctor_report["summary"]:
            doctor_report["summary"] = "Сформирован LLM+RAG отчёт по доступным подтверждённым источникам."
        doctor_report["summary_md"] = doctor_report["summary"]
        doctor_report["verification_summary"] = self._build_verification_summary(
            issues=pack_issues,
            insufficient_status=insufficient_status,
        )

        patient_context_report = dict(doctor_report)
        patient_context_report["kb_version"] = str(legacy_doctor_report.get("kb_version") or self._kb_version)
        patient_context_report["missing_data"] = [
            {"field": str(item.get("summary") or item.get("kind") or "missing_data"), "reason": str(item.get("details") or "")}
            for item in pack_issues
            if str(item.get("kind") or "").strip().lower() == "missing_data"
        ][:10]
        patient_context_report["insufficient_data"] = {"status": insufficient_status, "reason": insufficient_reason}

        strict_patient, patient_generation_path = build_patient_explain_with_fallback(
            doctor_report=patient_context_report,
            llm_router=self.llm_router,
            prompt_registry=self.prompt_registry,
            prompt_schema_strict=bool(self.settings.oncoai_prompt_schema_strict),
            strict_llm_only=True,
        )
        if str(patient_generation_path or "").strip().lower() != "llm":
            raise ValidationError("STRICT_LLM_RAG_ONLY: patient explain generation must use LLM path")

        patient_explain = map_strict_to_pack_patient_v1_2(
            strict_payload=strict_patient,
            request_id=request_id,
            source_ids=[str(item.get("source_id") or "") for item in citations if str(item.get("source_id") or "").strip()],
            drug_safety={"status": "unavailable", "important_risks": [], "questions_for_doctor": []},
        )

        effective_llm_path = self._resolve_run_meta_llm_path(llm_path, report_generation_path)
        run_meta = build_run_meta(
            request_id=request_id,
            retrieved=retrieved,
            reranked=reranked,
            citations_count=len(citations),
            llm_path=effective_llm_path,
            report_generation_path=self._map_generation_path_for_pack(report_generation_path),
            fallback_reason=self._map_fallback_reason_for_pack(fallback_reason),
            reasoning_mode=self._reasoning_mode,
            routing_meta=routing_meta,
            started_at_perf=started,
            settings=self.settings,
        )
        if str(run_meta.get("llm_path") or "").strip() != "primary":
            raise ValidationError("STRICT_LLM_RAG_ONLY: run_meta.llm_path must be primary")
        if str(run_meta.get("report_generation_path") or "").strip() != "primary":
            raise ValidationError("STRICT_LLM_RAG_ONLY: run_meta.report_generation_path must be primary")
        run_meta["fallback_reason"] = "none"
        run_meta["patient_generation_path"] = "llm"

        return {
            "schema_version": "0.2",
            "request_id": request_id,
            "doctor_report": doctor_report,
            "patient_explain": patient_explain,
            "run_meta": run_meta,
            "meta": self._build_execution_meta(),
            "insufficient_data": {"status": insufficient_status, "reason": insufficient_reason},
        }

    def _build_native_pack_v1_2_response(
        self,
        *,
        bridge_context: Any,
        normalized_payload: dict[str, Any],
        retrieved: list[dict[str, Any]],
        reranked: list[dict[str, Any]],
        llm_path: str,
        report_generation_path: str,
        fallback_reason: str | None,
        routing_meta: dict[str, Any],
        started: float,
        legacy_doctor_report: dict[str, Any],
        generated_plan_sections: list[dict[str, Any]],
    ) -> dict[str, Any]:
        if self._llm_rag_only_mode:
            return self._build_native_pack_v1_2_response_llm_only(
                bridge_context=bridge_context,
                normalized_payload=normalized_payload,
                retrieved=retrieved,
                reranked=reranked,
                llm_path=llm_path,
                report_generation_path=report_generation_path,
                fallback_reason=fallback_reason,
                routing_meta=routing_meta,
                started=started,
                legacy_doctor_report=legacy_doctor_report,
            )

        request_id = str(bridge_context.request_id or normalized_payload.get("request_id") or uuid.uuid4())
        query_type = str(bridge_context.query_type or "NEXT_STEPS").strip().upper()
        case_payload = normalized_payload.get("case") if isinstance(normalized_payload.get("case"), dict) else {}
        case_text = self._build_case_text_for_casefacts(normalized_payload=normalized_payload, case_json=bridge_context.case_json)
        cancer_type = self._infer_cancer_type(
            explicit_cancer_type=str(case_payload.get("cancer_type") or "unknown").strip(),
            case_text=case_text,
        )

        if bool(self.settings.oncoai_casefacts_enabled):
            case_facts = extract_case_facts(case_text=case_text, case_json=bridge_context.case_json).model_dump()
            case_facts["case_facts_v2"] = extract_case_facts_v2(
                case_text=case_text,
                case_json=bridge_context.case_json,
                drug_dictionary_entries=self._drug_dictionary_entries,
                drug_regimen_aliases=self._drug_regimen_aliases,
                drug_synonyms_extra=self._drug_synonyms_extra,
            ).model_dump()
        else:
            case_facts = {
                "initial_stage": {},
                "current_stage": {},
                "metastases": [],
                "biomarkers": {},
                "treatment_history": [],
                "complications": [],
                "case_facts_v2": {
                    "patient": {},
                    "labs": [],
                    "current_medications": [],
                    "comorbidities": [],
                    "tumor": {},
                    "therapy_timeline": [],
                    "key_unknowns": [
                        "CaseFactsV2 extraction disabled by feature flag ONCOAI_CASEFACTS_ENABLED=false",
                    ],
                },
                "key_unknowns": [
                    "CaseFacts extraction disabled by feature flag ONCOAI_CASEFACTS_ENABLED=false",
                ],
            }

        drug_safety = self._build_drug_safety(
            case_text=case_text,
            case_json=bridge_context.case_json,
        )

        disease_context = build_disease_context(
            normalized_payload=normalized_payload,
            case_json=bridge_context.case_json,
            case_facts=case_facts,
        )
        timeline = build_timeline(
            bridge_context.case_json,
            case_facts=case_facts,
        )
        timeline_reconciliation = reconcile_timeline_signals(
            case_text=case_text,
            case_facts=case_facts,
            timeline=timeline,
        )
        timeline = build_timeline(
            bridge_context.case_json,
            case_facts=case_facts,
            timeline_reconciliation=timeline_reconciliation,
        )
        min_requirements = evaluate_min_case_requirements(
            case_json=bridge_context.case_json if isinstance(bridge_context.case_json, dict) else {},
            case_facts=case_facts,
            disease_context=disease_context,
            case_payload=case_payload,
            routing_meta=routing_meta,
        )
        case_facts["minimum_dataset"] = min_requirements

        plan_sections = generated_plan_sections
        if bool(min_requirements.get("status")) and bool(min_requirements.get("no_ready_plan_without_minimum", True)):
            safe_intents = [
                str(item).strip()
                for item in min_requirements.get("safe_missing_data_plan_intents", [])
                if str(item).strip()
            ] or ["дозапрос данных", "уточнение", "безопасность"]
            safe_steps: list[dict[str, Any]] = []
            for index, intent in enumerate(safe_intents, start=1):
                safe_steps.append(
                    {
                        "step_id": str(uuid.uuid5(uuid.NAMESPACE_URL, f"oncoai:safe-plan:{request_id}:{index}:{intent}")),
                        "text": f"Выполнить: {intent}.",
                        "priority": "high",
                        "rationale": str(min_requirements.get("reason") or "Сначала нужно закрыть критические пробелы данных."),
                        "citation_ids": [],
                    }
                )
            safe_steps.append(
                {
                    "step_id": str(uuid.uuid5(uuid.NAMESPACE_URL, f"oncoai:safe-plan:{request_id}:reason")),
                    "text": "До восполнения минимального набора данных готовый лечебный план не формируется.",
                    "priority": "high",
                    "rationale": str(min_requirements.get("reason") or "Сначала нужно закрыть критические пробелы данных."),
                    "citation_ids": [],
                }
            )
            plan_sections = [
                {
                    "section": "diagnostics",
                    "title": "Уточнение минимально необходимого набора данных",
                    "steps": safe_steps,
                }
            ]
        plan_structured = flatten_plan_for_diff(plan_sections)

        diff_issues = compute_diff(cancer_type, plan_structured, query_type=query_type)
        deterministic_issues: list[dict[str, Any]] = [
            {
                "severity": self._pack_issue_severity(issue.severity),
                "kind": self._pack_issue_kind(category=issue.category, title=issue.title),
                "summary": issue.title,
                "details": issue.description,
                "field_path": "plan",
            }
            for issue in diff_issues
        ]
        for signal in drug_safety.signals:
            deterministic_issues.append(self._drug_signal_to_issue(signal))

        if cancer_type == "gastric_cancer":
            deterministic_issues.extend(
                apply_gastric_rules(
                    case_facts=case_facts,
                    disease_context=disease_context,
                    case_text=case_text,
                    plan_sections=plan_sections,
                )
            )

        sufficiency = evaluate_data_sufficiency(
            case_facts=case_facts,
            query_type=query_type,
            case_text=case_text,
            plan_sections=plan_sections,
        )
        if bool(min_requirements.get("status")):
            sufficiency["status"] = True
            existing_missing = (
                sufficiency.get("missing_critical_fields")
                if isinstance(sufficiency.get("missing_critical_fields"), list)
                else []
            )
            for field in min_requirements.get("missing_critical_fields", []):
                token = str(field).strip()
                if token and token not in existing_missing:
                    existing_missing.append(token)
            sufficiency["missing_critical_fields"] = existing_missing
            sufficiency["reason"] = str(min_requirements.get("reason") or sufficiency.get("reason") or "")
        if bool(routing_meta.get("ambiguous_brain_scope")):
            sufficiency["status"] = True
            sufficiency["reason"] = (
                "ambiguous_brain_scope: укажите ICD-10 и контур нозологии "
                "(C71 первичная опухоль мозга или C79.3 метастазы в ЦНС)."
            )
            critical_fields = (
                sufficiency.get("missing_critical_fields")
                if isinstance(sufficiency.get("missing_critical_fields"), list)
                else []
            )
            if "brain_scope_icd10" not in critical_fields:
                critical_fields.append("brain_scope_icd10")
            sufficiency["missing_critical_fields"] = critical_fields
        if sufficiency.get("status"):
            deterministic_issues.append(
                {
                    "severity": "warning",
                    "kind": "missing_data",
                    "summary": "Недостаточно данных для уверенного выбора следующего шага.",
                    "details": str(sufficiency.get("reason") or "Нужно уточнить недостающие клинические данные."),
                    "field_path": "case_facts",
                }
            )
        timeline_missing = (
            timeline_reconciliation.get("missing_items")
            if isinstance(timeline_reconciliation.get("missing_items"), list)
            else []
        )
        if bool(timeline_reconciliation.get("n5_profile")) and timeline_missing:
            deterministic_issues.append(
                {
                    "severity": "warning",
                    "kind": "missing_data",
                    "summary": "Неполная клиническая последовательность для кейса после прогрессирования.",
                    "details": "Не найдено в истории кейса: " + ", ".join(str(item) for item in timeline_missing),
                    "field_path": "timeline",
                }
            )

        llm_issues: list[dict[str, Any]] = []
        legacy_issues = legacy_doctor_report.get("issues") if isinstance(legacy_doctor_report.get("issues"), list) else []
        for item in legacy_issues:
            if not isinstance(item, dict):
                continue
            llm_issues.append(
                {
                    "severity": self._pack_issue_severity(str(item.get("severity") or "")),
                    "kind": self._pack_issue_kind(
                        category=str(item.get("category") or ""),
                        title=str(item.get("title") or item.get("summary") or ""),
                    ),
                    "summary": str(item.get("title") or item.get("summary") or "Potential issue").strip(),
                    "details": str(item.get("description") or item.get("details") or "").strip(),
                    "field_path": "legacy_report",
                }
            )

        merged_issues: list[dict[str, Any]] = []
        seen_issue_keys: set[tuple[str, str]] = set()
        for issue in deterministic_issues:
            key = (str(issue.get("kind") or ""), str(issue.get("summary") or ""))
            if key in seen_issue_keys:
                continue
            seen_issue_keys.add(key)
            merged_issues.append(issue)
        for issue in llm_issues:
            key = (str(issue.get("kind") or ""), str(issue.get("summary") or ""))
            if key in seen_issue_keys:
                continue
            seen_issue_keys.add(key)
            merged_issues.append(issue)

        max_citations = 40
        options = bridge_context.normalized_payload.get("options") if isinstance(bridge_context.normalized_payload, dict) else {}
        if isinstance(options, dict) and isinstance(options.get("max_citations"), int):
            max_citations = max(1, min(int(options.get("max_citations") or 40), 100))
        def _version_metadata_for_citation(doc_id: str, doc_version: str) -> dict[str, Any]:
            version = self.store.get_guideline_version_by_doc(doc_id=doc_id, doc_version=doc_version)
            metadata = version.get("metadata") if isinstance(version, dict) else {}
            return metadata if isinstance(metadata, dict) else {}

        citations, chunk_to_citation = build_citations_from_chunks(
            reranked_chunks=reranked,
            max_citations=max_citations,
            version_metadata_resolver=_version_metadata_for_citation,
        )
        citations, chunk_to_citation = self._sanitize_citations(
            citations=citations,
            chunk_to_citation=chunk_to_citation,
            reranked_chunks=reranked,
        )
        fallback_citation_ids: list[str] = []
        plan_sections = attach_plan_citations(
            plan_sections=plan_sections,
            reranked_chunks=reranked,
            chunk_to_citation=chunk_to_citation,
            fallback_citation_ids=fallback_citation_ids,
        )
        plan_sections = self._filter_plan_sections_with_citations(plan_sections)
        plan_sections = self._enrich_plan_steps_with_evidence_metadata(plan_sections)
        merged_issues = attach_issue_citations(
            issues=merged_issues,
            reranked_chunks=reranked,
            chunk_to_citation=chunk_to_citation,
            fallback_citation_ids=fallback_citation_ids,
        )
        merged_issues = self._guard_pack_issues_against_support(
            issues=merged_issues,
            case_facts=case_facts,
            citations=citations,
        )

        pack_issues: list[dict[str, Any]] = []
        valid_citation_ids = {
            str(item.get("citation_id") or "")
            for item in citations
            if isinstance(item, dict) and str(item.get("citation_id") or "").strip()
        }
        for idx, issue in enumerate(merged_issues, start=1):
            kind = str(issue.get("kind") or "other")
            linked_citation_ids = [
                str(item)
                for item in (issue.get("citation_ids") if isinstance(issue.get("citation_ids"), list) else [])
                if str(item).strip() in valid_citation_ids
            ]
            if kind != "missing_data" and not linked_citation_ids:
                continue
            pack_issues.append(
                {
                    "issue_id": self._pack_issue_id(f"{request_id}:{idx}:{issue.get('kind')}:{issue.get('summary')}"),
                    "severity": self._pack_issue_severity(str(issue.get("severity") or "")),
                    "kind": kind,
                    "summary": str(issue.get("summary") or "Potential issue"),
                    "details": str(issue.get("details") or ""),
                    "field_path": str(issue.get("field_path") or ""),
                    "suggested_questions": [],
                    "citation_ids": linked_citation_ids,
                }
            )
        pack_issues = self._normalize_pack_issue_texts(pack_issues)

        consilium_md = build_consilium_md(
            query_type=query_type,
            case_facts=case_facts,
            plan_sections=plan_sections,
            issues=pack_issues,
            sufficiency=sufficiency,
            source_ids=(
                bridge_context.source_ids
                if isinstance(getattr(bridge_context, "source_ids", None), list)
                else [str(item.get("source_id") or "") for item in citations]
            ),
            has_real_evidence=bool(reranked),
            timeline_reconciliation=timeline_reconciliation,
        )
        generated_at = datetime.now(timezone.utc).isoformat()
        doctor_report = {
            "schema_version": "1.2",
            "report_id": self._pack_issue_id(str(legacy_doctor_report.get("report_id") or uuid.uuid4())),
            "request_id": request_id,
            "query_type": query_type if query_type in {"NEXT_STEPS", "CHECK_LAST_TREATMENT"} else "NEXT_STEPS",
            "disease_context": disease_context,
            "case_facts": case_facts,
            "timeline": timeline,
            "consilium_md": consilium_md,
            "plan": plan_sections,
            "issues": pack_issues,
            "comparative_claims": [],
            "sanity_checks": [],
            "drug_safety": self._attach_drug_signal_citations(
                drug_safety_payload=drug_safety.model_dump(),
                citation_ids=[str(item.get("citation_id") or "") for item in citations if str(item.get("citation_id") or "").strip()],
                strict_fail_closed=self._strict_fail_closed,
            ),
            "citations": citations,
            "generated_at": generated_at,
        }
        if bool(self.settings.oncoai_casefacts_enabled):
            sanity_checks = run_sanity_checks(case_facts=case_facts, doctor_report=doctor_report)
            if any(item.get("status") == "fail" for item in sanity_checks):
                doctor_report = auto_repair_report(case_facts=case_facts, doctor_report=doctor_report)
                sanity_checks = run_sanity_checks(case_facts=case_facts, doctor_report=doctor_report)
        else:
            sanity_checks = [
                {
                    "check_id": "casefacts_feature_flag_disabled",
                    "status": "warn",
                    "details": "CaseFacts extraction disabled by feature flag ONCOAI_CASEFACTS_ENABLED=false",
                }
            ]
        doctor_report["sanity_checks"] = sanity_checks

        insufficient_status = bool(sufficiency.get("status"))
        insufficient_reason = str(sufficiency.get("reason") or "Критических пробелов данных не выявлено.")
        if not citations:
            insufficient_status = True
            insufficient_reason = "Недостаточно подтвержденных цитат из загруженных и одобренных источников."
            plan_sections = []
            pack_issues = [
                issue
                for issue in pack_issues
                if str(issue.get("kind") or "").strip().lower() == "missing_data"
            ]
            doctor_report["plan"] = []
            doctor_report["issues"] = pack_issues
            # Keep the structured consilium format even in fail-closed mode so
            # timeline/biomarker context and source-basis lines remain visible.
            doctor_report["consilium_md"] = build_consilium_md(
                query_type=query_type,
                case_facts=case_facts,
                plan_sections=doctor_report["plan"],
                issues=pack_issues,
                sufficiency=sufficiency,
                source_ids=(
                    bridge_context.source_ids
                    if isinstance(getattr(bridge_context, "source_ids", None), list)
                    else [str(item.get("source_id") or "") for item in citations]
                ),
                has_real_evidence=False,
                timeline_reconciliation=timeline_reconciliation,
            )
        if not reranked and not insufficient_status:
            insufficient_status = True
            insufficient_reason = "Недостаточно доказательств из предоставленных рекомендаций."
        insufficient_reason = normalize_ru_clinical_text(insufficient_reason)

        guided = build_guided_report(
            query_type=query_type,
            disease_context=disease_context,
            case_facts=case_facts,
            timeline=timeline,
            plan_sections=plan_sections,
            issues=pack_issues,
            citations=citations,
            insufficient_data={"status": insufficient_status, "reason": insufficient_reason},
        )
        summary_md = str(guided.get("doctor_summary_md") or "").strip()
        if not summary_md:
            summary_md = self._build_pack_summary_md(
                issues=pack_issues,
                query_type=query_type,
                insufficient_status=insufficient_status,
            )
        doctor_report["summary_md"] = summary_md

        llm_summary = ""
        if citations:
            llm_summary = normalize_ru_clinical_text(str(legacy_doctor_report.get("summary") or ""))
        guided_doctor_summary = normalize_ru_clinical_text(str(guided.get("doctor_summary_plain") or ""))
        doctor_report["summary"] = llm_summary or guided_doctor_summary or normalize_ru_clinical_text(str(summary_md))
        doctor_report["verification_summary"] = self._build_verification_summary(
            issues=pack_issues,
            insufficient_status=insufficient_status,
        )

        patient_context_report = dict(doctor_report)
        patient_context_report["kb_version"] = str(legacy_doctor_report.get("kb_version") or self._kb_version)
        patient_context_report["missing_data"] = [
            {
                "field": str(item.get("summary") or item.get("kind") or "missing_data"),
                "reason": str(item.get("details") or ""),
            }
            for item in pack_issues
            if str(item.get("kind") or "").strip().lower() == "missing_data"
        ][:10]
        patient_context_report["insufficient_data"] = {
            "status": insufficient_status,
            "reason": insufficient_reason,
        }

        strict_patient, patient_generation_path = build_patient_explain_with_fallback(
            doctor_report=patient_context_report,
            llm_router=self.llm_router,
            prompt_registry=self.prompt_registry,
            prompt_schema_strict=bool(self.settings.oncoai_prompt_schema_strict),
            strict_llm_only=self._llm_rag_only_mode,
        )
        if self._strict_fail_closed and str(patient_generation_path or "").strip().lower() != "llm":
            raise ValidationError("STRICT_FAIL_CLOSED: patient explain generation must use LLM path in strict_full")
        source_ids_for_patient = normalize_source_set_ids([str(item.get("source_id") or "") for item in citations])
        if not source_ids_for_patient:
            source_ids_for_patient = [str(item.get("source_id") or "") for item in citations if str(item.get("source_id") or "").strip()]
        checked_lines = [
            "Стадия, биомаркеры и история лечения сопоставлены с клиническими рекомендациями.",
            "Проверены риски и полнота данных перед выбором следующего этапа.",
        ]
        normalized_sources = set(normalize_source_set_ids(source_ids_for_patient))
        if "minzdrav" in normalized_sources:
            checked_lines.append("Использованы подтверждённые фрагменты клинических рекомендаций Минздрава РФ.")
        if "russco" in normalized_sources:
            checked_lines.append("Использованы подтверждённые фрагменты рекомендаций RUSSCO.")
        patient_explain = map_strict_to_pack_patient_v1_2(
            strict_payload=strict_patient,
            request_id=request_id,
            source_ids=source_ids_for_patient,
            what_was_checked=checked_lines,
            drug_safety=build_patient_drug_safety(drug_safety).model_dump(),
        )

        guided_patient_summary = normalize_ru_clinical_text(str(guided.get("patient_summary_plain") or ""))
        guided_patient_points = (
            guided.get("patient_key_points")
            if isinstance(guided.get("patient_key_points"), list)
            else []
        )
        if str(patient_generation_path or "").strip().lower() != "llm":
            if guided_patient_summary:
                patient_explain["summary_plain"] = guided_patient_summary
            if guided_patient_points:
                patient_explain["key_points"] = [
                    normalize_ru_clinical_text(str(item))
                    for item in guided_patient_points
                    if str(item).strip()
                ][:6]
        else:
            if len(str(patient_explain.get("summary_plain") or "").strip()) < 40 and guided_patient_summary:
                patient_explain["summary_plain"] = guided_patient_summary
            if not isinstance(patient_explain.get("key_points"), list) or not patient_explain.get("key_points"):
                if guided_patient_points:
                    patient_explain["key_points"] = [
                        normalize_ru_clinical_text(str(item))
                        for item in guided_patient_points
                        if str(item).strip()
                    ][:6]

        effective_llm_path = self._resolve_run_meta_llm_path(llm_path, report_generation_path)
        run_meta = build_run_meta(
            request_id=request_id,
            retrieved=retrieved,
            reranked=reranked,
            citations_count=len(citations),
            llm_path=effective_llm_path,
            report_generation_path=self._map_generation_path_for_pack(report_generation_path),
            fallback_reason=self._map_fallback_reason_for_pack(fallback_reason),
            reasoning_mode=self._reasoning_mode,
            routing_meta=routing_meta,
            started_at_perf=started,
            settings=self.settings,
        )
        if self._llm_rag_only_mode:
            if str(run_meta.get("llm_path") or "").strip() != "primary":
                raise ValidationError("STRICT_LLM_RAG_ONLY: run_meta.llm_path must be primary")
            if str(run_meta.get("report_generation_path") or "").strip() != "primary":
                raise ValidationError("STRICT_LLM_RAG_ONLY: run_meta.report_generation_path must be primary")
            run_meta["fallback_reason"] = "none"
        run_meta["patient_generation_path"] = str(patient_generation_path or "deterministic")

        return {
            "schema_version": "0.2",
            "request_id": request_id,
            "doctor_report": doctor_report,
            "patient_explain": patient_explain,
            "run_meta": run_meta,
            "meta": self._build_execution_meta(),
            "insufficient_data": {
                "status": insufficient_status,
                "reason": insufficient_reason,
            },
        }

    @staticmethod
    def _ensure_issue_evidence_source_coverage(
        doctor_report: dict[str, Any],
        reranked: list[dict[str, Any]],
    ) -> dict[str, Any]:
        issues = doctor_report.get("issues") if isinstance(doctor_report.get("issues"), list) else []
        if not issues or not reranked:
            return doctor_report

        represented_sources = {
            str(evidence.get("source_set") or "").strip()
            for issue in issues
            if isinstance(issue, dict)
            for evidence in issue.get("evidence", [])
            if isinstance(evidence, dict) and str(evidence.get("source_set") or "").strip()
        }
        chunk_by_source: dict[str, dict[str, Any]] = {}
        for chunk in reranked:
            source_set = str(chunk.get("source_set") or "").strip()
            if source_set and source_set not in chunk_by_source:
                chunk_by_source[source_set] = chunk

        missing_sources = [source for source in chunk_by_source.keys() if source not in represented_sources]
        if not missing_sources:
            return doctor_report

        first_issue = issues[0] if isinstance(issues[0], dict) else None
        if first_issue is None:
            return doctor_report
        evidence_list = first_issue.get("evidence") if isinstance(first_issue.get("evidence"), list) else []
        for source in missing_sources:
            chunk = chunk_by_source[source]
            evidence_list.append(
                {
                    "doc_id": chunk.get("doc_id", "unknown_doc"),
                    "doc_version": chunk.get("doc_version", "unknown_version"),
                    "source_set": chunk.get("source_set", source),
                    "cancer_type": chunk.get("cancer_type", ""),
                    "language": chunk.get("language", ""),
                    "section_title": chunk.get("section_title", "Guideline fragment"),
                    "pdf_page_index": int(chunk.get("pdf_page_index", 0) or 0),
                    "page_label": chunk.get("page_label", "1"),
                    "chunk_id": chunk.get("chunk_id", f"chunk:{source}"),
                    "quote": str(chunk.get("text", ""))[:220],
                }
            )
        first_issue["evidence"] = evidence_list
        return doctor_report

    def _step_4_build_doctor_report(
        self,
        schema_version: str,
        cancer_type: str,
        query_type: str,
        plan_structured: list[dict[str, Any]],
        reranked: list[dict[str, Any]],
    ) -> tuple[dict[str, Any], str, str | None]:
        diff_issues = [] if self._llm_rag_only_mode else compute_diff(cancer_type, plan_structured, query_type=query_type)
        try:
            strict_report, report_generation_path, fallback_reason = build_doctor_report_with_fallback(
                kb_version=self._kb_version,
                diff_issues=diff_issues,
                retrieved_chunks=reranked,
                llm_router=self.llm_router,
                prompt_registry=self.prompt_registry,
                prompt_schema_strict=bool(self.settings.oncoai_prompt_schema_strict),
                fail_closed=self._strict_fail_closed,
                strict_llm_only=self._llm_rag_only_mode,
            )
        except RuntimeError as exc:
            if self._strict_fail_closed:
                raise ValidationError(f"STRICT_FAIL_CLOSED: {exc}") from exc
            raise
        chunks_by_id = {chunk["chunk_id"]: chunk for chunk in reranked}
        doctor_report = map_strict_to_public_report(strict_report, chunks_by_id=chunks_by_id)
        doctor_report["schema_version"] = schema_version
        doctor_report = self._ensure_issue_evidence_source_coverage(doctor_report=doctor_report, reranked=reranked)
        doctor_report = enforce_retrieved_evidence(
            doctor_report,
            retrieved_chunk_ids=set(chunks_by_id.keys()),
            downgrade_invalid_to_data_quality=False,
        )

        doctor_report = self._sync_summary_with_visible_issues(doctor_report)
        doctor_report["summary"] = normalize_ru_clinical_text(str(doctor_report.get("summary") or ""))
        issues_raw = doctor_report.get("issues") if isinstance(doctor_report.get("issues"), list) else []
        normalized_issues: list[dict[str, Any]] = []
        for issue in issues_raw:
            if not isinstance(issue, dict):
                continue
            item = dict(issue)
            item["title"] = normalize_ru_clinical_text(str(item.get("title") or ""))
            item["description"] = normalize_ru_clinical_text(str(item.get("description") or ""))
            normalized_issues.append(item)
        doctor_report["issues"] = normalized_issues
        validate_doctor_report(doctor_report)
        return doctor_report, report_generation_path, fallback_reason

    def _step_5_build_patient_explain(
        self,
        payload: dict[str, Any],
        schema_version: str,
        doctor_report: dict[str, Any],
    ) -> dict[str, Any] | None:
        if payload.get("return_patient_explain", True):
            strict_patient, _patient_path = build_patient_explain_with_fallback(
                doctor_report=doctor_report,
                llm_router=self.llm_router,
                prompt_registry=self.prompt_registry,
                prompt_schema_strict=bool(self.settings.oncoai_prompt_schema_strict),
                strict_llm_only=self._llm_rag_only_mode,
            )
            if self._strict_fail_closed and str(_patient_path or "").strip().lower() != "llm":
                raise ValidationError("STRICT_FAIL_CLOSED: patient explain generation must use LLM path in strict_full")
            patient_explain = map_strict_to_public_patient(strict_patient)
            patient_explain["schema_version"] = schema_version
            return patient_explain
        return None

    def _step_6_finalize_response_and_persist(
        self,
        schema_version: str,
        doctor_report: dict[str, Any],
        patient_explain: dict[str, Any] | None,
        retrieved: list[dict[str, Any]],
        reranked: list[dict[str, Any]],
        llm_path: str,
        report_generation_path: str,
        fallback_reason: str | None,
        routing_meta: dict[str, Any] | None,
        started: float,
    ) -> dict[str, Any]:
        response: dict[str, Any] = {"doctor_report": doctor_report}
        if patient_explain is not None:
            response["patient_explain"] = patient_explain

        combined_fallback_reason = fallback_reason
        if self._retrieval_fallback_reason:
            combined_fallback_reason = (
                f"{self._retrieval_fallback_reason};{fallback_reason}"
                if fallback_reason
                else self._retrieval_fallback_reason
            )

        insufficient_reason = ""
        if not reranked:
            insufficient_reason = "Не найдено подтверждающих фрагментов рекомендаций для выбранных фильтров и запроса."
        elif not doctor_report.get("issues"):
            insufficient_reason = "По текущему кейсу не удалось сформировать утверждения с доказательной поддержкой."
        insufficient_reason = normalize_ru_clinical_text(insufficient_reason)

        response["insufficient_data"] = {
            "status": bool(insufficient_reason),
            "reason": insufficient_reason or "Достаточно подтверждающих данных.",
        }
        effective_llm_path = self._resolve_run_meta_llm_path(llm_path, report_generation_path)

        if schema_version == "0.2":
            report_path_value = str(report_generation_path or "").strip().lower()
            llm_path_value = str(effective_llm_path or "").strip().lower()
            fallback_value = str(combined_fallback_reason or "").strip()
            if self._llm_rag_only_mode:
                if llm_path_value != "primary":
                    raise ValidationError("STRICT_LLM_RAG_ONLY: llm_path must be primary")
                if report_path_value not in {"llm_primary", "primary"}:
                    raise ValidationError("STRICT_LLM_RAG_ONLY: report_generation_path must be llm_primary")
                report_path_value = "llm_primary"
                fallback_value = "none"
            response["run_meta"] = {
                "retrieval_k": len(retrieved),
                "rerank_n": len(reranked),
                "llm_path": "primary" if self._llm_rag_only_mode else effective_llm_path,
                "latency_ms_total": round((time.perf_counter() - started) * 1000.0, 2),
                "kb_version": self._kb_version,
                "vector_backend": self.settings.vector_backend,
                "embedding_backend": self.settings.embedding_backend,
                "reranker_backend": self.settings.reranker_backend,
                "report_generation_path": report_path_value,
                "reasoning_mode": self._reasoning_mode,
                "retrieval_engine": self._retrieval_engine,
            }
            if fallback_value:
                response["run_meta"]["fallback_reason"] = fallback_value
            if isinstance(routing_meta, dict):
                response["run_meta"]["routing_meta"] = {
                    "resolved_disease_id": str(routing_meta.get("resolved_disease_id") or "unknown_disease"),
                    "resolved_cancer_type": str(routing_meta.get("resolved_cancer_type") or "unknown"),
                    "match_strategy": str(routing_meta.get("match_strategy") or "default_sources_fallback"),
                    "ambiguous_brain_scope": bool(routing_meta.get("ambiguous_brain_scope")),
                    "source_ids": [
                        str(item).strip()
                        for item in (routing_meta.get("source_ids") if isinstance(routing_meta.get("source_ids"), list) else [])
                        if str(item).strip()
                    ],
                    "doc_ids": [
                        str(item).strip()
                        for item in (routing_meta.get("doc_ids") if isinstance(routing_meta.get("doc_ids"), list) else [])
                        if str(item).strip()
                    ],
                    "candidate_chunks": int(routing_meta.get("candidate_chunks") or 0),
                    "baseline_candidate_chunks": int(routing_meta.get("baseline_candidate_chunks") or 0),
                    "reduction_ratio": float(routing_meta.get("reduction_ratio") or 0.0),
                }
            response["meta"] = self._build_execution_meta()

        validate_analyze_response(response)
        self.store.save_report(
            report_id=doctor_report["report_id"],
            payload=response,
            created_at=datetime.now(timezone.utc).isoformat(),
        )
        safe_log(
            self.logger,
            "analyze.completed",
            {
                "schema_version": schema_version,
                "retrieval_k": len(retrieved),
                "rerank_n": len(reranked),
                "llm_path": effective_llm_path,
                "report_generation_path": report_generation_path,
                **({"fallback_reason": combined_fallback_reason} if combined_fallback_reason else {}),
            },
        )
        return response

    def _build_compat_projection_report(self, response: dict[str, Any]) -> dict[str, Any]:
        doctor_report = response.get("doctor_report") if isinstance(response.get("doctor_report"), dict) else {}
        patient_explain = response.get("patient_explain") if isinstance(response.get("patient_explain"), dict) else {}
        run_meta = response.get("run_meta") if isinstance(response.get("run_meta"), dict) else {}
        insufficient_data = response.get("insufficient_data") if isinstance(response.get("insufficient_data"), dict) else {}

        doctor_projection = project_doctor_report_v1_1(
            doctor_report_v1_2=doctor_report,
            run_meta=run_meta,
            insufficient_data=insufficient_data,
        )
        doctor_errors = validate_doctor_projection_v1_1(doctor_projection)

        patient_projection: dict[str, Any] | None = None
        patient_errors: list[str] = []
        if patient_explain:
            patient_projection = project_patient_explain_alt_profile(
                patient_v1_2=patient_explain,
                doctor_report_v1_2=doctor_report,
            )
            patient_errors = validate_patient_projection_alt(patient_projection)

        return {
            "doctor_report_v1_1": {
                "valid": len(doctor_errors) == 0,
                "errors": doctor_errors,
                "schema_path": "docs/contracts/external/doctor_report_v1_1.schema.json",
                "projection": doctor_projection,
            },
            "patient_explain_alt": {
                "valid": len(patient_errors) == 0,
                "errors": patient_errors,
                "schema_path": "docs/contracts/external/patient_explain_v1_2_alt.schema.json",
                "projection": patient_projection,
            },
        }

    def _enforce_compat_projection_gate(self, response: dict[str, Any]) -> None:
        if not bool(self.settings.oncoai_compat_v1_1_projection_enabled):
            return
        doctor_report = response.get("doctor_report") if isinstance(response.get("doctor_report"), dict) else {}
        if str(doctor_report.get("schema_version") or "") != "1.2":
            return

        report = self._build_compat_projection_report(response)
        doctor_status = report.get("doctor_report_v1_1") if isinstance(report.get("doctor_report_v1_1"), dict) else {}
        patient_status = report.get("patient_explain_alt") if isinstance(report.get("patient_explain_alt"), dict) else {}

        errors: list[str] = []
        if not bool(doctor_status.get("valid")):
            errors.extend([str(item) for item in (doctor_status.get("errors") if isinstance(doctor_status.get("errors"), list) else [])])
        if response.get("patient_explain") is not None and not bool(patient_status.get("valid")):
            errors.extend([str(item) for item in (patient_status.get("errors") if isinstance(patient_status.get("errors"), list) else [])])
        if errors:
            raise ValidationError("COMPAT_PROJECTION_VALIDATION_FAILED: " + "; ".join(errors[:5]))

    def admin_validate_contract_projections(
        self,
        role: str,
        payload: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        body = payload if isinstance(payload, dict) else {}
        include_projection = bool(body.get("include_projection", False))
        report_id = str(body.get("report_id") or "").strip()
        response = body.get("response") if isinstance(body.get("response"), dict) else None
        if response is None and report_id:
            response = self.store.get_report(report_id)
        if not isinstance(response, dict):
            raise ValidationError("Provide `report_id` or `response` payload for projection validation")

        projection_report = self._build_compat_projection_report(response)
        if not include_projection:
            for key in ("doctor_report_v1_1", "patient_explain_alt"):
                section = projection_report.get(key)
                if isinstance(section, dict) and "projection" in section:
                    section.pop("projection", None)

        doctor_valid = bool(projection_report.get("doctor_report_v1_1", {}).get("valid"))
        patient_section = projection_report.get("patient_explain_alt") if isinstance(projection_report.get("patient_explain_alt"), dict) else {}
        patient_valid = bool(patient_section.get("valid")) if response.get("patient_explain") is not None else True
        return {
            "status": "ok" if doctor_valid and patient_valid else "error",
            "report_id": report_id or str((response.get("doctor_report") or {}).get("report_id") or ""),
            "compatibility": projection_report,
        }

    def analyze(self, payload: dict[str, Any], role: str = "clinician", client_id: str = "anonymous") -> dict[str, Any]:
        started = time.perf_counter()
        resolved_payload = self._resolve_pack_case_reference(payload)
        bridge_context = normalize_analyze_request(resolved_payload)
        normalized_payload = bridge_context.normalized_payload
        query_type = str(bridge_context.query_type or "CHECK_LAST_TREATMENT").strip().upper()
        if query_type not in {"NEXT_STEPS", "CHECK_LAST_TREATMENT"}:
            query_type = "CHECK_LAST_TREATMENT"
        if bridge_context.dialect != DIALECT_PACK_V2:
            query_type = "NEXT_STEPS"

        schema_version = self._step_1_validate_and_guard_input(payload=normalized_payload, role=role, client_id=client_id)
        (
            plan_structured,
            filters,
            query,
            source_sets,
            doc_ids,
            route_pairs,
            route_triples,
            routing_meta,
        ) = self._step_2_normalize_and_prepare_retrieval(
            payload=normalized_payload,
            as_of_date=bridge_context.as_of_date,
            historical_assessment_requested=bool(bridge_context.historical_assessment),
        )

        case_text = self._build_case_text_for_casefacts(
            normalized_payload=normalized_payload,
            case_json=bridge_context.case_json,
        )
        case_facts_for_planner = (
            extract_case_facts(case_text=case_text, case_json=bridge_context.case_json).model_dump()
            if bool(self.settings.oncoai_casefacts_enabled)
            else {
                "initial_stage": {},
                "current_stage": {},
                "biomarkers": {},
                "metastases": [],
                "treatment_history": [],
                "complications": [],
                "key_unknowns": [],
            }
        )
        disease_context_hint = {
            "line": (
                normalized_payload.get("case", {}).get("diagnosis", {}).get("line")
                if isinstance(normalized_payload.get("case"), dict)
                else None
            )
        }
        if self._llm_rag_only_mode:
            generated_plan_sections = []
        else:
            generated_plan_sections = build_next_steps_plan_sections(
                query_type=query_type,
                case_facts=case_facts_for_planner,
                disease_context=disease_context_hint,
            )
            if query_type == "NEXT_STEPS":
                candidate_plan_structured = flatten_plan_for_diff(generated_plan_sections)
                if candidate_plan_structured:
                    plan_structured = candidate_plan_structured

        if self._llm_rag_only_mode:
            try:
                query_bundle = build_query_bundle_with_llm(
                    llm_router=self.llm_router,
                    base_query=query,
                    query_type=query_type,
                    cancer_type=str(normalized_payload.get("case", {}).get("cancer_type", "unknown"))
                    if isinstance(normalized_payload.get("case"), dict)
                    else "unknown",
                    case_payload=normalized_payload.get("case") if isinstance(normalized_payload.get("case"), dict) else {},
                    plan_sections=generated_plan_sections,
                )
            except RuntimeError as exc:
                raise ValidationError(f"STRICT_LLM_RAG_ONLY: query bundle planning failed: {exc}") from exc
        else:
            query_bundle = build_query_bundle(
                base_query=query,
                query_type=query_type,
                cancer_type=str(normalized_payload.get("case", {}).get("cancer_type", "unknown"))
                if isinstance(normalized_payload.get("case"), dict)
                else "unknown",
                case_facts=case_facts_for_planner,
                plan_sections=generated_plan_sections,
            )
        try:
            retrieved, reranked, llm_path, baseline_candidates = self._step_3_retrieve_rerank_and_probe_llm(
                query=query,
                query_bundle=query_bundle,
                filters=filters,
                source_sets=source_sets,
                doc_ids=doc_ids,
                route_pairs=route_pairs,
                route_triples=route_triples,
            )
        except RuntimeError as exc:
            if self._strict_fail_closed:
                raise ValidationError(f"STRICT_FAIL_CLOSED: {exc}") from exc
            raise
        routing_meta["candidate_chunks"] = len(retrieved)
        routing_meta["baseline_candidate_chunks"] = int(baseline_candidates)
        if baseline_candidates > 0:
            routing_meta["reduction_ratio"] = round(
                max(0.0, 1.0 - (float(len(retrieved)) / float(baseline_candidates))),
                4,
            )
        else:
            routing_meta["reduction_ratio"] = 0.0
        doctor_report, report_generation_path, fallback_reason = self._step_4_build_doctor_report(
            schema_version=schema_version,
            cancer_type=normalized_payload["case"]["cancer_type"],
            query_type=query_type,
            plan_structured=plan_structured,
            reranked=reranked,
        )

        use_native_pack_v1_2 = bridge_context.dialect == DIALECT_PACK_V2 and bool(self.settings.oncoai_doctor_schema_v1_2_enabled)
        if self._llm_rag_only_mode and not use_native_pack_v1_2:
            raise ValidationError("STRICT_LLM_RAG_ONLY: doctor schema v1.2 native pack path is required")
        if use_native_pack_v1_2:
            native_response = self._build_native_pack_v1_2_response(
                bridge_context=bridge_context,
                normalized_payload=normalized_payload,
                retrieved=retrieved,
                reranked=reranked,
                llm_path=llm_path,
                report_generation_path=report_generation_path,
                fallback_reason=fallback_reason,
                routing_meta=routing_meta,
                started=started,
                legacy_doctor_report=doctor_report,
                generated_plan_sections=generated_plan_sections,
            )
            if str(bridge_context.query_mode or "FULL_ANALYSIS").strip().upper() == "SOURCES_ONLY":
                native_response = self._apply_sources_only_mode(native_response)
            historical_assessment = self._build_historical_assessment(
                as_of_date=bridge_context.as_of_date,
                historical_assessment_requested=bool(bridge_context.historical_assessment),
                source_ids=bridge_context.source_ids,
                citations=(
                    native_response.get("doctor_report", {}).get("citations")
                    if isinstance(native_response.get("doctor_report"), dict)
                    else []
                ),
                routing_meta=routing_meta,
            )
            if historical_assessment is not None:
                native_response["historical_assessment"] = historical_assessment
            doctor_report_payload = native_response.get("doctor_report")
            if isinstance(doctor_report_payload, dict):
                self._validate_comparative_claims_policy(doctor_report_payload)
            self._enforce_compat_projection_gate(native_response)
            validate_analyze_response(native_response, allow_pack_legacy_v1_0=False)
            self.store.save_report(
                report_id=native_response["doctor_report"]["report_id"],
                payload=native_response,
                created_at=datetime.now(timezone.utc).isoformat(),
            )
            return native_response

        patient_explain = self._step_5_build_patient_explain(
            payload=normalized_payload,
            schema_version=schema_version,
            doctor_report=doctor_report,
        )
        legacy_response = self._step_6_finalize_response_and_persist(
            schema_version=schema_version,
            doctor_report=doctor_report,
            patient_explain=patient_explain,
            retrieved=retrieved,
            reranked=reranked,
            llm_path=llm_path,
            report_generation_path=report_generation_path,
            fallback_reason=fallback_reason,
            routing_meta=routing_meta,
            started=started,
        )

        response = serialize_analyze_response(
            context=bridge_context,
            legacy_response=legacy_response,
            doctor_schema_v1_2_enabled=bool(self.settings.oncoai_doctor_schema_v1_2_enabled),
            casefacts_enabled=bool(self.settings.oncoai_casefacts_enabled),
        )
        response["meta"] = self._build_execution_meta()
        if str(bridge_context.query_mode or "FULL_ANALYSIS").strip().upper() == "SOURCES_ONLY":
            response = self._apply_sources_only_mode(response)
        historical_assessment = self._build_historical_assessment(
            as_of_date=bridge_context.as_of_date,
            historical_assessment_requested=bool(bridge_context.historical_assessment),
            source_ids=bridge_context.source_ids,
            citations=(
                response.get("doctor_report", {}).get("citations")
                if isinstance(response.get("doctor_report"), dict)
                else []
            ),
            routing_meta=routing_meta,
        )
        if historical_assessment is not None:
            response["historical_assessment"] = historical_assessment
        doctor_report_payload = response.get("doctor_report")
        if isinstance(doctor_report_payload, dict):
            self._validate_comparative_claims_policy(doctor_report_payload)
        self._enforce_compat_projection_gate(response)
        if response is not legacy_response:
            validate_analyze_response(
                response,
                allow_pack_legacy_v1_0=not bool(self.settings.oncoai_doctor_schema_v1_2_enabled),
            )
            self.store.save_report(
                report_id=response["doctor_report"]["report_id"],
                payload=response,
                created_at=datetime.now(timezone.utc).isoformat(),
            )
        return response

    def case_import(self, role: str, payload: dict[str, Any]) -> dict[str, Any]:
        ensure_role(role, {"admin", "clinician"})
        raw_profile = str(payload.get("import_profile") or "UNKNOWN").strip().upper() if isinstance(payload, dict) else "UNKNOWN"
        raw_case_id = str(payload.get("case_id") or "").strip() if isinstance(payload, dict) else ""
        try:
            normalized_import_payload = normalize_case_import_payload(payload)
        except ValidationError as exc:
            started_at = datetime.now(timezone.utc).isoformat()
            finished_at = datetime.now(timezone.utc).isoformat()
            failed_run = self._build_failed_case_import_run(
                import_run_id=str(uuid.uuid4()),
                case_id=raw_case_id or str(uuid.uuid4()),
                import_profile=raw_profile or "UNKNOWN",
                started_at=started_at,
                finished_at=finished_at,
                error_code="INVALID_IMPORT_PAYLOAD",
                error_message=str(exc),
            )
            self.store.save_case_import_run(failed_run)
            return failed_run

        import_profile = str(normalized_import_payload.get("import_profile") or "").strip().upper()
        case_id = str(normalized_import_payload.get("case_id") or "").strip() or str(uuid.uuid4())
        data_mode = self._normalize_case_data_mode(normalized_import_payload.get("data_mode"))
        full_mode_acknowledged = bool(normalized_import_payload.get("full_mode_acknowledged"))
        started_at = datetime.now(timezone.utc).isoformat()
        finished_at = datetime.now(timezone.utc).isoformat()
        import_run_id = str(uuid.uuid4())

        if data_mode == "FULL" and not bool(getattr(self.settings, "case_import_allow_full_mode", False)):
            failed_run = self._build_failed_case_import_run(
                import_run_id=import_run_id,
                case_id=case_id,
                import_profile=import_profile or "UNKNOWN",
                started_at=started_at,
                finished_at=finished_at,
                error_code="FULL_MODE_DISABLED",
                error_message="FULL mode import is disabled for this environment (set CASE_IMPORT_ALLOW_FULL_MODE=true).",
            )
            self.store.save_case_import_run(failed_run)
            return failed_run

        if (
            data_mode == "FULL"
            and bool(getattr(self.settings, "case_import_full_require_ack", True))
            and not full_mode_acknowledged
        ):
            failed_run = self._build_failed_case_import_run(
                import_run_id=import_run_id,
                case_id=case_id,
                import_profile=import_profile or "UNKNOWN",
                started_at=started_at,
                finished_at=finished_at,
                error_code="FULL_MODE_ACK_REQUIRED",
                error_message="FULL mode import requires explicit full_mode_acknowledged=true.",
            )
            self.store.save_case_import_run(failed_run)
            return failed_run

        if import_profile not in SUPPORTED_CASE_IMPORT_PROFILES:
            failed_run = self._build_failed_case_import_run(
                import_run_id=import_run_id,
                case_id=case_id,
                import_profile=import_profile or "UNKNOWN",
                started_at=started_at,
                finished_at=finished_at,
                error_code="PROFILE_NOT_SUPPORTED",
                error_message=f"Import profile is not enabled in this stage: {import_profile or 'UNKNOWN'}",
            )
            self.store.save_case_import_run(failed_run)
            return failed_run

        raw_case_json = normalized_import_payload.get("case_json")
        warnings: list[dict[str, str]] = []
        raw_file_warnings = normalized_import_payload.get("__file_warnings")
        if isinstance(raw_file_warnings, list):
            for item in raw_file_warnings:
                if not isinstance(item, dict):
                    continue
                code = str(item.get("code") or "").strip()
                message = str(item.get("message") or "").strip()
                if code and message:
                    warnings.append({"code": code, "message": message})
        try:
            if isinstance(raw_case_json, dict):
                normalized_case, missing_required_fields = self._normalize_imported_case(
                    case_json=raw_case_json,
                    import_profile=import_profile,
                    case_id=case_id,
                    now=started_at,
                    data_mode=data_mode,
                )
            else:
                if import_profile == "FHIR_BUNDLE":
                    generated_case, missing_required_fields, warnings = self._build_case_json_from_fhir_bundle(
                        payload=normalized_import_payload,
                        case_id=case_id,
                        now=started_at,
                        data_mode=data_mode,
                    )
                elif self._llm_rag_only_mode and import_profile in {"FREE_TEXT", "CUSTOM_TEMPLATE"}:
                    generated_case, missing_required_fields = self._build_case_json_from_import_llm(
                        import_profile=import_profile,
                        payload=normalized_import_payload,
                        case_id=case_id,
                        now=started_at,
                        data_mode=data_mode,
                    )
                elif self._llm_rag_only_mode and import_profile == "KIN_PDF":
                    generated_case, missing_required_fields, warnings = self._build_case_json_from_kin_pdf_llm(
                        payload=normalized_import_payload,
                        case_id=case_id,
                        now=started_at,
                        data_mode=data_mode,
                    )
                elif import_profile in {"FREE_TEXT", "CUSTOM_TEMPLATE"}:
                    generated_case, missing_required_fields = self._build_case_json_from_import(
                        import_profile=import_profile,
                        payload=normalized_import_payload,
                        case_id=case_id,
                        now=started_at,
                        data_mode=data_mode,
                    )
                else:  # KIN_PDF
                    generated_case, missing_required_fields, warnings = self._build_case_json_from_kin_pdf(
                        payload=normalized_import_payload,
                        case_id=case_id,
                        now=started_at,
                        data_mode=data_mode,
                    )

                normalized_case, inherited_missing = self._normalize_imported_case(
                    case_json=generated_case,
                    import_profile=import_profile,
                    case_id=case_id,
                    now=started_at,
                    data_mode=data_mode,
                )
                missing_required_fields = sorted(set(missing_required_fields + inherited_missing))
        except ValidationError as exc:
            failed_run = self._build_failed_case_import_run(
                import_run_id=import_run_id,
                case_id=case_id,
                import_profile=import_profile,
                started_at=started_at,
                finished_at=finished_at,
                error_code="INVALID_IMPORT_PAYLOAD",
                error_message=str(exc),
            )
            self.store.save_case_import_run(failed_run)
            return failed_run

        if not self._llm_rag_only_mode:
            self._apply_icd10_inference_to_case(
                case_json=normalized_case,
                missing_required_fields=missing_required_fields,
                warnings=warnings,
            )

        if data_mode == "DEID":
            normalized_case, pii_redacted = self._apply_deid_redaction(case_json=normalized_case)
            if pii_redacted:
                warnings.append(
                    {
                        "code": "PII_REDACTED_DEID",
                        "message": "PII fragments were redacted in DEID mode.",
                    }
                )
        else:
            warnings.append(
                {
                    "code": "FULL_MODE_ENABLED",
                    "message": "FULL mode import stored non-redacted clinical payload.",
                }
            )

        missing_required_fields = sorted(set(str(item) for item in missing_required_fields if str(item)))

        stored_case_id = self.store.upsert_case(normalized_case)
        status = "PARTIAL_SUCCESS" if missing_required_fields else "SUCCESS"
        if isinstance(raw_case_json, dict):
            confidence = 1.0 if not missing_required_fields else 0.9
        elif import_profile in {"FHIR_BUNDLE", "KIN_PDF"}:
            confidence = 0.75 if not missing_required_fields else 0.65
        else:
            confidence = 0.8 if not missing_required_fields else 0.7

        deduped_warnings: list[dict[str, str]] = []
        warning_keys: set[tuple[str, str]] = set()
        for item in warnings:
            code = str(item.get("code") or "").strip()
            message = str(item.get("message") or "").strip()
            if not code or not message:
                continue
            key = (code, message)
            if key in warning_keys:
                continue
            warning_keys.add(key)
            deduped_warnings.append({"code": code, "message": message})

        result = {
            "schema_version": "1.0",
            "import_run_id": import_run_id,
            "case_id": stored_case_id,
            "import_profile": import_profile,
            "started_at": started_at,
            "finished_at": finished_at,
            "status": status,
            "confidence": confidence,
            "missing_required_fields": missing_required_fields,
            "warnings": deduped_warnings,
            "errors": [],
        }
        self.store.save_case_import_run(result)
        return result

    def case_import_file_base64(self, role: str, payload: dict[str, Any]) -> dict[str, Any]:
        ensure_role(role, {"admin", "clinician"})
        if not isinstance(payload, dict):
            raise ValidationError("Payload must be an object")
        import_payload = self._build_case_import_payload_from_file(payload)
        if payload.get("full_mode_acknowledged") is not None:
            import_payload["full_mode_acknowledged"] = bool(payload.get("full_mode_acknowledged"))
        return self.case_import(role=role, payload=import_payload)

    def _build_batch_merged_case(
        self,
        *,
        case_ids: list[str],
        data_mode: str = "DEID",
    ) -> dict[str, Any] | None:
        loaded_cases: list[dict[str, Any]] = []
        for case_id in case_ids:
            payload = self.store.get_case(case_id)
            if isinstance(payload, dict):
                loaded_cases.append(payload)
        if not loaded_cases:
            return None

        first_case = loaded_cases[0]
        first_patient = first_case.get("patient") if isinstance(first_case.get("patient"), dict) else {}
        first_diagnosis = {}
        diagnoses = first_case.get("diagnoses")
        if isinstance(diagnoses, list) and diagnoses and isinstance(diagnoses[0], dict):
            first_diagnosis = dict(diagnoses[0])

        merged_timeline: list[dict[str, Any]] = []
        timeline_seen: set[tuple[str, str, str]] = set()
        merged_notes_parts: list[str] = []
        merged_comorbidities: set[str] = set()
        merged_contraindications: set[str] = set()

        for idx, case_payload in enumerate(loaded_cases, start=1):
            notes = str(case_payload.get("notes") or "").strip()
            if notes:
                merged_notes_parts.append(f"[Doc {idx}] {notes}")
            for list_key, target in (
                ("comorbidities", merged_comorbidities),
                ("contraindications", merged_contraindications),
            ):
                values = case_payload.get(list_key)
                if isinstance(values, list):
                    for item in values:
                        text = str(item).strip()
                        if text:
                            target.add(text)

            diagnoses = case_payload.get("diagnoses")
            if not isinstance(diagnoses, list) or not diagnoses:
                continue
            primary = diagnoses[0] if isinstance(diagnoses[0], dict) else {}
            timeline = primary.get("timeline")
            if not isinstance(timeline, list):
                continue
            for event in timeline:
                if not isinstance(event, dict):
                    continue
                date_value = str(event.get("date") or "").strip()
                type_value = str(event.get("type") or "other").strip().lower() or "other"
                label_value = str(event.get("label") or event.get("details") or "").strip()
                if not label_value:
                    continue
                key = (date_value.lower(), type_value, label_value.lower())
                if key in timeline_seen:
                    continue
                timeline_seen.add(key)
                merged_timeline.append(
                    {
                        "date": date_value,
                        "precision": str(event.get("precision") or "unknown"),
                        "type": type_value,
                        "label": label_value,
                        "details": str(event.get("details") or "").strip(),
                    }
                )

        merged_timeline.sort(key=lambda item: str(item.get("date") or ""))

        now = datetime.now(timezone.utc).isoformat()
        merged_case_id = str(uuid.uuid4())
        merged_diagnosis_id = str(uuid.uuid4())
        disease_id = str(first_diagnosis.get("disease_id") or uuid.uuid4())
        diagnosis_entry: dict[str, Any] = {
            "diagnosis_id": merged_diagnosis_id,
            "disease_id": disease_id,
            "icd10": str(first_diagnosis.get("icd10") or ""),
            "histology": str(first_diagnosis.get("histology") or ""),
            "site": str(first_diagnosis.get("site") or ""),
            "biomarkers": first_diagnosis.get("biomarkers") if isinstance(first_diagnosis.get("biomarkers"), list) else [],
            "timeline": merged_timeline,
            "source_refs": [],
        }
        if isinstance(first_diagnosis.get("stage"), dict):
            diagnosis_entry["stage"] = first_diagnosis.get("stage")
        if isinstance(first_diagnosis.get("last_plan"), dict):
            diagnosis_entry["last_plan"] = first_diagnosis.get("last_plan")
        patient_entry: dict[str, Any] = {"sex": str(first_patient.get("sex") or "unknown")}
        birth_year = first_patient.get("birth_year")
        if isinstance(birth_year, int):
            patient_entry["birth_year"] = birth_year
        ecog = first_patient.get("ecog")
        if isinstance(ecog, int):
            patient_entry["ecog"] = ecog

        merged_case = {
            "schema_version": "1.0",
            "case_id": merged_case_id,
            "import_profile": "CUSTOM_TEMPLATE",
            "created_at": now,
            "updated_at": now,
            "data_mode": data_mode,
            "patient": patient_entry,
            "diagnoses": [diagnosis_entry],
            "attachments": [],
            "notes": "\n".join(merged_notes_parts)[:20000],
            "comorbidities": sorted(merged_comorbidities),
            "contraindications": sorted(merged_contraindications),
        }
        return merged_case

    def case_import_batch_file_base64(self, role: str, payload: dict[str, Any]) -> dict[str, Any]:
        ensure_role(role, {"admin", "clinician"})
        if not isinstance(payload, dict):
            raise ValidationError("Payload must be an object")
        files = payload.get("files")
        if not isinstance(files, list) or not files:
            raise ValidationError("files is required and must be non-empty array")

        batch_id = str(uuid.uuid4())
        runs: list[dict[str, Any]] = []
        successful_case_ids: list[str] = []
        for index, item in enumerate(files, start=1):
            if not isinstance(item, dict):
                runs.append(
                    {
                        "index": index,
                        "status": "ERROR",
                        "error": "file entry must be object",
                    }
                )
                continue
            filename = str(item.get("filename") or "").strip()
            try:
                import_result = self.case_import_file_base64(role=role, payload=item)
                status = str(import_result.get("status") or "ERROR")
                case_id = str(import_result.get("case_id") or "").strip()
                if status in {"SUCCESS", "PARTIAL_SUCCESS"} and case_id:
                    successful_case_ids.append(case_id)
                runs.append(
                    {
                        "index": index,
                        "filename": filename,
                        "status": status,
                        "import_run_id": str(import_result.get("import_run_id") or ""),
                        "case_id": case_id,
                        "missing_required_fields": import_result.get("missing_required_fields") or [],
                        "warnings": import_result.get("warnings") or [],
                        "errors": import_result.get("errors") or [],
                    }
                )
            except Exception as exc:  # noqa: BLE001
                runs.append(
                    {
                        "index": index,
                        "filename": filename,
                        "status": "ERROR",
                        "error": str(exc),
                    }
                )

        data_mode = self._normalize_case_data_mode(payload.get("data_mode") or "DEID")
        merged_case_payload = self._build_batch_merged_case(case_ids=successful_case_ids, data_mode=data_mode)
        merged_case_id = ""
        if isinstance(merged_case_payload, dict):
            merged_case_id = self.store.upsert_case(merged_case_payload)

        return {
            "schema_version": "1.0",
            "batch_id": batch_id,
            "total_files": len(files),
            "successful_imports": len(successful_case_ids),
            "failed_imports": len(files) - len(successful_case_ids),
            "merged_case_id": merged_case_id,
            "runs": runs,
            "status": "SUCCESS" if successful_case_ids else "FAILED",
        }

    def patient_analyze_file_base64(self, role: str, payload: dict[str, Any]) -> dict[str, Any]:
        ensure_role(role, {"admin", "clinician", "patient"})
        if not isinstance(payload, dict):
            raise ValidationError("Payload must be an object")

        elevated_role = role if role in {"admin", "clinician"} else "clinician"
        import_run = self.case_import_file_base64(role=elevated_role, payload=payload)
        case_id = str(import_run.get("case_id") or "").strip()
        if not case_id:
            raise ValidationError("Case import did not return case_id")

        query_type = str(payload.get("query_type") or "NEXT_STEPS").strip().upper()
        if query_type not in {"NEXT_STEPS", "CHECK_LAST_TREATMENT"}:
            query_type = "NEXT_STEPS"

        source_ids: list[str] = []
        sources_payload = payload.get("sources")
        if isinstance(sources_payload, dict):
            raw_source_ids = sources_payload.get("source_ids")
            if isinstance(raw_source_ids, list):
                source_ids = normalize_source_set_ids([str(item).strip() for item in raw_source_ids if str(item).strip()])
        if not source_ids and isinstance(payload.get("source_ids"), list):
            source_ids = normalize_source_set_ids([str(item).strip() for item in payload.get("source_ids", []) if str(item).strip()])
        if not source_ids:
            source_ids = list(DEFAULT_AUTO_SOURCE_IDS)

        source_mode = "AUTO" if len(source_ids) > 1 else "SINGLE"
        if isinstance(sources_payload, dict):
            declared_mode = str(sources_payload.get("mode") or "").strip().upper()
            if declared_mode in {"SINGLE", "AUTO"}:
                source_mode = declared_mode

        request_id = str(payload.get("request_id") or uuid.uuid4())
        language = str(payload.get("language") or "ru").strip().lower()
        if language not in {"ru", "en"}:
            language = "ru"
        query_mode = str(payload.get("query_mode") or "FULL_ANALYSIS").strip().upper()
        if query_mode not in {"FULL_ANALYSIS", "SOURCES_ONLY"}:
            query_mode = "FULL_ANALYSIS"
        as_of_date = str(
            payload.get("as_of_date") or payload.get("historical_reference_date") or ""
        ).strip()
        historical_assessment = bool(payload.get("historical_assessment")) or bool(as_of_date)

        analyze_payload = {
            "schema_version": "0.2",
            "request_id": request_id,
            "query_type": query_type,
            "query_mode": query_mode,
            "sources": {
                "mode": source_mode,
                "source_ids": source_ids,
            },
            "language": language,
            "as_of_date": as_of_date or None,
            "historical_assessment": historical_assessment,
            "case": {
                "case_id": case_id,
            },
        }
        analyze_response = self.analyze(
            payload=analyze_payload,
            role=elevated_role,
            client_id="patient-file-base64",
        )
        patient_explain = analyze_response.get("patient_explain")
        if not isinstance(patient_explain, dict):
            raise ValidationError("Patient explanation is unavailable for this request")
        patient_context = build_patient_context_from_analyze_response(analyze_response)

        response: dict[str, Any] = {
            "schema_version": "0.2",
            "request_id": str(analyze_response.get("request_id") or request_id),
            "case_id": case_id,
            "import_run_id": str(import_run.get("import_run_id") or ""),
            "patient_explain": patient_explain,
            "run_meta": analyze_response.get("run_meta") if isinstance(analyze_response.get("run_meta"), dict) else {},
            "insufficient_data": analyze_response.get("insufficient_data")
            if isinstance(analyze_response.get("insufficient_data"), dict)
            else {"status": False, "reason": "Sufficient evidence available."},
        }
        if isinstance(patient_context, dict) and patient_context:
            response["patient_context"] = patient_context
        return response

    def get_case(self, role: str, case_id: str) -> dict[str, Any]:
        ensure_role(role, {"admin", "clinician"})
        payload = self.store.get_case(case_id)
        if payload is None:
            raise NotFoundError(f"Case not found: {case_id}")
        return payload

    def get_case_import_run(self, role: str, import_run_id: str) -> dict[str, Any]:
        ensure_role(role, {"admin", "clinician"})
        payload = self.store.get_case_import_run(import_run_id)
        if payload is None:
            raise NotFoundError(f"Case import run not found: {import_run_id}")
        return payload

    def list_case_import_runs(self, role: str, limit: int = 20) -> list[dict[str, Any]]:
        ensure_role(role, {"admin", "clinician"})
        safe_limit = max(1, min(int(limit), 100))
        return self.store.list_case_import_runs(limit=safe_limit)

    def _record_admin_audit_event(
        self,
        *,
        role: str,
        action: str,
        doc_id: str | None = None,
        doc_version: str | None = None,
        details: dict[str, Any] | None = None,
    ) -> None:
        created_at = datetime.now(timezone.utc).isoformat()
        payload: dict[str, Any] = {
            "event_type": "admin_guideline_action",
            "role": role,
            "action": action,
            "doc_id": str(doc_id or ""),
            "doc_version": str(doc_version or ""),
            "created_at": created_at,
            "details": details or {},
        }
        try:
            self.store.save_admin_audit_event(payload, created_at=created_at)
        except Exception as exc:  # noqa: BLE001
            safe_log(self.logger, "warning", "Failed to persist admin audit event: %s", str(exc))

    def _transition_guideline_status(
        self,
        *,
        doc_id: str,
        doc_version: str,
        status: str,
        updated_at: str,
        metadata_patch: dict[str, Any] | None = None,
    ) -> str:
        normalized_status = str(status or "").strip().upper()
        version = self.store.get_guideline_version_by_doc(doc_id, doc_version)
        previous_status = str((version or {}).get("status") or "").strip().upper()
        metadata = version.get("metadata") if isinstance(version, dict) else {}
        history_raw = metadata.get("status_history") if isinstance(metadata, dict) else []
        history: list[str] = []
        if isinstance(history_raw, list):
            for item in history_raw:
                text = str(item or "").strip().upper()
                if text and (not history or history[-1] != text):
                    history.append(text)
        if not history and previous_status:
            history.append(previous_status)
        if normalized_status and (not history or history[-1] != normalized_status):
            history.append(normalized_status)

        patch = dict(metadata_patch or {})
        patch["status_history"] = history
        if previous_status and previous_status != normalized_status:
            patch["previous_status"] = previous_status
            patch["last_transition"] = f"{previous_status}->{normalized_status}"

        self.store.update_guideline_version_status(
            doc_id=doc_id,
            doc_version=doc_version,
            status=normalized_status,
            updated_at=updated_at,
            metadata_patch=patch,
        )
        return previous_status

    @staticmethod
    def _normalize_doc_kind(value: Any) -> str:
        token = str(value or "guideline").strip().lower()
        return token if token in {"guideline", "reference"} else "guideline"

    @staticmethod
    def _normalize_guideline_cancer_type(value: Any) -> tuple[str, str]:
        return apply_unknown_nosology_fallback(str(value or ""))

    @staticmethod
    def _extract_source_urls(metadata: dict[str, Any]) -> tuple[str, str, str]:
        source_page_url = str(metadata.get("source_page_url") or "").strip()
        source_pdf_url = str(metadata.get("source_pdf_url") or "").strip()
        legacy_source_url = str(metadata.get("source_url") or "").strip()
        if legacy_source_url and not source_page_url and not source_pdf_url:
            if ".pdf" in legacy_source_url.lower():
                source_pdf_url = legacy_source_url
            else:
                source_page_url = legacy_source_url
        source_url = resolve_primary_source_url(
            source_url=legacy_source_url,
            source_page_url=source_page_url,
            source_pdf_url=source_pdf_url,
        )
        return source_page_url, source_pdf_url, source_url

    def admin_upload(
        self,
        role: str,
        filename: str,
        content: bytes,
        metadata: dict[str, Any],
    ) -> dict[str, Any]:
        ensure_role(role, {"admin"})

        metadata = dict(metadata or {})
        required = {"doc_id", "doc_version", "source_set", "cancer_type", "language"}
        missing = [key for key in sorted(required) if not metadata.get(key)]
        if missing:
            raise ValidationError(f"Missing upload metadata fields: {', '.join(missing)}")

        source_set = normalize_source_set_id(str(metadata.get("source_set") or ""))
        metadata["source_set"] = source_set
        source_page_url, source_pdf_url, source_url = self._extract_source_urls(metadata)
        metadata["source_page_url"] = source_page_url
        metadata["source_pdf_url"] = source_pdf_url
        metadata["source_url"] = source_url
        doc_kind = self._normalize_doc_kind(metadata.get("doc_kind"))
        metadata["doc_kind"] = doc_kind
        allow_missing_source_url = bool(metadata.get("allow_missing_source_url"))
        if source_set in OFFICIAL_SOURCE_RULES and doc_kind == "guideline" and not source_url and not allow_missing_source_url:
            raise ValidationError("OFFICIAL_SOURCE_URL_REQUIRED: source_url is required for official guideline upload.")

        nosology_inference = ""
        if doc_kind == "guideline":
            normalized_cancer_type, nosology_inference = self._normalize_guideline_cancer_type(metadata.get("cancer_type"))
            metadata["cancer_type"] = normalized_cancer_type
            if not nosology_inference:
                nosology_inference = str(metadata.get("nosology_inference") or "").strip()

        sha = hashlib.sha256(content).hexdigest()
        existing_docs = self.store.list_docs()
        duplicate = None
        if source_set in OFFICIAL_SOURCE_RULES:
            duplicate = next(
                (
                    item
                    for item in existing_docs
                    if str(item.get("sha256") or "") == sha
                    and (
                        str(item.get("doc_id") or "") != str(metadata["doc_id"])
                        or str(item.get("doc_version") or "") != str(metadata["doc_version"])
                    )
                ),
                None,
            )
        if isinstance(duplicate, dict):
            alias_doc_id = str(duplicate.get("doc_id") or "")
            alias_doc_version = str(duplicate.get("doc_version") or "")
            self._record_admin_audit_event(
                role=role,
                action="upload_duplicate_skipped",
                doc_id=str(metadata["doc_id"]),
                doc_version=str(metadata["doc_version"]),
                details={
                    "sha256": sha,
                    "alias_doc_id": alias_doc_id,
                    "alias_doc_version": alias_doc_version,
                    "requested_filename": filename,
                    "source_url": source_url,
                    "source_page_url": source_page_url,
                    "source_pdf_url": source_pdf_url,
                },
            )
            return {
                "status": "duplicate_skipped",
                "doc_id": alias_doc_id,
                "doc_version": alias_doc_version,
                "sha256": sha,
                "source_url": source_url,
                "source_page_url": source_page_url,
                "source_pdf_url": source_pdf_url,
                "doc_kind": doc_kind,
                "alias_for": {
                    "doc_id": alias_doc_id,
                    "doc_version": alias_doc_version,
                },
            }

        doc_folder = (
            self.settings.docs_dir
            / metadata["source_set"]
            / metadata["doc_id"]
            / metadata["doc_version"]
        )
        doc_folder.mkdir(parents=True, exist_ok=True)
        target = doc_folder / filename
        target.write_bytes(content)

        sha = file_sha256(target)
        uploaded_at = datetime.now(timezone.utc).isoformat()

        record = DocRecord(
            doc_id=str(metadata["doc_id"]),
            doc_version=str(metadata["doc_version"]),
            source_set=str(metadata["source_set"]),
            cancer_type=str(metadata["cancer_type"]),
            language=str(metadata["language"]),
            file_path=str(target),
            sha256=sha,
            uploaded_at=uploaded_at,
        )
        self.store.upsert_doc(record)
        metadata_patch: dict[str, Any] = {
            "source_url": source_url,
            "source_page_url": source_page_url,
            "source_pdf_url": source_pdf_url,
            "doc_kind": doc_kind,
        }
        if nosology_inference:
            metadata_patch["nosology_inference"] = nosology_inference
        self._transition_guideline_status(
            doc_id=record.doc_id,
            doc_version=record.doc_version,
            status="NEW",
            updated_at=uploaded_at,
            metadata_patch=metadata_patch,
        )
        routes_created = (
            self._upsert_doc_nosology_routes(metadata=metadata, uploaded_at=uploaded_at)
            if doc_kind == "guideline"
            else 0
        )
        self._clear_routing_cache()

        self._kb_version = compute_kb_version(self.store.list_docs())
        self._record_admin_audit_event(
            role=role,
            action="upload",
            doc_id=record.doc_id,
            doc_version=record.doc_version,
            details={
                "sha256": record.sha256,
                "source_url": source_url,
                "source_page_url": source_page_url,
                "source_pdf_url": source_pdf_url,
                "doc_kind": doc_kind,
                "routes_created": routes_created,
            },
        )
        return {
            "status": "uploaded",
            "doc_id": record.doc_id,
            "doc_version": record.doc_version,
            "sha256": record.sha256,
            "uploaded_at": record.uploaded_at,
            "source_url": source_url,
            "source_page_url": source_page_url,
            "source_pdf_url": source_pdf_url,
            "doc_kind": doc_kind,
            "pipeline_status": "NEW",
            "routes_created": routes_created,
        }

    def admin_docs(self, role: str, *, valid_only: bool = False, kind: str = "guideline") -> dict[str, Any]:
        ensure_role(role, {"admin"})
        normalized_kind = str(kind or "guideline").strip().lower()
        if normalized_kind not in {"guideline", "reference", "all"}:
            raise ValidationError("kind must be one of guideline|reference|all")
        docs = self.store.list_docs()
        enriched_docs: list[dict[str, Any]] = []
        for doc in docs:
            version = self.store.get_guideline_version_by_doc(str(doc.get("doc_id")), str(doc.get("doc_version")))
            metadata = version.get("metadata") if isinstance(version, dict) else {}
            metadata = metadata if isinstance(metadata, dict) else {}
            enriched = dict(doc)
            enriched["source_set"] = normalize_source_set_id(str(enriched.get("source_set") or ""))
            enriched["status"] = str((version or {}).get("status") or "NEW")
            enriched["chunk_count"] = self.store.count_doc_chunks(str(doc.get("doc_id")), str(doc.get("doc_version")))
            source_page_url, source_pdf_url, source_url = self._extract_source_urls(metadata)
            enriched["source_page_url"] = source_page_url
            enriched["source_pdf_url"] = source_pdf_url
            enriched["source_url"] = source_url
            enriched["doc_kind"] = self._normalize_doc_kind(metadata.get("doc_kind"))
            enriched["last_error_code"] = str(metadata.get("last_error_code") or "")
            enriched["nosology_inference"] = str(metadata.get("nosology_inference") or "")
            enriched["updated_at"] = str((version or {}).get("updated_at") or doc.get("uploaded_at") or "")
            if enriched["doc_kind"] == "guideline":
                inferred_cancer_type = infer_cancer_type_for_guideline(
                    doc_id=str(enriched.get("doc_id") or ""),
                    source_url=str(enriched.get("source_url") or ""),
                    title=str(enriched.get("doc_id") or ""),
                    fallback=str(enriched.get("cancer_type") or ""),
                )
                normalized_cancer_type, nosology_inference = self._normalize_guideline_cancer_type(inferred_cancer_type)
                enriched["cancer_type"] = normalized_cancer_type
                if nosology_inference and not enriched["nosology_inference"]:
                    enriched["nosology_inference"] = nosology_inference
            validity = self._doc_release_validity(
                doc_id=str(enriched.get("doc_id") or ""),
                source_set=str(enriched.get("source_set") or ""),
                source_url=str(enriched.get("source_url") or ""),
                status=str(enriched.get("status") or ""),
                cancer_type=str(enriched.get("cancer_type") or ""),
            )
            enriched.update(validity)
            if normalized_kind != "all" and enriched["doc_kind"] != normalized_kind:
                continue
            enriched_docs.append(enriched)
        if valid_only:
            enriched_docs = [doc for doc in enriched_docs if bool(doc.get("is_valid"))]
        return {
            "docs": enriched_docs,
            "kb_version": self._kb_version,
            "governance": self._governance_snapshot(),
        }

    def admin_drug_dictionary_load(self, role: str, payload: dict[str, Any]) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        if not isinstance(payload, dict):
            raise ValidationError("Payload must be an object")

        content_base64 = str(payload.get("content_base64") or "").strip()
        content_text = str(payload.get("content") or "")
        filename = str(payload.get("filename") or "drug_dictionary.json").strip() or "drug_dictionary.json"
        if content_base64:
            try:
                raw_text = base64.b64decode(content_base64).decode("utf-8", errors="replace")
            except Exception as exc:  # noqa: BLE001
                raise ValidationError("Invalid content_base64 for drug dictionary load") from exc
        elif content_text.strip():
            raw_text = content_text
        else:
            raise ValidationError("Either content_base64 or content must be provided for drug dictionary load")

        bundle = load_drug_dictionary_bundle_from_text(raw_text)
        loaded = self._apply_drug_dictionary_bundle(bundle=bundle, origin=filename, loaded_by=role)
        self._record_admin_audit_event(
            role=role,
            action="drug_dictionary_load",
            details={
                "version": loaded["version"],
                "sha256": loaded["sha256"],
                "entries": loaded["entries"],
                "regimens": loaded["regimens"],
                "filename": filename,
            },
        )
        return {
            "status": "ok",
            "version": loaded["version"],
            "sha256": loaded["sha256"],
            "entries_loaded": loaded["entries"],
            "regimens_loaded": loaded["regimens"],
            "loaded_at": loaded["loaded_at"],
            "schema": bundle.schema,
            "notes": bundle.notes,
        }

    def admin_drug_safety_cache_warmup(self, role: str, payload: dict[str, Any] | None = None) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        body = payload if isinstance(payload, dict) else {}
        inns_raw = body.get("inns")
        inns: list[str] = []
        if isinstance(inns_raw, list):
            inns = [str(item).strip().lower() for item in inns_raw if str(item).strip()]
        elif isinstance(inns_raw, str):
            inns = [item.strip().lower() for item in inns_raw.split(",") if item.strip()]
        inns = sorted(set(inns))
        if not inns:
            default_limit = max(1, min(int(body.get("limit") or 60), 1000))
            inns = [
                str(item.get("inn") or "").strip().lower()
                for item in self._drug_dictionary_entries[:default_limit]
                if str(item.get("inn") or "").strip()
            ]
            inns = sorted(set(inns))

        result = self.drug_safety_provider.warmup_cache(inns)
        cache_state = self.admin_drug_safety_cache(role=role, limit=int(body.get("inspect_limit") or 200))
        self._record_admin_audit_event(
            role=role,
            action="drug_safety_cache_warmup",
            details={
                "requested": len(inns),
                "status": str(result.get("status") or "unknown"),
                "profiles": int(result.get("profiles") or 0),
                "warnings": len(result.get("warnings") if isinstance(result.get("warnings"), list) else []),
            },
        )
        return {
            **result,
            "requested_inn": inns,
            "cache_summary": cache_state.get("summary") if isinstance(cache_state, dict) else {},
        }

    def admin_drug_safety_cache(self, role: str, limit: int = 200) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        safe_limit = max(1, min(int(limit), 5000))
        items = self.store.list_drug_safety_cache(limit=safe_limit)
        now = datetime.now(timezone.utc)
        fresh_count = 0
        expired_count = 0
        status_counts: dict[str, int] = {}
        for item in items:
            status_key = str(item.get("status") or "unknown").strip().lower() or "unknown"
            status_counts[status_key] = status_counts.get(status_key, 0) + 1
            expires_at = str(item.get("expires_at") or "").strip()
            try:
                expires_dt = datetime.fromisoformat(expires_at.replace("Z", "+00:00")) if expires_at else None
            except ValueError:
                expires_dt = None
            if expires_dt is not None and expires_dt >= now:
                fresh_count += 1
            else:
                expired_count += 1

        dictionary_total = len(self._drug_dictionary_entries)
        unique_cached = len({str(item.get("inn") or "").strip().lower() for item in items if str(item.get("inn") or "").strip()})
        coverage_ratio = round((unique_cached / dictionary_total), 4) if dictionary_total > 0 else 0.0
        return {
            "items": items,
            "summary": {
                "dictionary_entries_total": dictionary_total,
                "cache_entries_total": len(items),
                "unique_cached_inn": unique_cached,
                "fresh_entries": fresh_count,
                "expired_entries": expired_count,
                "coverage_ratio": coverage_ratio,
                "status_counts": status_counts,
                "cache_ttl_hours": int(getattr(self.settings, "drug_safety_cache_ttl_hours", 24 * 14)),
            },
        }

    def admin_docs_cleanup_invalid(
        self,
        role: str,
        *,
        dry_run: bool = True,
        apply: bool = False,
        reason_allowlist: list[str] | None = None,
    ) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        safe_default_allowlist = {
            "non_official_source_url",
            "missing_source_url",
            "demo_document_excluded",
            "comparative_only_source_set",
        }
        normalized_allowlist = (
            {
                str(item).strip().lower()
                for item in reason_allowlist
                if str(item).strip()
            }
            if isinstance(reason_allowlist, list)
            else set(safe_default_allowlist)
        )
        docs_payload = self.admin_docs(role=role, valid_only=False, kind="all")
        docs = docs_payload.get("docs") if isinstance(docs_payload.get("docs"), list) else []
        candidates: list[dict[str, Any]] = []
        for item in docs:
            if not isinstance(item, dict):
                continue
            if bool(item.get("is_valid")):
                continue
            validity_reason = str(item.get("validity_reason") or "unknown").strip().lower()
            if normalized_allowlist and validity_reason not in normalized_allowlist:
                continue
            candidates.append(
                {
                    "doc_id": str(item.get("doc_id") or ""),
                    "doc_version": str(item.get("doc_version") or ""),
                    "source_set": str(item.get("source_set") or ""),
                    "status": str(item.get("status") or ""),
                    "validity_reason": validity_reason or "unknown",
                }
            )

        if dry_run and not apply:
            self._record_admin_audit_event(
                role=role,
                action="cleanup_invalid_dry_run",
                details={
                    "candidates": len(candidates),
                    "reason_allowlist": sorted(normalized_allowlist),
                },
            )
            return {
                "mode": "dry_run",
                "reason_allowlist": sorted(normalized_allowlist),
                "candidates": candidates,
                "deleted_count": 0,
                "error_count": 0,
                "errors": [],
            }

        deleted: list[dict[str, Any]] = []
        errors: list[dict[str, str]] = []
        for candidate in candidates:
            doc_id = str(candidate.get("doc_id") or "")
            doc_version = str(candidate.get("doc_version") or "")
            if not doc_id or not doc_version:
                continue
            try:
                stats = self._delete_doc_everywhere(doc_id=doc_id, doc_version=doc_version)
                deleted.append(
                    {
                        "doc_id": doc_id,
                        "doc_version": doc_version,
                        "validity_reason": str(candidate.get("validity_reason") or "unknown"),
                        "delete_stats": stats,
                    }
                )
            except Exception as exc:  # noqa: BLE001
                errors.append(
                    {
                        "doc_id": doc_id,
                        "doc_version": doc_version,
                        "error": str(exc),
                    }
                )

        self._record_admin_audit_event(
            role=role,
            action="cleanup_invalid_apply",
            details={
                "candidates": len(candidates),
                "deleted_count": len(deleted),
                "error_count": len(errors),
                "reason_allowlist": sorted(normalized_allowlist),
            },
        )

        return {
            "mode": "apply" if apply else "dry_run",
            "reason_allowlist": sorted(normalized_allowlist),
            "candidates": candidates,
            "deleted": deleted,
            "deleted_count": len(deleted),
            "error_count": len(errors),
            "errors": errors,
        }

    def admin_doc_rechunk(self, role: str, doc_id: str, doc_version: str) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        doc = self.store.get_doc(doc_id, doc_version)
        if not doc:
            raise NotFoundError(f"Document not found: {doc_id}:{doc_version}")
        version = self.store.get_guideline_version_by_doc(doc_id, doc_version)
        version_metadata = version.get("metadata") if isinstance(version, dict) else {}
        version_metadata = version_metadata if isinstance(version_metadata, dict) else {}
        chunk_metadata = dict(doc)
        chunk_metadata["source_url"] = resolve_primary_source_url(
            source_url=str(version_metadata.get("source_url") or ""),
            source_page_url=str(version_metadata.get("source_page_url") or ""),
            source_pdf_url=str(version_metadata.get("source_pdf_url") or ""),
        )
        chunk_metadata["doc_kind"] = self._normalize_doc_kind(version_metadata.get("doc_kind"))
        chunks = extract_pdf_chunks(
            Path(str(doc["file_path"])),
            metadata=chunk_metadata,
            structural_chunker_enabled=bool(self.settings.oncoai_structural_chunker_enabled),
        )
        self.store.replace_doc_chunks(doc_id=doc_id, doc_version=doc_version, chunks=[
            {
                **chunk,
                "vector_json": json.dumps(self.embedder.embed(str(chunk.get("text") or "")), ensure_ascii=False),
            }
            for chunk in chunks
        ])
        now = datetime.now(timezone.utc).isoformat()
        self._transition_guideline_status(
            doc_id=doc_id,
            doc_version=doc_version,
            status="CHUNKED",
            updated_at=now,
            metadata_patch={"chunk_count": len(chunks)},
        )
        previous_status = self._transition_guideline_status(
            doc_id=doc_id,
            doc_version=doc_version,
            status="PENDING_APPROVAL",
            updated_at=now,
            metadata_patch={"chunk_count": len(chunks)},
        )
        self._record_admin_audit_event(
            role=role,
            action="rechunk",
            doc_id=doc_id,
            doc_version=doc_version,
            details={
                "chunk_count": len(chunks),
                "intermediate_status": "CHUNKED",
                "previous_status": previous_status,
            },
        )
        return {
            "status": "PENDING_APPROVAL",
            "doc_id": doc_id,
            "doc_version": doc_version,
            "chunk_count": len(chunks),
            "intermediate_status": "CHUNKED",
            "previous_status": previous_status,
            "updated_at": now,
        }

    def admin_doc_approve(self, role: str, doc_id: str, doc_version: str) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        version = self.store.get_guideline_version_by_doc(doc_id, doc_version)
        status = str((version or {}).get("status") or "").strip().upper()
        if status not in {"CHUNKED", "PENDING_APPROVAL"}:
            raise ValidationError("APPROVE_REQUIRES_RECHUNK: document must be CHUNKED/PENDING_APPROVAL before approve.")
        now = datetime.now(timezone.utc).isoformat()
        previous_status = self._transition_guideline_status(
            doc_id=doc_id,
            doc_version=doc_version,
            status="APPROVED",
            updated_at=now,
        )
        self._record_admin_audit_event(role=role, action="approve", doc_id=doc_id, doc_version=doc_version)
        return {
            "status": "APPROVED",
            "doc_id": doc_id,
            "doc_version": doc_version,
            "previous_status": previous_status,
            "updated_at": now,
        }

    def admin_doc_verify_index(self, role: str, doc_id: str, doc_version: str) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        doc = self.store.get_doc(doc_id, doc_version)
        if not doc:
            raise NotFoundError(f"Document not found: {doc_id}:{doc_version}")

        sqlite_chunk_count = self.store.count_doc_chunks(doc_id=doc_id, doc_version=doc_version)
        vector_backend = str(self.settings.vector_backend or "local").strip().lower()
        qdrant_point_count = 0
        try:
            if hasattr(self.index, "count_doc_points"):
                qdrant_point_count = int(self.index.count_doc_points(doc_id=doc_id, doc_version=doc_version))  # type: ignore[attr-defined]
            else:
                qdrant_point_count = int(sqlite_chunk_count)
        except Exception:  # noqa: BLE001
            qdrant_point_count = 0

        status = "ok"
        if sqlite_chunk_count <= 0:
            status = "failed"
        elif vector_backend == "qdrant" and qdrant_point_count <= 0:
            status = "failed"

        return {
            "doc_id": doc_id,
            "doc_version": doc_version,
            "sqlite_chunk_count": int(sqlite_chunk_count),
            "qdrant_point_count": int(qdrant_point_count),
            "vector_backend": vector_backend,
            "status": status,
        }

    def _refresh_icd10_reference_from_doc(self, *, doc_id: str, doc_version: str, updated_at: str) -> int:
        chunks = self.store.list_chunks(filters={"doc_id": doc_id, "doc_version": doc_version})
        parsed_entries = parse_icd10_reference_entries_from_chunks(chunks)
        if not parsed_entries:
            return 0
        return self.store.replace_icd10_reference_entries(
            source_doc_id=doc_id,
            source_doc_version=doc_version,
            rows=parsed_entries,
            updated_at=updated_at,
        )

    def admin_doc_reject(self, role: str, doc_id: str, doc_version: str, reason: str | None = None) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        now = datetime.now(timezone.utc).isoformat()
        previous_status = self._transition_guideline_status(
            doc_id=doc_id,
            doc_version=doc_version,
            status="REJECTED",
            updated_at=now,
            metadata_patch={"last_error_code": str(reason or "rejected_by_admin")},
        )
        self._record_admin_audit_event(
            role=role,
            action="reject",
            doc_id=doc_id,
            doc_version=doc_version,
            details={"reason": str(reason or "rejected_by_admin")},
        )
        return {
            "status": "REJECTED",
            "doc_id": doc_id,
            "doc_version": doc_version,
            "previous_status": previous_status,
            "updated_at": now,
        }

    def admin_doc_index(self, role: str, doc_id: str, doc_version: str) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        version = self.store.get_guideline_version_by_doc(doc_id, doc_version)
        status = str((version or {}).get("status") or "")
        if status != "APPROVED":
            raise ValidationError("INDEX_REQUIRES_APPROVAL: document must be APPROVED before index.")
        version_metadata = version.get("metadata") if isinstance(version, dict) else {}
        version_metadata = version_metadata if isinstance(version_metadata, dict) else {}
        source_url = resolve_primary_source_url(
            source_url=str(version_metadata.get("source_url") or ""),
            source_page_url=str(version_metadata.get("source_page_url") or ""),
            source_pdf_url=str(version_metadata.get("source_pdf_url") or ""),
        )
        doc_kind = self._normalize_doc_kind(version_metadata.get("doc_kind"))
        doc = self.store.get_doc(doc_id, doc_version)
        if not doc:
            raise NotFoundError(f"Document not found: {doc_id}:{doc_version}")
        chunk_metadata = dict(doc)
        chunk_metadata["source_url"] = source_url
        chunk_metadata["doc_kind"] = doc_kind
        chunks = extract_pdf_chunks(
            Path(str(doc["file_path"])),
            metadata=chunk_metadata,
            structural_chunker_enabled=bool(self.settings.oncoai_structural_chunker_enabled),
        )
        self.index.replace_doc_chunks(doc_id=doc_id, doc_version=doc_version, chunks=chunks)
        verify = self.admin_doc_verify_index(role=role, doc_id=doc_id, doc_version=doc_version)
        if str(verify.get("status") or "").lower() != "ok":
            raise ValidationError(
                "INDEX_VERIFY_FAILED: chunk/vector counts are insufficient "
                f"(sqlite={verify.get('sqlite_chunk_count')}, qdrant={verify.get('qdrant_point_count')})"
            )

        now = datetime.now(timezone.utc).isoformat()
        icd10_reference_rows = 0
        if doc_kind == "reference" or str(doc.get("cancer_type") or "").strip().lower() == "reference_icd10":
            icd10_reference_rows = self._refresh_icd10_reference_from_doc(
                doc_id=doc_id,
                doc_version=doc_version,
                updated_at=now,
            )
        previous_status = self._transition_guideline_status(
            doc_id=doc_id,
            doc_version=doc_version,
            status="INDEXED",
            updated_at=now,
            metadata_patch={
                "chunk_count": len(chunks),
                "doc_kind": doc_kind,
                "source_url": source_url,
                "source_page_url": str(version_metadata.get("source_page_url") or ""),
                "source_pdf_url": str(version_metadata.get("source_pdf_url") or ""),
                "sqlite_chunk_count": int(verify.get("sqlite_chunk_count") or 0),
                "qdrant_point_count": int(verify.get("qdrant_point_count") or 0),
                "vector_backend": str(verify.get("vector_backend") or self.settings.vector_backend),
                "icd10_reference_rows": int(icd10_reference_rows),
            },
        )
        self._kb_version = compute_kb_version(self.store.list_docs())
        self._record_admin_audit_event(
            role=role,
            action="index",
            doc_id=doc_id,
            doc_version=doc_version,
            details={
                "chunk_count": len(chunks),
                "kb_version": self._kb_version,
                "verify_index": verify,
                "doc_kind": doc_kind,
                "icd10_reference_rows": int(icd10_reference_rows),
            },
        )
        return {
            "status": "INDEXED",
            "doc_id": doc_id,
            "doc_version": doc_version,
            "chunk_count": len(chunks),
            "kb_version": self._kb_version,
            "previous_status": previous_status,
            "updated_at": now,
            "verify_index": verify,
            "doc_kind": doc_kind,
            "icd10_reference_rows": int(icd10_reference_rows),
        }

    def admin_sync_russco(self, role: str) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        if not bool(self.settings.oncoai_guideline_sync_enabled):
            raise ValidationError("GUIDELINE_SYNC_DISABLED: ONCOAI_GUIDELINE_SYNC_ENABLED=false")

        documents = discover_russco_2025_documents()
        synced: list[dict[str, Any]] = []
        errors: list[dict[str, str]] = []
        for item in documents:
            try:
                payload = download_russco_pdf(item.url)
                source_set = normalize_source_set_id(item.source_set)
                hints = resolve_official_doc_hints(
                    source_set=source_set,
                    doc_id=item.doc_id,
                    source_url=item.url,
                    fallback_cancer_type=item.cancer_type,
                    fallback_doc_kind="guideline",
                )
                inferred_cancer_type = infer_cancer_type_for_guideline(
                    doc_id=item.doc_id,
                    source_url=item.url,
                    title=item.filename,
                    fallback=str(hints.get("cancer_type") or item.cancer_type),
                )
                cancer_type, nosology_inference = self._normalize_guideline_cancer_type(inferred_cancer_type)
                uploaded = self.admin_upload(
                    role=role,
                    filename=item.filename,
                    content=payload,
                    metadata={
                        "doc_id": item.doc_id,
                        "doc_version": item.doc_version,
                        "source_set": source_set,
                        "cancer_type": cancer_type,
                        "language": item.language,
                        "source_page_url": item.url,
                        "source_pdf_url": item.url,
                        "doc_kind": str(hints.get("doc_kind") or "guideline"),
                        "icd10_prefixes": list(hints.get("icd10_prefixes") or []),
                        "nosology_inference": nosology_inference,
                    },
                )
                rechunked = self.admin_doc_rechunk(role=role, doc_id=item.doc_id, doc_version=item.doc_version)
                validity = self._doc_release_validity(
                    doc_id=item.doc_id,
                    source_set=source_set,
                    source_url=item.url,
                    status=str(rechunked.get("status") or ""),
                    cancer_type=cancer_type,
                )
                synced.append(
                    {
                        "doc_id": item.doc_id,
                        "doc_version": item.doc_version,
                        "source_url": item.url,
                        "source_page_url": item.url,
                        "source_pdf_url": item.url,
                        "cancer_type": cancer_type,
                        "nosology_inference": nosology_inference,
                        "sha256": uploaded.get("sha256"),
                        "status": rechunked.get("status"),
                        "is_valid": validity.get("is_valid"),
                        "validity_reason": validity.get("validity_reason"),
                        "official_source": validity.get("official_source"),
                    }
                )
            except Exception as exc:  # noqa: BLE001
                errors.append({"doc_id": item.doc_id, "error": str(exc)})
        validation_summary = self._build_sync_validation_summary(synced)
        payload = {
            "source": "russco",
            "status": "ok" if not errors else "partial",
            "synced": synced,
            "errors": errors,
            "count": len(synced),
            "validation_summary": validation_summary,
        }
        self._record_admin_audit_event(
            role=role,
            action="sync_russco",
            details={
                "status": payload["status"],
                "count": payload["count"],
                "errors": len(errors),
                "validation_summary": validation_summary,
            },
        )
        return payload

    def admin_sync_minzdrav(self, role: str) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        if not bool(self.settings.oncoai_guideline_sync_enabled):
            raise ValidationError("GUIDELINE_SYNC_DISABLED: ONCOAI_GUIDELINE_SYNC_ENABLED=false")

        synced: list[dict[str, Any]] = []
        errors: list[dict[str, str]] = []
        for item in KNOWN_MINZDRAV_PDFS:
            try:
                payload, resolved_download_url = download_minzdrav_pdf_with_url(item)
                source_set = normalize_source_set_id(item.source_set)
                source_page_url = str(item.source_page_url or "").strip()
                source_pdf_url = str(resolved_download_url or resolve_minzdrav_pdf_from_page(item) or item.source_pdf_url or "").strip()
                effective_source_url = resolve_primary_source_url(
                    source_page_url=source_page_url,
                    source_pdf_url=source_pdf_url,
                )
                hints = resolve_official_doc_hints(
                    source_set=source_set,
                    doc_id=item.doc_id,
                    source_url=effective_source_url,
                    fallback_cancer_type=item.cancer_type,
                    fallback_doc_kind="guideline",
                )
                inferred_cancer_type = infer_cancer_type_for_guideline(
                    doc_id=item.doc_id,
                    source_url=effective_source_url or source_page_url,
                    title=item.filename,
                    fallback=str(hints.get("cancer_type") or item.cancer_type),
                )
                cancer_type, nosology_inference = self._normalize_guideline_cancer_type(inferred_cancer_type)
                uploaded = self.admin_upload(
                    role=role,
                    filename=item.filename,
                    content=payload,
                    metadata={
                        "doc_id": item.doc_id,
                        "doc_version": item.doc_version,
                        "source_set": source_set,
                        "cancer_type": cancer_type,
                        "language": item.language,
                        "source_page_url": source_page_url,
                        "source_pdf_url": source_pdf_url,
                        "doc_kind": str(hints.get("doc_kind") or "guideline"),
                        "icd10_prefixes": list(hints.get("icd10_prefixes") or []),
                        "nosology_inference": nosology_inference,
                    },
                )
                rechunked = self.admin_doc_rechunk(role=role, doc_id=item.doc_id, doc_version=item.doc_version)
                validity = self._doc_release_validity(
                    doc_id=item.doc_id,
                    source_set=source_set,
                    source_url=effective_source_url,
                    status=str(rechunked.get("status") or ""),
                    cancer_type=cancer_type,
                )
                synced.append(
                    {
                        "doc_id": item.doc_id,
                        "doc_version": item.doc_version,
                        "source_url": effective_source_url,
                        "source_page_url": source_page_url,
                        "source_pdf_url": source_pdf_url,
                        "cancer_type": cancer_type,
                        "nosology_inference": nosology_inference,
                        "sha256": uploaded.get("sha256"),
                        "status": rechunked.get("status"),
                        "is_valid": validity.get("is_valid"),
                        "validity_reason": validity.get("validity_reason"),
                        "official_source": validity.get("official_source"),
                    }
                )
            except Exception as exc:  # noqa: BLE001
                message = str(exc)
                error_code = "MINZDRAV_SYNC_ERROR"
                if "MINZDRAV_NON_PDF_CONTENT" in message:
                    error_code = "MINZDRAV_NON_PDF_CONTENT"
                elif "MINZDRAV_DOWNLOAD_FAILED" in message:
                    error_code = "MINZDRAV_DOWNLOAD_FAILED"
                errors.append({"doc_id": item.doc_id, "error": message, "error_code": error_code})
        payload = {
            "source": "minzdrav",
            "status": "ok" if not errors else "partial",
            "synced": synced,
            "errors": errors,
            "count": len(synced),
            "validation_summary": self._build_sync_validation_summary(synced),
        }
        self._record_admin_audit_event(
            role=role,
            action="sync_minzdrav",
            details={
                "status": payload["status"],
                "count": payload["count"],
                "errors": len(errors),
                "validation_summary": payload.get("validation_summary"),
            },
        )
        return payload

    def admin_routing_routes(self, role: str, language: str | None = None) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        normalized_language = str(language or "").strip().lower()
        routes = self.store.list_nosology_routes(
            language=normalized_language if normalized_language in {"ru", "en"} else None,
            active_only=False,
        )
        return {
            "routes": routes,
            "count": len(routes),
        }

    def admin_routing_rebuild(self, role: str) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        docs = self.store.list_docs()
        active_pairs = {
            (normalize_source_set_id(str(doc.get("source_set") or "")), str(doc.get("doc_id") or "").strip())
            for doc in docs
            if str(doc.get("source_set") or "").strip() and str(doc.get("doc_id") or "").strip()
        }
        sync_result = self.store.sync_nosology_routes_active_docs(active_pairs=active_pairs)
        active_routes = self.store.list_nosology_routes(active_only=True)
        active_route_pairs = {
            (normalize_source_set_id(str(route.get("source_id") or "")), str(route.get("doc_id") or "").strip())
            for route in active_routes
            if str(route.get("source_id") or "").strip() and str(route.get("doc_id") or "").strip()
        }

        generated_routes = 0
        now = datetime.now(timezone.utc).isoformat()
        for doc in docs:
            pair = (normalize_source_set_id(str(doc.get("source_set") or "")), str(doc.get("doc_id") or "").strip())
            if pair in active_route_pairs:
                continue
            generated_routes += self._upsert_doc_nosology_routes(metadata=dict(doc), uploaded_at=now)

        self._clear_routing_cache()
        active_after = self.store.list_nosology_routes(active_only=True)
        return {
            "status": "ok",
            "routes_total": len(self.store.list_nosology_routes(active_only=False)),
            "active_routes": len(active_after),
            "generated_routes": generated_routes,
            "synced_pairs": len(active_pairs),
            "sync": sync_result,
            "updated_at": now,
        }

    def admin_reindex(self, role: str) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        docs = self.store.list_docs()
        total_docs = len(docs)
        processed_docs = 0
        started_at = datetime.now(timezone.utc).isoformat()
        job_id = self.store.create_reindex_job(
            started_at=started_at,
            total_docs=total_docs,
        )
        ingestion_run_id = self.store.create_ingestion_run(
            started_at=started_at,
            total_docs=total_docs,
            metadata={"kind": "admin_reindex"},
        )
        self.store.attach_ingestion_run_to_reindex(job_id=job_id, run_id=ingestion_run_id)
        try:
            if not docs:
                raise ValidationError("No uploaded documents found. Upload at least one guideline before reindex.")
            for idx, doc in enumerate(docs, start=1):
                path = Path(doc["file_path"])
                chunks = extract_pdf_chunks(
                    path,
                    metadata=doc,
                    structural_chunker_enabled=bool(self.settings.oncoai_structural_chunker_enabled),
                )
                self.index.replace_doc_chunks(doc_id=doc["doc_id"], doc_version=doc["doc_version"], chunks=chunks)
                processed_docs = idx
                self.store.update_reindex_job_progress(
                    job_id=job_id,
                    processed_docs=processed_docs,
                    total_docs=total_docs,
                )
                self.store.update_ingestion_run_progress(
                    run_id=ingestion_run_id,
                    processed_docs=processed_docs,
                    total_docs=total_docs,
                )

            self._kb_version = compute_kb_version(docs)
            finished_at = datetime.now(timezone.utc).isoformat()
            self.store.finish_reindex_job(
                job_id,
                finished_at=finished_at,
                processed_docs=processed_docs,
                total_docs=total_docs,
            )
            self.store.finish_ingestion_run(
                ingestion_run_id,
                finished_at=finished_at,
                processed_docs=processed_docs,
                total_docs=total_docs,
                kb_version=self._kb_version,
            )
        except Exception as exc:  # noqa: BLE001
            error_code = "VALIDATION_ERROR" if isinstance(exc, ValidationError) else "REINDEX_ERROR"
            finished_at = datetime.now(timezone.utc).isoformat()
            self.store.finish_reindex_job(
                job_id,
                finished_at=finished_at,
                error_message=str(exc),
                error_code=error_code,
                processed_docs=processed_docs,
                total_docs=total_docs,
            )
            self.store.finish_ingestion_run(
                ingestion_run_id,
                finished_at=finished_at,
                error_message=str(exc),
                error_code=error_code,
                processed_docs=processed_docs,
                total_docs=total_docs,
                kb_version=self._kb_version,
            )
            raise

        routing_rebuild = self.admin_routing_rebuild(role=role)
        self._clear_routing_cache()
        return {
            "job_id": job_id,
            "ingestion_run_id": ingestion_run_id,
            "status": "done",
            "kb_version": self._kb_version,
            "processed_docs": processed_docs,
            "total_docs": total_docs,
            "last_error_code": None,
            "routing_rebuild": routing_rebuild,
        }

    def admin_reindex_status(self, role: str, job_id: str) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        job = self.store.get_reindex_job(job_id)
        if not job:
            raise NotFoundError(f"Reindex job not found: {job_id}")
        ingestion_run_id = job.get("ingestion_run_id")
        if ingestion_run_id:
            run = self.store.get_ingestion_run(str(ingestion_run_id))
            if run:
                job["ingestion_run"] = run
        job["kb_version"] = self._kb_version
        return job

    def admin_doc_pdf(self, role: str, doc_id: str, doc_version: str) -> tuple[bytes, str]:
        ensure_role(role, {"admin", "clinician"})
        doc = next(
            (
                item
                for item in self.store.list_docs()
                if item["doc_id"] == doc_id and item["doc_version"] == doc_version
            ),
            None,
        )
        if doc is None:
            raise NotFoundError(f"Document not found: {doc_id}:{doc_version}")

        path = Path(doc["file_path"])
        if not path.exists():
            raise NotFoundError(f"Document file not found: {doc_id}:{doc_version}")
        return path.read_bytes(), path.name

    @staticmethod
    def _sanitize_session_text(value: Any, *, max_len: int = 240) -> str:
        text = str(value or "").strip()
        if not text:
            return ""
        return text[:max_len]

    @staticmethod
    def _now_epoch_sec() -> int:
        return int(time.time())

    def session_check_access(self, *, session_id: str, user_id: str, issued_at: int) -> dict[str, Any]:
        clean_session_id = self._sanitize_session_text(session_id, max_len=160)
        clean_user_id = self._sanitize_session_text(user_id, max_len=160)
        if not clean_session_id or not clean_user_id:
            return {"allowed": False, "reason": "invalid_session_identity"}

        now_epoch = self._now_epoch_sec()
        if self.store.is_session_revoked(clean_session_id, now_epoch_sec=now_epoch):
            return {"allowed": False, "reason": "session_revoked"}

        forced_after = self.store.get_forced_logout_after(clean_user_id)
        if forced_after and (int(issued_at) <= 0 or int(issued_at) < int(forced_after)):
            return {
                "allowed": False,
                "reason": "forced_logout",
                "forced_logout_after": int(forced_after),
            }

        return {
            "allowed": True,
            "reason": "ok",
            "forced_logout_after": int(forced_after) if forced_after else None,
        }

    def session_revoke(self, *, role: str, payload: dict[str, Any]) -> dict[str, Any]:
        ensure_role(role, {"admin", "clinician", "patient"})
        scope = self._sanitize_session_text(payload.get("scope"), max_len=20).lower() or "self"

        if scope == "user":
            ensure_role(role, {"admin"})
            target_user_id = self._sanitize_session_text(payload.get("user_id"), max_len=160)
            if not target_user_id:
                raise ValidationError("user_id is required for scope=user")
            actor_user_id = self._sanitize_session_text(payload.get("actor_user_id"), max_len=160)
            reason = self._sanitize_session_text(payload.get("reason"), max_len=400) or "admin_forced_logout"
            forced_after = self._now_epoch_sec()
            self.store.force_logout_user(
                user_id=target_user_id,
                forced_after_epoch=forced_after,
                updated_at=datetime.now(timezone.utc).isoformat(),
                actor_user_id=actor_user_id,
                reason=reason,
            )
            return {
                "ok": True,
                "scope": "user",
                "user_id": target_user_id,
                "forced_logout_after": forced_after,
            }

        sessions_payload = payload.get("sessions")
        sessions = sessions_payload if isinstance(sessions_payload, list) else []
        revoked_ids: set[str] = set()
        user_ids: set[str] = set()
        reason = self._sanitize_session_text(payload.get("reason"), max_len=400) or "logout"
        revoked_at = datetime.now(timezone.utc).isoformat()
        for item in sessions:
            if not isinstance(item, dict):
                continue
            session_id = self._sanitize_session_text(item.get("session_id"), max_len=160)
            user_id = self._sanitize_session_text(item.get("user_id"), max_len=160)
            if not session_id or not user_id:
                continue
            exp_raw = item.get("exp")
            exp = None
            if isinstance(exp_raw, (int, float)):
                exp = int(exp_raw)
            role_name = self._sanitize_session_text(item.get("role"), max_len=40)
            self.store.revoke_session(
                session_id=session_id,
                user_id=user_id,
                role=role_name,
                revoked_at=revoked_at,
                expires_at=exp,
                reason=reason,
            )
            revoked_ids.add(session_id)
            user_ids.add(user_id)

        return {
            "ok": True,
            "scope": "self",
            "revoked_session_ids": len(revoked_ids),
            "user_ids": sorted(user_ids),
        }

    def session_reserve_idp_jti(self, *, payload: dict[str, Any]) -> dict[str, Any]:
        jti_hash = self._sanitize_session_text(payload.get("jti_hash"), max_len=160).lower()
        user_id = self._sanitize_session_text(payload.get("user_id"), max_len=160)
        exp_raw = payload.get("exp")
        exp = 0
        if isinstance(exp_raw, (int, float)):
            exp = int(exp_raw)
        elif isinstance(exp_raw, str) and exp_raw.strip().isdigit():
            exp = int(exp_raw.strip())

        if not jti_hash or len(jti_hash) < 8:
            return {"allowed": False, "reason": "idp_jti_missing"}
        if not user_id:
            return {"allowed": False, "reason": "idp_claims_missing_identity_or_role_not_allowed"}

        now_epoch = self._now_epoch_sec()
        if exp <= now_epoch:
            return {"allowed": False, "reason": "idp_token_expired"}

        allowed = self.store.reserve_idp_token_jti(
            jti_hash=jti_hash,
            user_id=user_id,
            first_seen_epoch=now_epoch,
            expires_at=exp,
            now_epoch_sec=now_epoch,
        )
        if not allowed:
            return {"allowed": False, "reason": "idp_token_replay_detected"}
        return {"allowed": True, "reason": "ok", "expires_at": exp}

    def session_record_audit(self, *, role: str, payload: dict[str, Any]) -> dict[str, Any]:
        ensure_role(role, {"admin", "clinician", "patient"})
        self._purge_expired_session_audit_events()
        event = self._sanitize_session_text(payload.get("event"), max_len=100) or "unknown"
        outcome = self._sanitize_session_text(payload.get("outcome"), max_len=16) or "info"
        if outcome not in {"allow", "deny", "info", "error"}:
            outcome = "info"
        record = {
            "event": event,
            "outcome": outcome,
            "role": self._sanitize_session_text(payload.get("role"), max_len=40),
            "user_id": self._sanitize_session_text(payload.get("user_id"), max_len=160),
            "session_id": self._sanitize_session_text(payload.get("session_id"), max_len=160),
            "actor_user_id": self._sanitize_session_text(payload.get("actor_user_id"), max_len=160),
            "reason": self._sanitize_session_text(payload.get("reason"), max_len=400),
            "reason_group": self._session_reason_group(self._sanitize_session_text(payload.get("reason"), max_len=400)),
            "path": self._sanitize_session_text(payload.get("path"), max_len=240),
            "correlation_id": self._sanitize_session_text(payload.get("correlation_id"), max_len=120),
        }
        created_at = datetime.now(timezone.utc).isoformat()
        event_id = self.store.save_session_audit_event(record, created_at=created_at)
        return {"ok": True, "event_id": event_id, "timestamp": created_at}

    @staticmethod
    def _encode_session_audit_cursor(created_at: str, event_id: str) -> str:
        if not created_at or not event_id:
            return ""
        return f"{created_at}|{event_id}"

    @staticmethod
    def _decode_session_audit_cursor(cursor: str) -> tuple[str, str]:
        value = str(cursor or "").strip()
        if not value or "|" not in value:
            return "", ""
        created_at, event_id = value.split("|", 1)
        return created_at.strip(), event_id.strip()

    @staticmethod
    def _session_reason_group(reason: str | None) -> str:
        value = str(reason or "").strip().lower()
        if not value:
            return "other"
        if "idp_" in value or "credentials" in value or "auth" in value:
            return "auth"
        if "replay" in value or "refresh_rotation" in value:
            return "token"
        if "revoked" in value or "forced_logout" in value:
            return "revocation"
        if "config" in value or "missing" in value:
            return "config"
        if "rate_limit" in value:
            return "rate_limit"
        return "other"

    @staticmethod
    def _normalize_audit_timestamp(value: Any, *, field_name: str) -> str:
        raw = str(value or "").strip()
        if not raw:
            return ""
        candidate = raw
        if raw.endswith("Z"):
            candidate = f"{raw[:-1]}+00:00"
        try:
            parsed = datetime.fromisoformat(candidate)
        except ValueError as exc:
            raise ValidationError(f"Invalid {field_name}: expected ISO timestamp") from exc
        if parsed.tzinfo is None:
            parsed = parsed.replace(tzinfo=timezone.utc)
        return parsed.astimezone(timezone.utc).isoformat()

    def _purge_expired_session_audit_events(self) -> int:
        retention_days = int(getattr(self.settings, "session_audit_retention_days", 90) or 0)
        if retention_days <= 0:
            return 0
        cutoff = (datetime.now(timezone.utc) - timedelta(days=retention_days)).isoformat()
        return self.store.purge_session_audit_before(created_before=cutoff)

    def session_audit(
        self,
        *,
        role: str,
        limit: int = 50,
        filters: dict[str, Any] | None = None,
        cursor: str = "",
    ) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        self._purge_expired_session_audit_events()
        safe_limit = max(1, min(int(limit), 500))
        raw_filters = filters or {}
        safe_filters = {
            "outcome": self._sanitize_session_text(raw_filters.get("outcome"), max_len=16).lower(),
            "event": self._sanitize_session_text(raw_filters.get("event"), max_len=100),
            "reason": self._sanitize_session_text(raw_filters.get("reason"), max_len=200),
            "reason_group": self._sanitize_session_text(raw_filters.get("reason_group"), max_len=40).lower(),
            "user_id": self._sanitize_session_text(raw_filters.get("user_id"), max_len=160),
            "correlation_id": self._sanitize_session_text(raw_filters.get("correlation_id"), max_len=120),
            "from_ts": self._normalize_audit_timestamp(raw_filters.get("from_ts"), field_name="from_ts"),
            "to_ts": self._normalize_audit_timestamp(raw_filters.get("to_ts"), field_name="to_ts"),
        }
        if safe_filters["from_ts"] and safe_filters["to_ts"] and safe_filters["from_ts"] > safe_filters["to_ts"]:
            raise ValidationError("Invalid audit time window: from_ts must be <= to_ts")
        safe_cursor = self._sanitize_session_text(cursor, max_len=300)
        cursor_created_at, cursor_event_id = self._decode_session_audit_cursor(safe_cursor)
        events_raw = self.store.list_session_audit_events(
            limit=safe_limit + 1,
            filters=safe_filters,
            cursor_created_at=cursor_created_at,
            cursor_event_id=cursor_event_id,
        )
        has_more = len(events_raw) > safe_limit
        events = events_raw[:safe_limit]
        next_cursor = ""
        if has_more and events:
            last = events[-1]
            next_cursor = self._encode_session_audit_cursor(
                str(last.get("created_at") or ""),
                str(last.get("event_id") or ""),
            )
        for event in events:
            event["timestamp"] = event.pop("created_at", "")
            event["reason_group"] = str(event.get("reason_group") or "").strip() or self._session_reason_group(
                str(event.get("reason") or "")
            )
        revoked_sample = self.store.list_revoked_session_ids(limit=50, now_epoch_sec=self._now_epoch_sec())
        return {
            "count": len(events),
            "limit": safe_limit,
            "filters": safe_filters,
            "cursor": safe_cursor,
            "next_cursor": next_cursor,
            "events": events,
            "revoked_session_ids_sample": revoked_sample,
        }

    def session_audit_summary(
        self,
        *,
        role: str,
        window_hours: int = 24,
        from_ts: str = "",
        to_ts: str = "",
    ) -> dict[str, Any]:
        ensure_role(role, {"admin"})
        self._purge_expired_session_audit_events()

        safe_window = max(1, min(int(window_hours), 24 * 7))
        safe_from_ts = self._normalize_audit_timestamp(from_ts, field_name="from_ts")
        safe_to_ts = self._normalize_audit_timestamp(to_ts, field_name="to_ts")
        if safe_from_ts and safe_to_ts and safe_from_ts > safe_to_ts:
            raise ValidationError("Invalid audit time window: from_ts must be <= to_ts")
        if not safe_from_ts:
            safe_from_ts = (datetime.now(timezone.utc) - timedelta(hours=safe_window)).isoformat()
        if not safe_to_ts:
            safe_to_ts = datetime.now(timezone.utc).isoformat()

        summary = self.store.session_audit_summary(
            from_ts=safe_from_ts,
            to_ts=safe_to_ts,
            top_limit=10,
        )
        outcome_counts_raw = summary.get("outcome_counts") if isinstance(summary, dict) else {}
        outcome_counts = {
            "allow": int((outcome_counts_raw or {}).get("allow") or 0),
            "deny": int((outcome_counts_raw or {}).get("deny") or 0),
            "info": int((outcome_counts_raw or {}).get("info") or 0),
            "error": int((outcome_counts_raw or {}).get("error") or 0),
        }
        total_events = int(summary.get("total_events") or 0)
        deny_count = int(outcome_counts.get("deny") or 0)
        error_count = int(outcome_counts.get("error") or 0)
        deny_rate = float(deny_count / total_events) if total_events > 0 else 0.0

        tracked_reason_counts = self.store.session_audit_reason_counts(
            reasons=[
                "idp_token_replay_detected",
                "idp_config_incomplete",
                "credentials_mode_without_users",
            ],
            from_ts=safe_from_ts,
            to_ts=safe_to_ts,
        )
        replay_detected_count = int(tracked_reason_counts.get("idp_token_replay_detected") or 0)
        config_error_count = int(tracked_reason_counts.get("idp_config_incomplete") or 0) + int(
            tracked_reason_counts.get("credentials_mode_without_users") or 0
        )

        min_events = max(1, int(getattr(self.settings, "session_audit_alert_min_events", 10) or 10))
        deny_warn = float(getattr(self.settings, "session_audit_alert_deny_rate_warn", 0.35) or 0.35)
        deny_critical = float(getattr(self.settings, "session_audit_alert_deny_rate_critical", 0.60) or 0.60)
        deny_warn = max(0.0, min(deny_warn, 1.0))
        deny_critical = max(deny_warn, min(deny_critical, 1.0))

        error_warn = max(1, int(getattr(self.settings, "session_audit_alert_error_count_warn", 5) or 5))
        error_critical = max(
            error_warn,
            int(getattr(self.settings, "session_audit_alert_error_count_critical", 20) or 20),
        )

        replay_warn = max(1, int(getattr(self.settings, "session_audit_alert_replay_count_warn", 1) or 1))
        replay_critical = max(
            replay_warn,
            int(getattr(self.settings, "session_audit_alert_replay_count_critical", 3) or 3),
        )

        config_warn = max(1, int(getattr(self.settings, "session_audit_alert_config_error_count_warn", 1) or 1))
        config_critical = max(
            config_warn,
            int(getattr(self.settings, "session_audit_alert_config_error_count_critical", 3) or 3),
        )

        alerts: list[dict[str, Any]] = []

        def add_alert(code: str, metric: str, level: str, value: float | int, threshold: float | int, message: str) -> None:
            alerts.append(
                {
                    "code": code,
                    "metric": metric,
                    "level": level,
                    "value": value,
                    "threshold": threshold,
                    "message": message,
                }
            )

        if total_events >= min_events:
            if deny_rate >= deny_critical:
                add_alert(
                    "deny_rate_exceeded",
                    "deny_rate",
                    "critical",
                    round(deny_rate, 4),
                    deny_critical,
                    "High deny-rate in session auth traffic.",
                )
            elif deny_rate >= deny_warn:
                add_alert(
                    "deny_rate_exceeded",
                    "deny_rate",
                    "warn",
                    round(deny_rate, 4),
                    deny_warn,
                    "Elevated deny-rate in session auth traffic.",
                )

        if error_count >= error_critical:
            add_alert(
                "error_count_exceeded",
                "error_count",
                "critical",
                error_count,
                error_critical,
                "High volume of session auth errors.",
            )
        elif error_count >= error_warn:
            add_alert(
                "error_count_exceeded",
                "error_count",
                "warn",
                error_count,
                error_warn,
                "Elevated volume of session auth errors.",
            )

        if replay_detected_count >= replay_critical:
            add_alert(
                "replay_detected_count_exceeded",
                "replay_detected_count",
                "critical",
                replay_detected_count,
                replay_critical,
                "Multiple replay-detection events found.",
            )
        elif replay_detected_count >= replay_warn:
            add_alert(
                "replay_detected_count_exceeded",
                "replay_detected_count",
                "warn",
                replay_detected_count,
                replay_warn,
                "Replay-detection event found.",
            )

        if config_error_count >= config_critical:
            add_alert(
                "config_error_count_exceeded",
                "config_error_count",
                "critical",
                config_error_count,
                config_critical,
                "Session auth configuration errors detected.",
            )
        elif config_error_count >= config_warn:
            add_alert(
                "config_error_count_exceeded",
                "config_error_count",
                "warn",
                config_error_count,
                config_warn,
                "Session auth configuration warning detected.",
            )

        critical_alerts = sum(1 for item in alerts if str(item.get("level") or "") == "critical")
        warn_alerts = sum(1 for item in alerts if str(item.get("level") or "") == "warn")
        incident_score = min(100, critical_alerts * 50 + warn_alerts * 25)
        if critical_alerts > 0:
            incident_level = "high"
        elif warn_alerts >= 2:
            incident_level = "medium"
        elif warn_alerts == 1:
            incident_level = "low"
        else:
            incident_level = "none"

        return {
            "window_hours": safe_window,
            "from_ts": safe_from_ts,
            "to_ts": safe_to_ts,
            "total_events": total_events,
            "unique_users": int(summary.get("unique_users") or 0),
            "outcome_counts": outcome_counts,
            "reason_group_counts": summary.get("reason_group_counts") or {},
            "top_reasons": summary.get("top_reasons") or [],
            "top_events": summary.get("top_events") or [],
            "incident_level": incident_level,
            "incident_score": incident_score,
            "incident_signals": {
                "deny_rate": round(deny_rate, 4),
                "error_count": error_count,
                "replay_detected_count": replay_detected_count,
                "config_error_count": config_error_count,
                "min_events_for_deny_rate_alert": min_events,
            },
            "alerts": alerts,
        }

    def report_json(self, role: str, report_id: str) -> dict[str, Any]:
        ensure_role(role, {"admin", "clinician"})
        payload = self.store.get_report(report_id)
        if not payload:
            raise NotFoundError(f"Report not found: {report_id}")
        return payload

    @staticmethod
    def _build_report_export_sections(payload: dict[str, Any]) -> list[tuple[str, list[str]]]:
        doctor_report = payload.get("doctor_report") if isinstance(payload.get("doctor_report"), dict) else {}
        sections: list[tuple[str, list[str]]] = []

        summary = str(doctor_report.get("summary_md") or doctor_report.get("summary") or doctor_report.get("consilium_md") or "").strip()
        sections.append(("Краткое резюме", [summary or "Резюме не сформировано."]))

        timeline_lines: list[str] = []
        timeline = doctor_report.get("timeline") if isinstance(doctor_report.get("timeline"), list) else []
        for item in timeline[:20]:
            if isinstance(item, dict):
                date_part = str(item.get("date") or "").strip()
                label = str(item.get("label") or item.get("details") or "").strip()
                if label:
                    timeline_lines.append(f"{date_part}: {label}" if date_part else label)
            elif isinstance(item, str) and item.strip():
                timeline_lines.append(item.strip())
        if timeline_lines:
            sections.append(("Таймлайн", timeline_lines))

        plan_lines: list[str] = []
        plan = doctor_report.get("plan") if isinstance(doctor_report.get("plan"), list) else []
        for section in plan:
            if not isinstance(section, dict):
                continue
            section_title = str(section.get("title") or section.get("section") or "План").strip()
            if section_title:
                plan_lines.append(f"[{section_title}]")
            steps = section.get("steps") if isinstance(section.get("steps"), list) else []
            for step in steps:
                if not isinstance(step, dict):
                    continue
                text = str(step.get("text") or "").strip()
                if text:
                    plan_lines.append(f"- {text}")
        sections.append(("План действий", plan_lines or ["План действий отсутствует."]))

        issue_lines: list[str] = []
        issues = doctor_report.get("issues") if isinstance(doctor_report.get("issues"), list) else []
        for issue in issues[:30]:
            if not isinstance(issue, dict):
                continue
            severity = str(issue.get("severity") or "info").upper()
            summary_line = str(issue.get("summary") or "").strip()
            details = str(issue.get("details") or "").strip()
            if summary_line:
                line = f"[{severity}] {summary_line}"
                if details:
                    line += f" — {details}"
                issue_lines.append(line)
        sections.append(("Клинические замечания", issue_lines or ["Существенных замечаний не выявлено."]))

        citation_lines: list[str] = []
        citations = doctor_report.get("citations") if isinstance(doctor_report.get("citations"), list) else []
        for citation in citations[:30]:
            if not isinstance(citation, dict):
                continue
            source = str(citation.get("source_id") or "unknown").strip()
            section_path = str(citation.get("section_path") or "").strip()
            quote = str(citation.get("quote") or "").strip()
            line = f"{source}: {section_path}" if section_path else source
            if quote:
                line = f"{line} — {quote}"
            citation_lines.append(line)
        sections.append(("Цитаты и источники", citation_lines or ["Цитаты не найдены."]))
        return sections

    @staticmethod
    def _render_minimal_pdf_bytes(title: str, sections: list[tuple[str, list[str]]]) -> bytes:
        # Deterministic single-page text PDF for environments without external rendering engines.
        def esc(value: str) -> str:
            return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

        lines: list[str] = [title]
        for section_title, section_lines in sections:
            lines.append("")
            lines.append(section_title)
            lines.extend(section_lines if section_lines else ["(пусто)"])

        content_ops = ["BT", "/F1 11 Tf", "40 800 Td"]
        first = True
        for line in lines[:55]:
            safe_line = esc(line[:180])
            if first:
                content_ops.append(f"({safe_line}) Tj")
                first = False
            else:
                content_ops.append("0 -14 Td")
                content_ops.append(f"({safe_line}) Tj")
        content_ops.append("ET")
        stream = "\n".join(content_ops).encode("latin-1", errors="replace")

        objects: list[bytes] = []
        objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
        objects.append(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
        objects.append(b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>")
        objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
        objects.append(f"<< /Length {len(stream)} >>\nstream\n".encode("latin-1") + stream + b"\nendstream")

        chunks: list[bytes] = [b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"]
        offsets = [0]
        for index, obj in enumerate(objects, start=1):
            offsets.append(sum(len(chunk) for chunk in chunks))
            chunks.append(f"{index} 0 obj\n".encode("latin-1") + obj + b"\nendobj\n")
        xref_offset = sum(len(chunk) for chunk in chunks)
        xref_lines = [b"xref", f"0 {len(objects) + 1}".encode("latin-1"), b"0000000000 65535 f "]
        for offset in offsets[1:]:
            xref_lines.append(f"{offset:010d} 00000 n ".encode("latin-1"))
        trailer = (
            b"trailer\n"
            + f"<< /Size {len(objects) + 1} /Root 1 0 R >>\n".encode("latin-1")
            + b"startxref\n"
            + f"{xref_offset}\n".encode("latin-1")
            + b"%%EOF\n"
        )
        chunks.append(b"\n".join(xref_lines) + b"\n" + trailer)
        return b"".join(chunks)

    @staticmethod
    def _render_minimal_docx_bytes(title: str, sections: list[tuple[str, list[str]]]) -> bytes:
        def xml_escape(value: str) -> str:
            return (
                value.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )

        paragraphs: list[str] = [title]
        for section_title, lines in sections:
            paragraphs.append(section_title)
            paragraphs.extend(lines if lines else ["(пусто)"])

        body_xml = "".join(
            (
                "<w:p><w:r><w:t>"
                + xml_escape(paragraph[:500])
                + "</w:t></w:r></w:p>"
            )
            for paragraph in paragraphs
        )
        document_xml = (
            "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
            "<w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
            f"<w:body>{body_xml}<w:sectPr/></w:body></w:document>"
        )
        content_types = (
            "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
            "<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">"
            "<Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>"
            "<Default Extension=\"xml\" ContentType=\"application/xml\"/>"
            "<Override PartName=\"/word/document.xml\" "
            "ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml\"/>"
            "</Types>"
        )
        rels_xml = (
            "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
            "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
            "<Relationship Id=\"rId1\" "
            "Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" "
            "Target=\"word/document.xml\"/>"
            "</Relationships>"
        )

        buffer = BytesIO()
        with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", content_types)
            archive.writestr("_rels/.rels", rels_xml)
            archive.writestr("word/document.xml", document_xml)
        return buffer.getvalue()

    def report_pdf(self, role: str, report_id: str) -> bytes:
        payload = self.report_json(role, report_id)
        sections = self._build_report_export_sections(payload)
        title = "ОнкоНавигатор: консилиумный пакет"
        return self._render_minimal_pdf_bytes(title=title, sections=sections)

    def report_docx(self, role: str, report_id: str) -> bytes:
        payload = self.report_json(role, report_id)
        sections = self._build_report_export_sections(payload)
        try:
            from docx import Document
        except ModuleNotFoundError as exc:  # pragma: no cover
            _ = exc
            return self._render_minimal_docx_bytes("ОнкоНавигатор: консилиумный пакет", sections)

        doc = Document()
        doc.add_heading("ОнкоНавигатор: консилиумный пакет", level=1)
        for title, lines in sections:
            doc.add_heading(title, level=2)
            for line in lines:
                doc.add_paragraph(str(line))

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def report_html(self, role: str, report_id: str) -> str:
        payload = self.report_json(role, report_id)
        report = payload["doctor_report"]

        if report.get("schema_version") == "1.0":
            issue_html = "".join(
                (
                    f"<li><strong>{issue['severity'].upper()}:</strong> {issue['summary']}"
                    f"<br/><small>{issue.get('details', '')}</small></li>"
                )
                for issue in report.get("issues", [])
            )
            return (
                "<!DOCTYPE html><html><head><meta charset='utf-8'/>"
                "<title>OncoAI Report</title>"
                "<style>body{font-family:Arial,Helvetica,sans-serif;margin:24px;}"
                "h1{margin-bottom:4px;} .meta{color:#666;} li{margin:8px 0;}</style></head><body>"
                f"<h1>Doctor Report</h1><p class='meta'>Report ID: {report['report_id']} | Request: {payload.get('request_id','')}</p>"
                f"<p>{report.get('summary_md', '')}</p><ul>{issue_html}</ul></body></html>"
            )

        issue_html = "".join(
            (
                f"<li><strong>{str(issue.get('severity') or 'info').upper()}:</strong> "
                f"{issue.get('title') or issue.get('summary') or ''}"
                f"<br/><small>{issue.get('description') or issue.get('details') or ''}</small></li>"
            )
            for issue in report.get("issues", [])
        )
        run_meta = payload.get("run_meta") if isinstance(payload.get("run_meta"), dict) else {}
        report_id_value = str(report.get("report_id") or report_id)
        kb_version_value = str(report.get("kb_version") or run_meta.get("kb_version") or "unknown")
        summary_text = str(report.get("summary") or report.get("summary_md") or report.get("consilium_md") or "")
        return (
            "<!DOCTYPE html><html><head><meta charset='utf-8'/>"
            "<title>OncoAI Report</title>"
            "<style>body{font-family:Arial,Helvetica,sans-serif;margin:24px;}"
            "h1{margin-bottom:4px;} .meta{color:#666;} li{margin:8px 0;}</style></head><body>"
            f"<h1>Doctor Report</h1><p class='meta'>Report ID: {report_id_value} | KB: {kb_version_value}</p>"
            f"<p>{summary_text}</p><ul>{issue_html}</ul></body></html>"
        )
